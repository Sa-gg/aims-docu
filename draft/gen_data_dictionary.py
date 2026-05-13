#!/usr/bin/env python3
"""Generate AIMS Data Dictionary DOCX matching the academic thesis format."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def simplify(dtype):
    d = dtype.upper()
    if "VARCHAR(36)" in d:  return "varchar", "36"
    if "VARCHAR(50)" in d:  return "varchar", "50"
    if "VARCHAR"     in d:  return "varchar", "255"
    if "TEXT"        in d:  return "text",    "-"
    if "BOOLEAN"     in d:  return "tinyint", "1"
    if "INTEGER"     in d:  return "int",     "11"
    if "DECIMAL"     in d:  return "decimal", "-"
    if "TIMESTAMP"   in d:  return "datetime","-"
    if "JSONB"       in d:  return "json",    "-"
    if "ENUM"        in d:  return "varchar", "50"
    return dtype.lower(), "-"

def key_type(k):
    if "PK" in k:  return "PK"
    if "FK" in k:  return "FK"
    if "UQ" in k:  return "UQ"
    return ""

def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcB  = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "000000")
                tcB.append(el)
            tcPr.append(tcB)

def fix_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def cell_write(cell, text, bold=False, italic=False,
               red=False, underline=False, size=10,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size = Pt(size)
    if red:
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "80")
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)

COL_W = [3.5, 2.3, 1.7, 8.42]

TABLES = [
    {
        "table": "users", "model": "User",
        "explanation": "Table 1 stores the complete profile and authentication credentials for every system user. It serves as the central identity record supporting three account types: students, teachers, and administrators. Authentication may be performed using a local email-and-password pair, a numeric DepEd employee ID (for teachers synced from EnrollPro), or a linked Google account via OAuth 2.0. Teachers synced from the EnrollPro integration have their enrollproId and employeeId populated for cross-system identification. Each user is scoped to a single school through the schoolId foreign key, enforcing the multi-tenant data isolation of the platform.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique user identifier (UUID v4)"),
            ("email","VARCHAR","UQ","NO","—","Unique email address used for authentication"),
            ("emailVerified","BOOLEAN","","NO","false","Whether the email address has been verified"),
            ("emailVerifyToken","VARCHAR","UQ","YES","NULL","One-time token sent for email verification"),
            ("emailVerifyExpiry","TIMESTAMP","","YES","NULL","Expiry datetime of the email verification token"),
            ("passwordHash","VARCHAR","","NO","—","Bcrypt-hashed local password"),
            ("firstName","VARCHAR","","NO","—","User first name"),
            ("lastName","VARCHAR","","NO","—","User last name"),
            ("role","VARCHAR(50)","","NO","STUDENT","Access role: STUDENT, TEACHER, or ADMIN"),
            ("googleSub","VARCHAR","UQ","YES","NULL","Google OAuth unique subject identifier"),
            ("googleEmail","VARCHAR","","YES","NULL","Email from the linked Google account"),
            ("googleName","VARCHAR","","YES","NULL","Display name from the linked Google account"),
            ("googlePicture","VARCHAR","","YES","NULL","Profile photo URL from the linked Google account"),
            ("googleAccessToken","VARCHAR","","YES","NULL","OAuth2 access token for Google API calls"),
            ("googleRefreshToken","VARCHAR","","YES","NULL","OAuth2 refresh token for renewing access tokens"),
            ("googleTokenExpiry","TIMESTAMP","","YES","NULL","Expiry datetime of the current Google access token"),
            ("driveConnected","BOOLEAN","","NO","false","Whether Google Drive integration is active"),
            ("driveFolderId","VARCHAR","","YES","NULL","Root Google Drive folder ID for user file storage"),
            ("fileStorageBackend","VARCHAR(50)","","NO","LOCAL","Storage backend in use: LOCAL or GDRIVE"),
            ("schoolId","VARCHAR(36)","FK->schools","YES","NULL","Foreign key referencing the user's school"),
            ("employeeId","VARCHAR","UQ","YES","NULL","DepEd employee ID (teachers only, from EnrollPro)"),
            ("enrollproId","INTEGER","","YES","NULL","EnrollPro internal teacher record ID"),
            ("enrollproSyncedAt","TIMESTAMP","","YES","NULL","Timestamp of the last successful EnrollPro sync"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
            ("updatedAt","TIMESTAMP","","NO","auto","Timestamp of the last record update"),
        ],
    },
    {
        "table": "refresh_tokens", "model": "RefreshToken",
        "explanation": "Table 2 stores the JWT refresh tokens issued to users upon successful login. Each record is linked to a specific user and carries an expiry datetime. Tokens may be explicitly revoked on logout or when a password change is performed, enabling fine-grained session management. This design allows users to maintain persistent login sessions across multiple devices while still providing a mechanism for immediate session invalidation when required.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique token record identifier"),
            ("token","VARCHAR","UQ","NO","—","JWT refresh token string"),
            ("userId","VARCHAR(36)","FK->users","NO","—","Foreign key to the token owner"),
            ("expiresAt","TIMESTAMP","","NO","—","Expiry datetime of the token (7 days from issuance)"),
            ("createdAt","TIMESTAMP","","NO","now()","Token creation timestamp"),
            ("revoked","BOOLEAN","","NO","false","Whether the token has been revoked"),
        ],
    },
    {
        "table": "schools", "model": "School",
        "explanation": "Table 3 stores the tenant records for each school using the AIMS platform. Every piece of user, course, and academic data in the system is scoped to a school record to enforce strict multi-tenant data isolation. Each school record contains branding configuration including primary, secondary, and accent colors along with a logo URL, which are applied throughout the user interface for that tenant. Contact and address information are also stored for administrative reference.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique school identifier"),
            ("name","VARCHAR","","NO","—","Full school name"),
            ("shortName","VARCHAR","","NO","—","Abbreviated school name used in UI headers"),
            ("logoUrl","VARCHAR","","YES","NULL","URL of the school logo image"),
            ("primaryColor","VARCHAR","","NO","#1d4ed8","Brand primary hex color"),
            ("secondaryColor","VARCHAR","","NO","#7c3aed","Brand secondary hex color"),
            ("accentColor","VARCHAR","","NO","#0891b2","Brand accent hex color"),
            ("address","TEXT","","YES","NULL","Physical address of the school"),
            ("contactEmail","VARCHAR","","YES","NULL","School contact email address"),
            ("isActive","BOOLEAN","","NO","true","Whether the school tenant is currently active"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
            ("updatedAt","TIMESTAMP","","NO","auto","Timestamp of the last record update"),
        ],
    },
    {
        "table": "courses", "model": "Course",
        "explanation": "Table 4 stores the subject classes created by teachers. Each course has a unique enrollment code that students use to join. Courses support DepEd grading categories, configurable mastery-lock gating that prevents students from accessing future quizzes until they pass prerequisite ones, and AI-powered remedial quiz generation. The course belongs to a school and a teacher, and aggregates all quizzes, tasks, resources, topics, and student enrollments associated with that class.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique course identifier"),
            ("name","VARCHAR","","NO","—","Course display name"),
            ("description","TEXT","","YES","NULL","Optional course description"),
            ("code","VARCHAR","UQ","NO","—","Unique join code for student enrollment"),
            ("subject","VARCHAR","","NO","General","Subject label (e.g., English, Math)"),
            ("subjectType","VARCHAR(50)","","NO","TLE","DepEd subject category classification"),
            ("gradeLevel","VARCHAR","","NO","Grade 11","Target grade level (e.g., Grade 10)"),
            ("schoolYear","VARCHAR","","NO","2024-2025","Academic school year label"),
            ("color","VARCHAR","","NO","#2563eb","Hex color used for the course card display"),
            ("softLockEnabled","BOOLEAN","","NO","false","Whether mastery soft-lock gating is active"),
            ("passingThreshold","INTEGER","","NO","70","Minimum percentage score required to pass quizzes"),
            ("masteryLockMode","VARCHAR(50)","","NO","PROSPECTIVE","Mastery enforcement mode: PROSPECTIVE or RETROACTIVE"),
            ("masteryLockActivatedAt","TIMESTAMP","","YES","NULL","Datetime mastery lock was first activated"),
            ("archived","BOOLEAN","","NO","false","Whether the course has been archived"),
            ("remedialDifficulty","VARCHAR","","NO","MEDIUM","Default difficulty for AI-generated remedial quizzes"),
            ("remedialNumQuestions","INTEGER","","YES","NULL","Override for number of questions in remedial quizzes"),
            ("schoolId","VARCHAR(36)","FK->schools","YES","NULL","Foreign key to the owning school"),
            ("teacherId","VARCHAR(36)","FK->users","YES","NULL","Foreign key to the teacher who owns the course"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
            ("updatedAt","TIMESTAMP","","NO","auto","Timestamp of the last record update"),
        ],
    },
    {
        "table": "enrollments", "model": "Enrollment",
        "explanation": "Table 5 is a junction table that records which students are enrolled in which courses. A student may be enrolled in multiple courses and a course may have many enrolled students. A unique constraint on the (userId, courseId) pair prevents duplicate enrollment records. Archived enrollments are soft-deleted rather than permanently removed, preserving historical grading data while hiding the student from active course rosters.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique enrollment record identifier"),
            ("userId","VARCHAR(36)","FK->users","NO","—","Foreign key to the enrolled student"),
            ("courseId","VARCHAR(36)","FK->courses","NO","—","Foreign key to the course enrolled in"),
            ("archived","BOOLEAN","","NO","false","Whether this enrollment has been archived"),
            ("printedCopy","BOOLEAN","","NO","false","Whether the student's grade sheet has been printed"),
            ("createdAt","TIMESTAMP","","NO","now()","Enrollment timestamp"),
        ],
    },
    {
        "table": "topics", "model": "Topic",
        "explanation": "Table 6 allows teachers to organize course content into named topic or unit groupings. Quizzes, tasks, and resources can be tagged to a topic so that students see them grouped by chapter or theme rather than in a flat list. Each topic has a configurable display order within its parent course, and topic names are unique within a course to prevent duplication.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","CUID()","Unique topic identifier"),
            ("courseId","VARCHAR(36)","FK->courses","NO","—","Foreign key to the parent course"),
            ("name","VARCHAR","","NO","—","Topic name (unique within a course)"),
            ("orderIndex","INTEGER","","NO","0","Display sort order within the course"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
        ],
    },
    {
        "table": "resources", "model": "Resource",
        "explanation": "Table 7 stores the learning materials that teachers upload or link for student access. Resources may be uploaded files (PDF, DOCX, TXT), external hyperlinks, or YouTube video embeds. Files are stored either locally on the server or on Google Drive, depending on the teacher's configured storage backend. Resources support scheduling for deferred publication and may be linked to a specific topic or quiz to provide contextual reading material before or during an assessment.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique resource identifier"),
            ("title","VARCHAR","","NO","—","Resource display title"),
            ("description","TEXT","","YES","NULL","Optional description of the resource"),
            ("topicId","VARCHAR(36)","FK->topics","YES","NULL","Foreign key to the parent topic"),
            ("courseId","VARCHAR(36)","FK->courses","YES","NULL","Foreign key to the parent course"),
            ("teacherId","VARCHAR(36)","FK->users","YES","NULL","Foreign key to the owning teacher"),
            ("uploadedById","VARCHAR(36)","FK->users","NO","—","Foreign key to the user who uploaded the resource"),
            ("fileType","VARCHAR(50)","","YES","NULL","File type classification: PDF, DOCX, or TXT"),
            ("fileName","VARCHAR","","YES","NULL","Original uploaded file name"),
            ("storedName","VARCHAR","","YES","NULL","File name as stored on disk or in storage"),
            ("filePath","VARCHAR","","YES","NULL","Relative storage path for local files"),
            ("fileSize","INTEGER","","YES","NULL","File size in bytes"),
            ("mimeType","VARCHAR","","YES","NULL","MIME type of the uploaded file"),
            ("storageKey","VARCHAR","","YES","NULL","Storage provider object key"),
            ("storageUrl","VARCHAR","","YES","NULL","Public or signed URL of the stored file"),
            ("storageBackend","VARCHAR(50)","","NO","LOCAL","Storage backend used: LOCAL or GDRIVE"),
            ("linkUrl","VARCHAR","","YES","NULL","External hyperlink URL (link-type resources)"),
            ("youtubeUrl","VARCHAR","","YES","NULL","YouTube video URL (video-type resources)"),
            ("status","VARCHAR(50)","","NO","PUBLISHED","Publication status: DRAFT, SCHEDULED, or PUBLISHED"),
            ("scheduledAt","TIMESTAMP","","YES","NULL","Datetime to auto-publish a scheduled resource"),
            ("uploadedAt","TIMESTAMP","","NO","now()","Datetime the file was first uploaded"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
        ],
    },
    {
        "table": "resource_attachments", "model": "ResourceAttachment",
        "explanation": "Table 8 stores supplementary files or links attached to a resource record. Unlike the primary resource file, attachments are additional supporting materials added to enrich the resource. Attachments may be uploaded files from local storage or Google Drive, external hyperlinks, or YouTube video links. Each attachment is uniquely associated with its parent resource record.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique attachment identifier"),
            ("resourceId","VARCHAR(36)","FK->resources","NO","—","Foreign key to the parent resource"),
            ("kind","VARCHAR(50)","","NO","FILE","Attachment kind: FILE, LINK, or YOUTUBE"),
            ("fileName","VARCHAR","","YES","NULL","Original file name"),
            ("storedName","VARCHAR","","YES","NULL","Stored file name on disk"),
            ("filePath","VARCHAR","","YES","NULL","Relative path of the stored file"),
            ("fileSize","INTEGER","","YES","NULL","File size in bytes"),
            ("mimeType","VARCHAR","","YES","NULL","MIME type of the file"),
            ("storageKey","VARCHAR","","YES","NULL","Storage provider object key"),
            ("storageUrl","VARCHAR","","YES","NULL","URL of the stored file"),
            ("storageBackend","VARCHAR(50)","","NO","LOCAL","Storage backend: LOCAL or GDRIVE"),
            ("driveFileId","VARCHAR","","YES","NULL","Google Drive file ID"),
            ("webViewLink","VARCHAR","","YES","NULL","Google Drive web view URL"),
            ("linkUrl","VARCHAR","","YES","NULL","External hyperlink URL"),
            ("youtubeUrl","VARCHAR","","YES","NULL","YouTube video URL"),
            ("createdAt","TIMESTAMP","","NO","now()","Attachment creation timestamp"),
        ],
    },
    {
        "table": "resource_deployments", "model": "ResourceDeployment",
        "explanation": "Table 9 tracks when a library-level resource has been deployed (copied) into a specific course. This enables teachers to build a shared resource library and reuse materials across multiple courses without duplication of content. The deployedAt timestamp records when the deployment occurred, and the courseResourceId is unique to ensure a library resource is deployed to a given course only once.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique deployment record identifier"),
            ("libraryResourceId","VARCHAR(36)","FK->resources","NO","—","Foreign key to the original library resource"),
            ("courseResourceId","VARCHAR(36)","FK->resources","NO","—","Foreign key to the course-scoped resource copy (unique)"),
            ("deployedAt","TIMESTAMP","","NO","now()","Timestamp when the resource was deployed"),
        ],
    },
    {
        "table": "quizzes", "model": "Quiz",
        "explanation": "Table 10 stores all quiz records in the system. Quizzes may be manually authored by a teacher or AI-generated as remedial assessments for students who did not meet the passing threshold on a previous quiz. Each quiz supports a configurable time limit, due date, point value, and DepEd grading category. The mastery-lock feature uses the passingThreshold value to determine whether a student may proceed to subsequent quizzes in the course. Remedial quizzes carry a forStudentId and sourceQuizId that link them back to the target student and the original quiz that triggered generation.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique quiz identifier"),
            ("title","VARCHAR","","NO","—","Quiz title"),
            ("description","TEXT","","YES","NULL","Optional quiz description or instructions"),
            ("courseId","VARCHAR(36)","FK->courses","YES","NULL","Foreign key to the parent course"),
            ("teacherId","VARCHAR(36)","FK->users","NO","—","Foreign key to the teacher who created the quiz"),
            ("topicId","VARCHAR(36)","FK->topics","YES","NULL","Foreign key to the parent topic"),
            ("resourceId","VARCHAR(36)","FK->resources","YES","NULL","Resource linked to this quiz"),
            ("sourceResourceId","VARCHAR(36)","FK->resources","YES","NULL","Source resource used for AI quiz generation"),
            ("orderIndex","INTEGER","","NO","0","Sort order within the course or topic"),
            ("status","VARCHAR(50)","","NO","DRAFT","Quiz status: DRAFT, PUBLISHED, CLOSED, or ARCHIVED"),
            ("scheduledAt","TIMESTAMP","","YES","NULL","Datetime to automatically publish the quiz"),
            ("dueDate","TIMESTAMP","","YES","NULL","Submission due date"),
            ("closeBeforeDue","BOOLEAN","","NO","false","Whether the quiz auto-closes at the due date"),
            ("points","INTEGER","","NO","0","Total point value of the quiz"),
            ("passingScore","INTEGER","","NO","70","Minimum score (%) required to pass the quiz"),
            ("passingThreshold","INTEGER","","NO","70","Threshold used for mastery-lock evaluation"),
            ("showScore","BOOLEAN","","NO","true","Whether students can see their score after submission"),
            ("showAnswers","BOOLEAN","","NO","true","Whether correct answers are revealed after submission"),
            ("timeLimit","INTEGER","","YES","NULL","Time limit in minutes (NULL means unlimited)"),
            ("isRemedial","BOOLEAN","","NO","false","Whether this is an AI-generated remedial quiz"),
            ("forStudentId","VARCHAR(36)","FK->users","YES","NULL","Target student for a personalised remedial quiz"),
            ("sourceQuizId","VARCHAR(36)","FK->quizzes","YES","NULL","Source quiz this remedial was generated from"),
            ("taskType","VARCHAR(50)","","NO","ASSIGNMENT","DepEd grading category for this quiz"),
            ("gradingWeightId","VARCHAR(36)","FK->grading_weights","YES","NULL","Foreign key to the grading weight category"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
            ("updatedAt","TIMESTAMP","","NO","auto","Timestamp of the last record update"),
        ],
    },
    {
        "table": "quiz_attachments", "model": "QuizAttachment",
        "explanation": "Table 11 stores files or links attached to a quiz by the teacher to provide reference materials during the assessment. Attachments may be uploaded documents, external hyperlinks, or YouTube videos, and may be stored locally or on Google Drive. This allows teachers to include context-setting materials or data tables that students need to view in order to answer the quiz questions.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique attachment identifier"),
            ("quizId","VARCHAR(36)","FK->quizzes","NO","—","Foreign key to the parent quiz"),
            ("kind","VARCHAR(50)","","NO","FILE","Attachment kind: FILE, LINK, or YOUTUBE"),
            ("fileName","VARCHAR","","YES","NULL","Original file name"),
            ("storedName","VARCHAR","","YES","NULL","Stored file name on disk"),
            ("filePath","VARCHAR","","YES","NULL","Relative path of the stored file"),
            ("fileSize","INTEGER","","YES","NULL","File size in bytes"),
            ("mimeType","VARCHAR","","YES","NULL","MIME type of the file"),
            ("storageKey","VARCHAR","","YES","NULL","Storage provider object key"),
            ("storageUrl","VARCHAR","","YES","NULL","URL of the stored file"),
            ("storageBackend","VARCHAR(50)","","NO","LOCAL","Storage backend: LOCAL or GDRIVE"),
            ("driveFileId","VARCHAR","","YES","NULL","Google Drive file ID"),
            ("webViewLink","VARCHAR","","YES","NULL","Google Drive web view URL"),
            ("linkUrl","VARCHAR","","YES","NULL","External hyperlink URL"),
            ("youtubeUrl","VARCHAR","","YES","NULL","YouTube video URL"),
            ("createdAt","TIMESTAMP","","NO","now()","Attachment creation timestamp"),
        ],
    },
    {
        "table": "questions", "model": "Question",
        "explanation": "Table 12 stores the individual questions that belong to each quiz. The system supports five question types: multiple choice, true/false, short answer, identification, and enumeration. Multiple-choice questions store their options in the options JSON field. Questions carry optional metadata such as difficulty level and a skill tag used by the AI engine when generating remedial quizzes. Questions that require manual teacher grading (open-ended types) have the requiresTeacherReview flag set to true.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique question identifier"),
            ("quizId","VARCHAR(36)","FK->quizzes","NO","—","Foreign key to the parent quiz"),
            ("questionText","TEXT","","NO","—","Question prompt text"),
            ("questionType","VARCHAR(50)","","NO","—","Type: MULTIPLE_CHOICE, TRUE_FALSE, SHORT_ANSWER, IDENTIFICATION, ENUMERATION"),
            ("options","json","","YES","NULL","Answer choices array (for multiple choice / true-false)"),
            ("correctAnswer","VARCHAR","","NO","—","Correct answer string or choice key"),
            ("modelAnswer","TEXT","","YES","NULL","Model answer or rubric for open-ended questions"),
            ("explanation","TEXT","","YES","NULL","Explanation shown to the student after submission"),
            ("skillTag","VARCHAR","","YES","NULL","Learning skill or competency tag"),
            ("difficulty","VARCHAR","","YES","NULL","Difficulty level: EASY, MEDIUM, or HARD"),
            ("requiresTeacherReview","BOOLEAN","","NO","false","Whether the question requires manual teacher grading"),
            ("points","INTEGER","","NO","1","Point value of the question"),
            ("orderIndex","INTEGER","","NO","—","Display order within the quiz"),
        ],
    },
    {
        "table": "submissions", "model": "Submission",
        "explanation": "Table 13 records each student's quiz attempt, tracking the full lifecycle from the moment the student opens the quiz to grading and score return. The score is stored as a decimal percentage computed upon grading. A submission with a MISSING status indicates the student did not submit by the due date, while EXCUSED means the teacher has waived the requirement. Each quiz-student pair results in at most one active submission record.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique submission identifier"),
            ("userId","VARCHAR(36)","FK->users","NO","—","Foreign key to the student who made the submission"),
            ("quizId","VARCHAR(36)","FK->quizzes","NO","—","Foreign key to the submitted quiz"),
            ("score","decimal","","YES","NULL","Computed percentage score"),
            ("status","VARCHAR(50)","","NO","IN_PROGRESS","Status: IN_PROGRESS, SUBMITTED, GRADED, RETURNED, MISSING, or EXCUSED"),
            ("startedAt","TIMESTAMP","","NO","now()","Datetime the student started the quiz"),
            ("submittedAt","TIMESTAMP","","YES","NULL","Datetime the student submitted the quiz"),
            ("gradedAt","TIMESTAMP","","YES","NULL","Datetime the submission was graded"),
            ("returnedAt","TIMESTAMP","","YES","NULL","Datetime the graded submission was returned to the student"),
        ],
    },
    {
        "table": "answers", "model": "Answer",
        "explanation": "Table 14 stores the individual answers that a student provided for each question in a quiz submission. The isCorrect flag is populated automatically for objective question types (multiple choice, true/false, identification) and remains NULL for questions that require manual teacher review. The pointsEarned field supports partial credit and is set during the grading process. A unique constraint on (submissionId, questionId) ensures each question is answered only once per submission.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique answer identifier"),
            ("submissionId","VARCHAR(36)","FK->submissions","NO","—","Foreign key to the parent submission"),
            ("questionId","VARCHAR(36)","FK->questions","NO","—","Foreign key to the answered question"),
            ("answerText","TEXT","","NO","—","Student-provided answer text"),
            ("isCorrect","BOOLEAN","","YES","NULL","Whether the answer was marked correct (NULL = pending review)"),
            ("pointsEarned","INTEGER","","YES","NULL","Points awarded for this answer"),
        ],
    },
    {
        "table": "tasks", "model": "Task",
        "explanation": "Table 15 stores the performance tasks and written work assignments created by teachers. Unlike quizzes, tasks require students to upload files as their response, which the teacher then grades manually—optionally using a rubric. Tasks support configurable point values, DepEd grading categories, and due dates. Each task belongs to a course and may be categorized under a topic for organized display in the student's classwork view.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique task identifier"),
            ("title","VARCHAR","","NO","—","Task title"),
            ("instructions","TEXT","","YES","NULL","Task instructions shown to students"),
            ("taskType","VARCHAR(50)","","NO","ASSIGNMENT","DepEd grading category: ASSIGNMENT, WRITTEN_WORK, or PERFORMANCE_TASK"),
            ("gradingWeightId","VARCHAR(36)","FK->grading_weights","YES","NULL","Foreign key to the grading weight category"),
            ("points","INTEGER","","NO","100","Maximum point value"),
            ("dueDate","TIMESTAMP","","YES","NULL","Submission due date"),
            ("topicId","VARCHAR(36)","FK->topics","YES","NULL","Foreign key to the parent topic"),
            ("courseId","VARCHAR(36)","FK->courses","NO","—","Foreign key to the parent course"),
            ("teacherId","VARCHAR(36)","FK->users","NO","—","Foreign key to the teacher who created the task"),
            ("status","VARCHAR(50)","","NO","DRAFT","Task status: DRAFT, PUBLISHED, or CLOSED"),
            ("linkUrl","VARCHAR","","YES","NULL","Linked external URL"),
            ("youtubeUrl","VARCHAR","","YES","NULL","Linked YouTube video URL"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
            ("updatedAt","TIMESTAMP","","NO","auto","Timestamp of the last record update"),
        ],
    },
    {
        "table": "task_attachments", "model": "TaskAttachment",
        "explanation": "Table 16 stores files, links, or YouTube videos attached to a task by the teacher to provide supplementary information, rubrics, or reference materials for students. Attachments may be served from local storage or Google Drive. These materials are visible to all enrolled students when they open the task detail page.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique attachment identifier"),
            ("taskId","VARCHAR(36)","FK->tasks","NO","—","Foreign key to the parent task"),
            ("kind","VARCHAR(50)","","NO","FILE","Attachment kind: FILE, LINK, or YOUTUBE"),
            ("fileName","VARCHAR","","YES","NULL","Original file name"),
            ("storedName","VARCHAR","","YES","NULL","Stored file name on disk"),
            ("filePath","VARCHAR","","YES","NULL","Relative path of the stored file"),
            ("fileSize","INTEGER","","YES","NULL","File size in bytes"),
            ("mimeType","VARCHAR","","YES","NULL","MIME type of the file"),
            ("storageBackend","VARCHAR(50)","","NO","LOCAL","Storage backend: LOCAL or GDRIVE"),
            ("driveFileId","VARCHAR","","YES","NULL","Google Drive file ID"),
            ("webViewLink","VARCHAR","","YES","NULL","Google Drive web view URL"),
            ("linkUrl","VARCHAR","","YES","NULL","External hyperlink URL"),
            ("youtubeUrl","VARCHAR","","YES","NULL","YouTube video URL"),
            ("createdAt","TIMESTAMP","","NO","now()","Attachment creation timestamp"),
        ],
    },
    {
        "table": "task_submissions", "model": "TaskSubmission",
        "explanation": "Table 17 records each student's submission for a task. A unique constraint on (taskId, studentId) ensures only one submission record exists per student per task. The record tracks the complete grading lifecycle from initial assignment through submission, grading, and return. The grade and previousGrade fields allow teachers to update scores while preserving a history of changes. Teacher feedback is stored alongside the grade for direct communication with the student.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique task submission identifier"),
            ("taskId","VARCHAR(36)","FK->tasks","NO","—","Foreign key to the submitted task"),
            ("studentId","VARCHAR(36)","FK->users","NO","—","Foreign key to the submitting student"),
            ("status","VARCHAR(50)","","NO","ASSIGNED","Status: ASSIGNED, SUBMITTED, RETURNED, GRADED, MISSING, or EXCUSED"),
            ("grade","decimal","","YES","NULL","Current grade or score"),
            ("previousGrade","decimal","","YES","NULL","Previous grade before the most recent update"),
            ("feedback","TEXT","","YES","NULL","Teacher feedback text"),
            ("submittedAt","TIMESTAMP","","YES","NULL","Datetime the student submitted"),
            ("gradedAt","TIMESTAMP","","YES","NULL","Datetime the submission was graded"),
            ("returnedAt","TIMESTAMP","","YES","NULL","Datetime the graded work was returned to the student"),
            ("gradeUpdatedAt","TIMESTAMP","","YES","NULL","Datetime the grade was last modified"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
            ("updatedAt","TIMESTAMP","","NO","auto","Timestamp of the last record update"),
        ],
    },
    {
        "table": "task_submission_files", "model": "TaskSubmissionFile",
        "explanation": "Table 18 stores the individual files uploaded by a student as part of a task submission. A single task submission may have multiple files attached. Files are stored in either the local server file system or on Google Drive depending on the student's configured storage backend. The table preserves file metadata including the original name, size, MIME type, and the storage key or Drive file ID required to retrieve the file.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique file record identifier"),
            ("submissionId","VARCHAR(36)","FK->task_submissions","NO","—","Foreign key to the parent task submission"),
            ("fileName","VARCHAR","","NO","—","Original uploaded file name"),
            ("storedName","VARCHAR","","NO","—","Stored file name on disk"),
            ("filePath","VARCHAR","","NO","—","Relative path of the stored file"),
            ("fileSize","INTEGER","","YES","NULL","File size in bytes"),
            ("mimeType","VARCHAR","","YES","NULL","MIME type of the uploaded file"),
            ("storageBackend","VARCHAR(50)","","NO","LOCAL","Storage backend: LOCAL or GDRIVE"),
            ("driveFileId","VARCHAR","","YES","NULL","Google Drive file ID"),
            ("webViewLink","VARCHAR","","YES","NULL","Google Drive web view URL"),
            ("createdAt","TIMESTAMP","","NO","now()","File upload timestamp"),
        ],
    },
    {
        "table": "rubrics", "model": "Rubric",
        "explanation": "Table 19 stores scoring rubrics that may be attached to either a task or a quiz. Each rubric contains an ordered set of criteria defining the dimensions of evaluation. Each criterion in turn defines multiple performance levels with associated descriptors and point values. Once grading has begun on a rubric, it may be locked to prevent modifications that would retroactively alter existing grades.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique rubric identifier"),
            ("taskId","VARCHAR(36)","FK->tasks","YES","NULL","Foreign key to the task this rubric is attached to (unique)"),
            ("quizId","VARCHAR(36)","FK->quizzes","YES","NULL","Foreign key to the quiz this rubric is attached to (unique)"),
            ("isLocked","BOOLEAN","","NO","false","Whether the rubric is locked from further editing"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
            ("updatedAt","TIMESTAMP","","NO","auto","Timestamp of the last record update"),
        ],
    },
    {
        "table": "rubric_criteria", "model": "RubricCriterion",
        "explanation": "Table 20 stores the individual criteria within a rubric, representing the specific dimensions evaluated when grading a student's work. Examples of criteria include Content, Grammar, Organization, and Presentation. Each criterion has a title, an optional description clarifying what is being evaluated, and a display order index for rendering the rubric in a consistent layout.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique criterion identifier"),
            ("rubricId","VARCHAR(36)","FK->rubrics","NO","—","Foreign key to the parent rubric"),
            ("title","VARCHAR","","NO","—","Criterion title (e.g., Content, Grammar)"),
            ("description","TEXT","","YES","NULL","Optional description of what is being evaluated"),
            ("orderIndex","INTEGER","","NO","0","Display sort order within the rubric"),
        ],
    },
    {
        "table": "rubric_levels", "model": "RubricLevel",
        "explanation": "Table 21 defines the performance levels for each rubric criterion. Typical levels include Excellent, Satisfactory, Developing, and Beginning, each carrying a descriptor text and a point value used to compute the student's score. Levels are ordered within their parent criterion for consistent display from highest to lowest performance.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique level identifier"),
            ("criterionId","VARCHAR(36)","FK->rubric_criteria","NO","—","Foreign key to the parent rubric criterion"),
            ("title","VARCHAR","","NO","—","Level title (e.g., Excellent, Satisfactory)"),
            ("description","TEXT","","YES","NULL","Descriptor text for this performance level"),
            ("points","decimal","","YES","NULL","Point value awarded at this level"),
            ("orderIndex","INTEGER","","NO","0","Display sort order within the criterion"),
        ],
    },
    {
        "table": "rubric_grades", "model": "RubricGrade",
        "explanation": "Table 22 records the rubric evaluation outcome for each criterion in a student's task submission. For each criterion a specific performance level is selected, and an optional pointsOverride field allows the teacher to award a custom score that differs from the level's default point value. A unique constraint on (submissionId, criterionId) ensures each criterion is scored only once per submission.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique rubric grade record identifier"),
            ("submissionId","VARCHAR(36)","FK->task_submissions","NO","—","Foreign key to the task submission being graded"),
            ("criterionId","VARCHAR(36)","FK->rubric_criteria","NO","—","Foreign key to the rubric criterion being evaluated"),
            ("levelId","VARCHAR(36)","FK->rubric_levels","YES","NULL","Selected performance level (NULL = no level selected)"),
            ("pointsOverride","decimal","","YES","NULL","Manual point override, ignoring the level's default points"),
        ],
    },
    {
        "table": "grading_weights", "model": "GradingWeight",
        "explanation": "Table 23 defines the DepEd grading weight categories for each course. Common categories include Written Works (25%), Performance Tasks (50%), and Quarterly Assessment (25%), mirroring the DepEd K-12 grading policy. These weights are applied by the grading engine to compute a student's final transmuted grade from raw quiz and task scores. Category names are unique within a course to prevent duplicate weight entries.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique grading weight identifier"),
            ("courseId","VARCHAR(36)","FK->courses","NO","—","Foreign key to the parent course"),
            ("name","VARCHAR","","NO","—","Category name (e.g., Written Works, Performance Tasks)"),
            ("percentage","decimal","","NO","—","Weight as a percentage (e.g., 25.0 for 25%)"),
            ("color","VARCHAR","","NO","#2563eb","Hex color used for this category in grade charts"),
            ("orderIndex","INTEGER","","NO","0","Display sort order within the course"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
            ("updatedAt","TIMESTAMP","","NO","auto","Timestamp of the last record update"),
        ],
    },
    {
        "table": "mastery_unlocks", "model": "MasteryUnlock",
        "explanation": "Table 24 records teacher-granted overrides that allow a specific student to access a mastery-locked quiz within a course. When mastery-lock gating is enabled, students who have not met the passing threshold on prerequisite quizzes are blocked from proceeding. This table stores exceptions manually granted by the teacher. A unique constraint on (courseId, studentId, quizId) ensures each unlock is issued only once per student per quiz per course.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","CUID()","Unique unlock record identifier"),
            ("courseId","VARCHAR(36)","FK->courses","NO","—","Foreign key to the course where the unlock applies"),
            ("studentId","VARCHAR(36)","FK->users","NO","—","Foreign key to the student whose access is unlocked"),
            ("quizId","VARCHAR(36)","FK->quizzes","NO","—","Foreign key to the quiz being unlocked"),
            ("unlockedBy","VARCHAR(36)","FK->users","NO","—","Foreign key to the teacher who granted the unlock"),
            ("createdAt","TIMESTAMP","","NO","now()","Datetime the unlock was granted"),
        ],
    },
    {
        "table": "task_comments", "model": "TaskComment",
        "explanation": "Table 25 stores comments attached to tasks, enabling direct communication between teachers and students about task instructions, clarifications, or feedback. Comments may be marked private, meaning they are visible only to the author and the specified recipient. This supports both class-wide discussion and private teacher-to-student feedback within the context of a specific task.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique comment identifier"),
            ("taskId","VARCHAR(36)","FK->tasks","NO","—","Foreign key to the parent task"),
            ("authorId","VARCHAR(36)","FK->users","NO","—","Foreign key to the user who wrote the comment"),
            ("recipientId","VARCHAR(36)","FK->users","YES","NULL","Foreign key to the intended recipient (optional)"),
            ("content","TEXT","","NO","—","Comment body text"),
            ("isPrivate","BOOLEAN","","NO","false","Whether the comment is private"),
            ("createdAt","TIMESTAMP","","NO","now()","Comment creation timestamp"),
        ],
    },
    {
        "table": "quiz_comments", "model": "QuizComment",
        "explanation": "Table 26 stores comments attached to quizzes, mirroring the task comment design for quiz-related communication. Teachers may post clarifications about quiz instructions, and students may ask questions. Private comments are visible only to the author and specified recipient, enabling direct teacher-to-student feedback within the context of a specific quiz.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique comment identifier"),
            ("quizId","VARCHAR(36)","FK->quizzes","NO","—","Foreign key to the parent quiz"),
            ("authorId","VARCHAR(36)","FK->users","NO","—","Foreign key to the user who wrote the comment"),
            ("recipientId","VARCHAR(36)","FK->users","YES","NULL","Foreign key to the intended recipient (optional)"),
            ("content","TEXT","","NO","—","Comment body text"),
            ("isPrivate","BOOLEAN","","NO","false","Whether the comment is private"),
            ("createdAt","TIMESTAMP","","NO","now()","Comment creation timestamp"),
        ],
    },
    {
        "table": "announcements", "model": "Announcement",
        "explanation": "Table 27 stores course announcements posted by teachers to keep enrolled students informed. Announcements support a draft mode for editing before publication and a scheduling feature for deferred release at a specified datetime. A single file attachment may be included for supplementary documents. The recipientIds JSON field enables teachers to target announcements at a specific subset of students rather than the entire class.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","CUID()","Unique announcement identifier"),
            ("courseId","VARCHAR(36)","FK->courses","NO","—","Foreign key to the parent course"),
            ("authorId","VARCHAR(36)","FK->users","NO","—","Foreign key to the teacher who created the announcement"),
            ("content","TEXT","","NO","—","Announcement body text"),
            ("isDraft","BOOLEAN","","NO","false","Whether the announcement is still a draft"),
            ("scheduledAt","TIMESTAMP","","YES","NULL","Datetime to auto-publish the announcement"),
            ("attachmentName","VARCHAR","","YES","NULL","Name of the attached file"),
            ("attachmentPath","VARCHAR","","YES","NULL","Storage path of the attached file"),
            ("attachmentSize","INTEGER","","YES","NULL","File size in bytes"),
            ("attachmentDriveFileId","VARCHAR","","YES","NULL","Google Drive file ID for the attachment"),
            ("attachmentWebViewLink","VARCHAR","","YES","NULL","Google Drive web view URL for the attachment"),
            ("recipientIds","json","","YES","NULL","JSON array of targeted student user IDs (NULL = all students)"),
            ("createdAt","TIMESTAMP","","NO","now()","Record creation timestamp"),
        ],
    },
    {
        "table": "announcement_comments", "model": "AnnouncementComment",
        "explanation": "Table 28 stores replies to course announcements from both teachers and students. This creates a threaded discussion channel around each announcement, allowing students to ask questions or react to posted information. All comments on an announcement are visible to all enrolled students, fostering open class communication.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","CUID()","Unique comment identifier"),
            ("announcementId","VARCHAR(36)","FK->announcements","NO","—","Foreign key to the parent announcement"),
            ("authorId","VARCHAR(36)","FK->users","NO","—","Foreign key to the user who wrote the comment"),
            ("content","TEXT","","NO","—","Comment body text"),
            ("createdAt","TIMESTAMP","","NO","now()","Comment creation timestamp"),
        ],
    },
    {
        "table": "activity_logs", "model": "ActivityLog",
        "explanation": "Table 29 provides an audit trail of user interactions with resources in the system. Actions such as viewing or downloading a learning material are recorded with the acting user's identity and a timestamp. The metadata JSON field stores additional contextual information such as the resource title or file name, supporting analytics and monitoring of student engagement with course materials.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","UUID()","Unique activity log entry identifier"),
            ("userId","VARCHAR(36)","FK->users","NO","—","Foreign key to the user who performed the action"),
            ("action","VARCHAR","","NO","—","Action performed (e.g., VIEW, DOWNLOAD)"),
            ("resourceId","VARCHAR(36)","FK->resources","YES","NULL","Foreign key to the resource that was acted upon"),
            ("metadata","json","","YES","NULL","Additional contextual metadata for the action"),
            ("createdAt","TIMESTAMP","","NO","now()","Datetime the action was recorded"),
        ],
    },
    {
        "table": "notifications", "model": "Notification",
        "explanation": "Table 30 stores in-app notifications sent to users to alert them of important events. The primary notification type is REMEDIAL_PENDING, which notifies a student that an AI-generated remedial quiz has been created for them following a failing score. Each notification tracks its read status so the UI can display an unread badge count. The optional data JSON field carries contextual information such as a quiz ID or course ID that the application uses to navigate the user to the relevant content when the notification is tapped.",
        "fields": [
            ("id","VARCHAR(36)","PK","NO","CUID()","Unique notification identifier"),
            ("userId","VARCHAR(36)","FK->users","NO","—","Foreign key to the recipient user"),
            ("type","VARCHAR(50)","","NO","—","Notification type (e.g., REMEDIAL_PENDING)"),
            ("title","VARCHAR","","NO","—","Short notification title"),
            ("body","TEXT","","NO","—","Notification body text"),
            ("data","json","","YES","NULL","Optional JSON payload with contextual data"),
            ("read","BOOLEAN","","NO","false","Whether the notification has been read"),
            ("createdAt","TIMESTAMP","","NO","now()","Notification creation timestamp"),
        ],
    },
]


def build():
    doc = Document()

    # Portrait A4, standard margins
    sec = doc.sections[0]
    sec.page_width   = Cm(21.0)
    sec.page_height  = Cm(29.7)
    sec.left_margin  = sec.right_margin = Cm(2.54)
    sec.top_margin   = sec.bottom_margin = Cm(2.54)

    # Heading
    h = doc.add_paragraph()
    r = h.add_run("Data Dictionary")
    r.bold      = True
    r.font.size = Pt(14)

    # Intro paragraph
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.first_line_indent = Cm(1.25)
    intro.paragraph_format.space_after = Pt(6)
    ir = intro.add_run(
        "A data dictionary documents and explains the structure of the AIMS "
        "(AI-Integrated Monitoring System) database. It provides a comprehensive "
        "reference of every table and field in the system, including data types, "
        "field lengths, key relationships, and descriptions. This document is "
        "intended to guide developers, database administrators, and academic "
        "reviewers in understanding how data is organized, stored, and related "
        "within the system. Each table entry is accompanied by an explanation "
        "describing the table's purpose and how it contributes to the overall "
        "functionality of the platform."
    )
    ir.font.size = Pt(11)

    doc.add_paragraph()

    for num, tdef in enumerate(TABLES, 1):
        tname       = tdef["table"]
        fields      = tdef["fields"]
        explanation = tdef["explanation"]

        # "Table N"
        th = doc.add_paragraph()
        th.paragraph_format.space_after = Pt(0)
        tr = th.add_run(f"Table {num}")
        tr.bold      = True
        tr.font.size = Pt(12)

        # Underlined table name
        tn = doc.add_paragraph()
        tn.paragraph_format.space_after = Pt(2)
        tnr = tn.add_run(tname)
        tnr.underline = True
        tnr.font.size = Pt(11)

        # 4-column table
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style     = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        hrow = tbl.rows[0]
        for i, hdr in enumerate(["Field Name", "Field Type", "Length", "Description"]):
            cell_write(hrow.cells[i], hdr, bold=True, size=10)

        for fname, dtype, kstr, _nullable, _default, desc in fields:
            ftype, flen = simplify(dtype)
            kt  = key_type(kstr)
            row = tbl.add_row()

            if kt == "PK":
                cell_write(row.cells[0], fname, bold=True, underline=True, size=10)
            elif kt == "FK":
                cell_write(row.cells[0], fname, red=True, underline=True, size=10)
            elif kt == "UQ":
                cell_write(row.cells[0], fname, underline=True, size=10)
            else:
                cell_write(row.cells[0], fname, size=10)

            cell_write(row.cells[1], ftype, size=10)
            cell_write(row.cells[2], flen,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_write(row.cells[3], desc,  size=10)

        set_table_borders(tbl)
        fix_col_widths(tbl, COL_W)

        # Explanation paragraph
        doc.add_paragraph()
        ep = doc.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ep.paragraph_format.first_line_indent = Cm(1.25)
        ep.paragraph_format.space_after = Pt(6)
        er = ep.add_run(explanation)
        er.font.size = Pt(11)

        doc.add_paragraph()

    base = os.path.dirname(os.path.abspath(__file__))
    out  = os.path.join(base, "AIMS_Data_Dictionary.docx")
    # If file is locked (open in Word), fall back to a versioned name
    try:
        doc.save(out)
    except PermissionError:
        out = os.path.join(base, "AIMS_Data_Dictionary_new.docx")
        doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
