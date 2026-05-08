#!/usr/bin/env python3
"""build_v13.py

Builds papers/AIMS/revisions/cha1&2_v13.docx from v12.

Scope (minimal, Chapter 2 only):
- Update the 5 RRL section headings to match the finalized Chapter 1 objectives wording.
- Fix the Dynamic Remedial section heading formatting to match the other headings (bold).
- Update the Chapter 2 synthesis paragraphs to use the exact objective phrases.
- Update the Matrix (Table 1) rotated feature column headers to the exact objective phrases.

Constraints:
- Preserve v8 Chapter 2 formatting conventions (left-aligned bold headings, spacing).
- Do not introduce the word “dashboard”.
- Do not rewrite the RRL body paragraphs beyond the synthesis block.
- Never overwrite existing revisions.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v12.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v13.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    m = _leading_ws_re.match(text)
    return m.group(0) if m else ""


def replace_paragraph_text_preserve_runs(p, new_text: str) -> None:
    """Replace paragraph text without destroying run formatting.

    Strategy:
    - Preserve leading whitespace from existing paragraph.
    - Put the whole new text in the first run (creating one if missing).
    - Blank out remaining runs.
    """

    new_text = sanitize_xml_text(new_text.strip())
    prefix = leading_ws(p.text)
    final_text = prefix + new_text

    if not p.runs:
        p.add_run(final_text)
        return

    p.runs[0].text = final_text
    for r in p.runs[1:]:
        r.text = ""


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(doc: Document, needle: str, start: int = 0, end: int | None = None) -> int:
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def ensure_all_runs_bold(p) -> None:
    for r in p.runs:
        if r.text.strip():
            r.font.bold = True


def set_cell_text_preserve_runs(cell, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    if not cell.paragraphs:
        cell.text = new_text
        return
    p = cell.paragraphs[0]
    replace_paragraph_text_preserve_runs(p, new_text)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    # Target objective phrases (match Chapter 1 objectives wording)
    obj_11 = "Centralized learning portal with material distribution and task submission"
    obj_12 = "AI-Powered Quiz Generator with instructor validation and approval controls"
    obj_13 = "Automated assessment and grading engine supporting teacher validation"
    obj_14 = "Dynamic remedial quiz generator for targeted student intervention"
    obj_15 = "Mastery-based learning module with configurable progression locks"

    # ── Chapter 2 section headings ──────────────────────────────────────────
    ch2_start = find_paragraph_index(doc, "Chapter 2")
    ref_start = find_paragraph_index(doc, "REFERENCES")

    heading_map = {
        "Centralized Learning Portal for Module Distribution and Management": obj_11,
        "AI-Powered Quiz Generator with Instructor Validation Controls": obj_12,
        "Automated Assessment and Grading Engine": obj_13,
        "Dynamic Remedial Quiz Generator for Targeted Student Intervention": obj_14,
        "Mastery-Based Learning Module and Configurable Progression": obj_15,
    }

    # Update known headings by exact match
    for i in range(ch2_start, ref_start):
        t = doc.paragraphs[i].text.strip()
        if t in heading_map:
            replace_paragraph_text_preserve_runs(doc.paragraphs[i], heading_map[t])

    # Ensure the Dynamic Remedial section heading is bold (v12 had it unbolded)
    try:
        dyn_i = find_paragraph_index(doc, "Dynamic Remedial Quiz Generator for Targeted Student Intervention")
        # If it still exists, update then bold
        replace_paragraph_text_preserve_runs(doc.paragraphs[dyn_i], obj_14)
        ensure_all_runs_bold(doc.paragraphs[dyn_i])
    except ValueError:
        # In v12, this line exists; after update it becomes obj_14, so find by new text
        dyn_i = find_paragraph_index(doc, obj_14)
        ensure_all_runs_bold(doc.paragraphs[dyn_i])

    # ── Synthesis paragraphs (rewrite block only; keep citations; remove legacy terms) ──
    syn_h = find_paragraph_index(doc, "Synthesis of the Reviewed Literature")

    # Paragraphs immediately following synthesis heading, up to Table 1.
    table_i = None
    for i in range(syn_h + 1, ref_start):
        if doc.paragraphs[i].text.strip() in {"Table 1.", "Table 1"}:
            table_i = i
            break
    if table_i is None:
        raise RuntimeError("Could not find Table 1 after synthesis heading")

    syn_p1 = (
        "The reviewed literature collectively demonstrates that each of the five core features of A.I.M.S. is grounded in a robust and growing body of empirical evidence. "
        f"{obj_11} addresses persistent access and distribution inequities in public school environments (Bradley, 2021; Bustillo & Aguilos, 2022), while {obj_12} reduces the cognitive and administrative burden of assessment creation through validated neural and LLM-based generation methods (Bulathwela et al., 2023; Alamoudi et al., 2025). "
        f"{obj_13} promotes evaluation consistency and scalability through machine learning and rubric-guided assessment (Ramesh & Sanampudi, 2022; García-Varela et al., 2025). "
        f"{obj_14} supports targeted intervention by operationalizing corrective feedback and structured remediation cycles aligned with identified learning gaps (Persky & Hughes, 2022; Jumao-as et al., 2025). "
        f"Finally, {obj_15} provides structured, competency-based progression that supports mastery attainment and reduces learning gaps (Persky & Hughes, 2022; Jumao-as et al., 2025)."
    )

    syn_p2 = (
        "A consistent theme across both foreign and local literature is that the greatest educational impact is achieved when multiple intelligent features are integrated cohesively within a single platform, "
        "rather than deployed as fragmented standalone tools that create workflow disruptions and adoption resistance (Balase & Paglinawan, 2025; Colegado, 2025). "
        "The reviewed studies demonstrate that while individual functionalities such as AI-assisted quiz generation, automated assessment, targeted remediation, and mastery-based progression have been validated in isolation, "
        "no existing platform fully integrates the complete instructional and assessment cycle in the manner envisioned in A.I.M.S. This integration gap represents the primary research and design opportunity that the present study addresses."
    )

    syn_p3 = (
        "Overall, the literature reviewed in this chapter provides strong theoretical, empirical, and contextual justification for the design and development of A.I.M.S. as a unified, AI-driven educational platform tailored to the Philippine public school context. "
        "The evidence base confirms that the proposed features are not merely technically feasible but also pedagogically sound, contextually responsive, and aligned with contemporary best practices in educational technology research. "
        "By operationalizing verified AI capabilities within a coherent, teacher-centered system architecture, A.I.M.S. represents a research-supported response to documented gaps in Philippine digital learning environments, integrating access management, automated assessment, targeted remediation, and mastery-based progression into a single, scalable educational system."
    )

    # Replace the first 3 non-empty synthesis paragraphs before Table 1
    syn_targets = []
    for i in range(syn_h + 1, table_i):
        if doc.paragraphs[i].text.strip():
            syn_targets.append(i)
    if len(syn_targets) < 3:
        raise RuntimeError("Unexpected synthesis block; not enough paragraphs")

    replace_paragraph_text_preserve_runs(doc.paragraphs[syn_targets[0]], syn_p1)
    replace_paragraph_text_preserve_runs(doc.paragraphs[syn_targets[1]], syn_p2)
    replace_paragraph_text_preserve_runs(doc.paragraphs[syn_targets[2]], syn_p3)

    # ── Post-table 'Synthesis' paragraphs (update feature list wording) ──────
    # Find the 'Synthesis' heading after the table caption.
    syn2_h = None
    for i in range(table_i, ref_start):
        if doc.paragraphs[i].text.strip() == "Synthesis":
            syn2_h = i
            break

    if syn2_h is not None:
        # Replace the next two non-empty paragraphs.
        p_idxs = []
        for i in range(syn2_h + 1, ref_start):
            t = doc.paragraphs[i].text.strip()
            if t in {"REFERENCES", "Chapter 3"}:
                break
            if t:
                p_idxs.append(i)
            if len(p_idxs) >= 2:
                break

        if len(p_idxs) >= 2:
            post1 = (
                "The matrix of related literature and studies above presents a comparative overview of existing systems and studies in relation to the five core technical features of the proposed A.I.M.S. platform. "
                f"The studies mapped in the matrix span both foreign and local literature, covering {obj_11} (Bradley, 2021; Bustillo & Aguilos, 2022), {obj_12} (Bulathwela et al., 2023; Alamoudi et al., 2025), {obj_13} (Ramesh & Sanampudi, 2022; García-Varela et al., 2025), {obj_14} (Persky & Hughes, 2022; Jumao-as et al., 2025), and {obj_15} (Persky & Hughes, 2022; Jumao-as et al., 2025)."
            )
            post2 = (
                "Overall, the synthesis reveals that while the individual features addressed by A.I.M.S. are present in existing systems and have been independently validated in the literature, "
                f"no single platform currently integrates all five functionalities: {obj_11}, {obj_12}, {obj_13}, {obj_14}, and {obj_15}, into a unified, AI-driven educational environment. "
                "This integration gap constitutes the primary research contribution of A.I.M.S. and affirms the novelty and practical significance of the proposed system as a response to documented needs in the Philippine educational technology landscape."
            )
            replace_paragraph_text_preserve_runs(doc.paragraphs[p_idxs[0]], post1)
            replace_paragraph_text_preserve_runs(doc.paragraphs[p_idxs[1]], post2)

    # ── Matrix table header row (rotated feature columns) ───────────────────
    if not doc.tables:
        raise RuntimeError("No tables found; expected Matrix table")
    tbl = doc.tables[0]

    # The rotated feature headers are in row 1, columns 1..5 in this document.
    header_row = tbl.rows[1]
    header_values = [obj_11, obj_12, obj_13, obj_14, obj_15]
    for col_i, text in enumerate(header_values, start=1):
        set_cell_text_preserve_runs(header_row.cells[col_i], text)

    # Safety check: do not allow the forbidden word
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "dashboard" in full_text.lower():
        raise RuntimeError("Forbidden word 'dashboard' found after updates")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()
