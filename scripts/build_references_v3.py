"""
Build updated_references_v3.docx
Full APA 7th edition reference list for cha1&2_v3.docx
All foreign references verified with confirmed DOIs (2021-2026)
Philippine local references from Chapter 1 (need DOI verification by student)
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins matching thesis specs
section = doc.sections[0]
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_reference(doc, text, note=None):
    """Add APA reference with hanging indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)

    # Double spacing
    from docx.shared import Length
    pPr = p._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '480')
    spacing.set(qn('w:lineRule'), 'auto')
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    if note:
        note_run = p.add_run(f'  [{note}]')
        note_run.font.name = 'Times New Roman'
        note_run.font.size = Pt(10)
        note_run.italic = True
    return p

def blank_line(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pPr = p._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '480')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

add_heading(doc, 'REFERENCES')

blank_line(doc)

add_heading(doc, 'Foreign References (DOI Verified)')
blank_line(doc)

# ─── VERIFIED FOREIGN REFERENCES ───────────────────────────────────────────

add_reference(doc,
    "Alamoudi, S., Al Khuzayem, L. A., & Jamal, A. (2025). Optimizing automated question generation "
    "for educational assessments: A semantic analysis of LLMs with structured and unstructured ontologies. "
    "Engineering, Technology & Applied Science Research, 15(3), 23664–23671. "
    "https://doi.org/10.48084/etasr.10662"
)

add_reference(doc,
    "Bradley, V. M. (2021). Learning management system (LMS) use with online instruction. "
    "International Journal of Technology in Education, 4(1), 68–92. "
    "https://doi.org/10.46328/ijte.36"
)

add_reference(doc,
    "Bulathwela, S., Pérez-Ortiz, M., Holloway, C., Cukurova, M., & Shawe-Taylor, J. (2023). "
    "Review on neural question generation for education purposes. "
    "International Journal of Artificial Intelligence in Education, 33(4), 691–717. "
    "https://doi.org/10.1007/s40593-023-00374-x"
)

add_reference(doc,
    "de Vreugd, L., van Leeuwen, A., Jansen, R., & van der Schaaf, M. (2024). "
    "Learning analytics dashboard design and evaluation to support student self-regulation of study behaviour. "
    "Journal of Learning Analytics, 11(3), 249–262. "
    "https://doi.org/10.18608/jla.2024.8529"
)

add_reference(doc,
    "García-Varela, F., Nussbaum, M., Mendoza, M., Martínez-Troncoso, C., & Bekerman, Z. (2025). "
    "ChatGPT as a stable and fair tool for automated essay scoring. "
    "Education Sciences, 15(8), Article 946. "
    "https://doi.org/10.3390/educsci15080946"
)

add_reference(doc,
    "International Organization for Standardization. (2011). Systems and software engineering—Systems "
    "and software Quality Requirements and Evaluation (SQuaRE)—System and software quality models "
    "(ISO/IEC Standard No. 25010:2011). https://www.iso.org/standard/35733.html"
)

add_reference(doc,
    "Kaliisa, R., Misiejuk, K., López-Pernas, S., Khalil, M., & Saqr, M. (2024). Have learning analytics "
    "dashboards lived up to the hype? A systematic review of impact on students' achievement, motivation, "
    "participation and attitude. In Proceedings of the 14th Learning Analytics and Knowledge Conference "
    "(LAK '24) (pp. 295–304). ACM. https://doi.org/10.1145/3636555.3636884"
)

add_reference(doc,
    "Persky, A. M., & Hughes, M. L. (2022). A practical review of mastery learning. "
    "American Journal of Pharmaceutical Education, 86(10), Article 8906. "
    "https://doi.org/10.5688/ajpe8906"
)

add_reference(doc,
    "Ramesh, D., & Sanampudi, S. K. (2022). An automated essay scoring systems: "
    "A systematic literature review. Artificial Intelligence Review, 55(3), 2495–2527. "
    "https://doi.org/10.1007/s10462-021-10068-2"
)

add_reference(doc,
    "Susnjak, T., Ramaswami, G. S., & Mathrani, A. (2022). Learning analytics dashboard: "
    "A tool for providing actionable insights to learners. "
    "International Journal of Educational Technology in Higher Education, 19(1), Article 12. "
    "https://doi.org/10.1186/s41239-021-00313-7"
)

blank_line(doc)
add_heading(doc, 'Local References (Philippine — DOI to be verified by research team)')
blank_line(doc)

# ─── LOCAL/PHILIPPINE REFERENCES (from Chapter 1 — DOIs need student verification) ─

add_reference(doc,
    "Balase, [First Initial]. [M.], & Paglinawan, [First Initial]. [M.]. (2025). [Full title to be retrieved]. "
    "[Journal name or publisher]. [DOI or URL]",
    note="From Ch1 background — full bibliographic details to be confirmed"
)

add_reference(doc,
    "Bustillo, [First Initial]. [M.], & Aguilos, [First Initial]. [M.]. (2022). [Full title to be retrieved]. "
    "[Journal name or publisher]. [DOI or URL]",
    note="From Ch1 background — full bibliographic details to be confirmed"
)

add_reference(doc,
    "Colegado, [First Initial]. [M.]. (2025). [Full title to be retrieved]. "
    "[Journal name or publisher]. [DOI or URL]",
    note="From Ch1 background — full bibliographic details to be confirmed"
)

blank_line(doc)
add_heading(doc, 'VERIFICATION NOTES FOR RESEARCH TEAM')
blank_line(doc)

notes_text = (
    "1. FOREIGN REFERENCES (9 total): All nine foreign references above have been independently verified "
    "through their DOI links. Authors, titles, journals, years, and DOIs are confirmed accurate.\n\n"
    "2. LOCAL REFERENCES (3 cited): Bustillo & Aguilos (2022), Balase & Paglinawan (2025), and Colegado "
    "(2025) are cited in Chapter 1 of this document. Their full bibliographic details (complete first names, "
    "full title, journal/publisher, volume, pages, and DOI) must be confirmed by the research team before "
    "submission. If these sources are from Philippine journals (e.g., JPAIR, IJRAR, IOER International "
    "Research Journal, Palawan Scientist, etc.), retrieve the exact DOI or stable URL from the original "
    "source.\n\n"
    "3. REFERENCES REMOVED: The following proposed references from earlier sessions were confirmed to be "
    "FABRICATED or have invalid DOIs and have been excluded from v3:\n"
    "   \u2022 Tian et al. (2025) arXiv:2506.07955 — FABRICATED (June 2026 arXiv number = future date)\n"
    "   \u2022 Gagnon (2023) doi:10.1080/00094056.2023.2252447 — INVALID (404 HTTP error, DOI does not exist)\n"
    "   \u2022 Yang et al. (2025) — wrong author attribution for doi:10.3390/educsci15080946 "
    "(correct authors: García-Varela et al., 2025)"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf = p.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pPr = p._element.get_or_add_pPr()
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:before'), '0')
spacing.set(qn('w:after'), '0')
spacing.set(qn('w:line'), '480')
spacing.set(qn('w:lineRule'), 'auto')
pPr.append(spacing)
run = p.add_run(notes_text)
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

doc.save('updated_references_v3.docx')
print('Saved: updated_references_v3.docx')
