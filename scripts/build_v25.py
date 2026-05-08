#!/usr/bin/env python3
"""build_v25.py

Builds papers/AIMS/revisions/cha1&2_v25.docx from v24.

Scope:
- Replace the ISO/IEC 25010 contextual citation with a 2021-2026 source.
- Keep ISO/IEC 25010 as the operational evaluation framework.
- Expand Chapter 1 Definition of Terms with cross-cutting A.I.M.S. concepts.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v24.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v25.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    match = _leading_ws_re.match(text)
    return match.group(0) if match else ""


def replace_paragraph_text_preserve_runs(paragraph, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    prefix = leading_ws(paragraph.text)
    final_text = prefix + new_text

    if not paragraph.runs:
        paragraph.add_run(final_text)
        return

    paragraph.runs[0].text = final_text
    for run in paragraph.runs[1:]:
        run.text = ""


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(
    doc: Document,
    needle: str,
    start: int = 0,
    end: int | None = None,
) -> int:
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def insert_paragraphs_before(target_paragraph, texts: list[str]) -> None:
    for text in texts:
        new_paragraph = target_paragraph.insert_paragraph_before(sanitize_xml_text(text))
        new_paragraph.style = target_paragraph.style


def insert_reference_before_trailing_blank(doc: Document, reference_text: str) -> None:
    references_i = find_paragraph_index(doc, "REFERENCES")
    for i in range(references_i + 1, len(doc.paragraphs)):
        if not doc.paragraphs[i].text.strip():
            new_paragraph = doc.paragraphs[i].insert_paragraph_before(
                sanitize_xml_text(reference_text)
            )
            new_paragraph.style = doc.paragraphs[i - 1].style
            return

    doc.add_paragraph(sanitize_xml_text(reference_text))


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    chapter1 = find_paragraph_index(doc, "Chapter 1")
    chapter2 = find_paragraph_index(doc, "Chapter 2")

    iso_context_i = find_paragraph_index_contains(
        doc,
        "ISO/IEC 25010. Contextually",
        start=chapter1,
        end=chapter2,
    )
    iso_context_new = (
        "ISO/IEC 25010. Contextually, it refers to a software product quality model that organizes system evaluation into core characteristics and measurable attributes for judging how well a software product satisfies defined quality requirements across different use contexts (Zhang et al., 2023)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[iso_context_i], iso_context_new)

    definition_terms = [
        (
            "Instructor Validation. Contextually, it refers to the human review and approval step used to verify AI-assisted assessment outputs for relevance, accuracy, and pedagogical appropriateness before instructional use (Gardner et al., 2021; Gorgun & Bulut, 2024)."
        ),
        (
            "Operationally, it is the mandatory teacher review workflow in A.I.M.S. through which AI-generated quizzes and remedial interventions remain pending until the teacher edits, approves, or rejects them before release to students."
        ),
        (
            "Objective Quiz Items. Contextually, these are structured assessment items with fixed or clearly delimited correct answers, such as multiple-choice questions, which are suitable for automated scoring and rapid feedback in digital learning environments (Alamoudi et al., 2025; Ouyang et al., 2023)."
        ),
        (
            "Operationally, they are the machine-scoreable quiz questions generated and administered by A.I.M.S., including multiple-choice items with predefined answer keys used for immediate preliminary scoring."
        ),
        (
            "Progression Locks. Contextually, these are rule-based access controls in mastery learning environments that restrict advancement to subsequent content until learners satisfy prerequisite performance criteria or complete corrective activities (Persky & Hughes, 2022; Toti et al., 2023)."
        ),
        (
            "Operationally, they are the configurable access restrictions in A.I.M.S. that keep later lessons, activities, or materials unavailable until the learner meets the teacher-set mastery threshold or receives instructor override approval."
        ),
        (
            "Teacher Override. Contextually, it refers to the instructor's discretionary authority in digital learning systems to supersede automated progression or intervention rules when professional judgment indicates a need for contextual adjustment or exception handling (Gardner et al., 2021; Persky & Hughes, 2022)."
        ),
        (
            "Operationally, it is the control in A.I.M.S. that allows teachers to manually unlock progression, release pending remedial activities, or finalize validated scores when learner context warrants an exception."
        ),
        (
            "Remedial Intervention. Contextually, it refers to targeted follow-up instruction or assessment designed to address identified learning gaps after formative evaluation and support subsequent mastery achievement (Hegde et al., 2024; Jumao-as et al., 2025)."
        ),
        (
            "Operationally, it is the teacher-approved corrective activity in A.I.M.S., typically delivered as a remedial quiz linked to unmet competencies before reassessment and re-entry into the mastery progression flow."
        ),
    ]
    insert_paragraphs_before(doc.paragraphs[iso_context_i], definition_terms)

    zhang_reference = (
        "Zhang, C., Li, B., Wang, L., Xu, H., & Shao, T. (2023). A hierarchical model for quality evaluation of mixed source software based on ISO/IEC 25010. International Journal of Software Engineering and Knowledge Engineering, 33(2), 181-205. https://doi.org/10.1142/S021819402250070X"
    )
    insert_reference_before_trailing_blank(doc, zhang_reference)

    full_ch1 = "\n".join(p.text for p in doc.paragraphs[chapter1:chapter2])
    required_terms = [
        "Instructor Validation.",
        "Objective Quiz Items.",
        "Progression Locks.",
        "Teacher Override.",
        "Remedial Intervention.",
        "Zhang et al., 2023",
    ]
    for term in required_terms:
        if term not in full_ch1:
            raise RuntimeError(f"Missing Chapter 1 content: {term}")

    references_text = "\n".join(p.text for p in doc.paragraphs[find_paragraph_index(doc, 'REFERENCES') :])
    if "10.1142/S021819402250070X" not in references_text:
        raise RuntimeError("Missing Zhang reference DOI")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()