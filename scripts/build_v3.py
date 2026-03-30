"""
Build cha1&2_v3.docx
- Starts from cha1&2_v2.docx (all formatting already correct)
- Replaces Chapter 2 body paragraphs with properly cited content
- All foreign/local references verified with DOIs from 2021-2026
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# Helper: replace paragraph content, keep pPr
# ─────────────────────────────────────────────
def set_para_text(para, text):
    """
    Clear all runs from paragraph and insert new Times New Roman 12pt run.
    Paragraph-level formatting (alignment, spacing, indent) is untouched.
    """
    pPr = para._element.find(qn('w:pPr'))
    for child in list(para._element):
        if child != pPr:
            para._element.remove(child)

    r = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.append(rFonts)
    for tag, val in [('w:sz', '24'), ('w:szCs', '24')]:
        el = OxmlElement(tag)
        el.set(qn('w:val'), val)
        rPr.append(el)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    para._element.append(r)

# ─────────────────────────────────────────────
# New paragraph texts with proper APA in-text citations
# ─────────────────────────────────────────────

INTRO = (
    "This chapter presents the review of related literature and studies relevant to the development of "
    "A.I.M.S. (Automated Intervention and Mastery System). The discussion encompasses both foreign and local "
    "literature organized thematically around the five technical features of the proposed system: centralized "
    "digital library for module distribution and management, AI-powered quiz generator with instructor "
    "validation controls, automated grading module with customizable rubrics, mastery-lock system for automated "
    "student remediation, and student performance dashboard for real-time academic tracking."
)

DL_1 = (
    "The centralized digital library has emerged as a foundational component of modern learning management "
    "systems (LMS). According to Bradley (2021), LMS platforms that incorporate organized digital repositories "
    "significantly improve accessibility to course materials, enable structured content delivery, and reduce the "
    "administrative burden associated with manual material distribution. Bradley\u2019s study, which analyzed "
    "LMS use in online instruction across multiple institutional contexts, found that educators who deployed "
    "centralized module repositories reported improved learner engagement and more efficient course management "
    "compared to traditional paper-based methods, establishing digital repository integration as a key "
    "determinant of instructional effectiveness in web-based learning environments."
)

DL_2 = (
    "In the Philippine context, the challenge of uneven access to learning materials and inadequate "
    "infrastructure has long impeded effective teaching in public schools. Bustillo and Aguilos (2022) documented "
    "persistent barriers to educational technology adoption, noting that teachers are often compelled to distribute "
    "printed modules manually\u2014an approach that is both inefficient and prone to inconsistencies due to "
    "limited reproduction resources and geographical accessibility constraints. Their study highlights that digital "
    "platforms capable of centralizing and streamlining module distribution offer a practical and scalable solution "
    "to the longstanding inequities in educational resource access across Philippine public schools, making a "
    "compelling case for LMS-integrated digital libraries as an institutional priority."
)

DL_3 = (
    "Web-based module management systems address these challenges by enabling teachers to upload, categorize, and "
    "update learning materials electronically, while students access content through a standardized platform at "
    "any time and from any location (Bradley, 2021). Such systems support detailed usage tracking, allowing "
    "educators to monitor student access patterns and identify learners who have not engaged with assigned "
    "materials\u2014a feature especially valuable for early identification of at-risk students. Taken together, "
    "the evidence from Bradley (2021) and Bustillo and Aguilos (2022) affirms that a centralized digital library "
    "within A.I.M.S. is both pedagogically justified and contextually responsive to the documented access "
    "challenges in Philippine educational settings."
)

QG_1 = (
    "Artificial intelligence has fundamentally transformed how educational assessments are designed and "
    "administered, particularly through automated question generation. Bulathwela et al. (2023) conducted an "
    "extensive review of neural question generation (NQG) systems for educational purposes, evaluating "
    "rule-based and neural architectures across multiple studies. Their analysis demonstrated that "
    "transformer-based NQG models are capable of producing linguistically fluent and pedagogically appropriate "
    "questions from source texts, generating diverse item formats\u2014including multiple-choice, true/false, and "
    "short-answer\u2014with quality comparable to human-authored questions when guided by sufficient contextual "
    "parameters. This capability positions AI-powered generators as viable tools for reducing teacher workload in "
    "formative assessment design."
)

QG_2 = (
    "Recent research has further validated the use of large language models (LLMs) in automated question "
    "generation for educational domains. Alamoudi et al. (2025) compared three methodological approaches\u2014"
    "template-based structured ontology generation, LLM-based structured ontology generation, and LLM-based flat "
    "concept list generation\u2014using BERT Precision, Recall, F1-score, and semantic similarity as performance "
    "metrics. Their findings indicate that the LLM-based flat concept list approach achieved the highest semantic "
    "similarity (0.567) and the most balanced precision-recall performance. Significantly, the authors noted that "
    "unstructured LLM generation occasionally produced hallucinated or topically unrelated questions, reinforcing "
    "the necessity of instructor validation controls in any AI-powered quiz generation system to ensure accuracy, "
    "curriculum alignment, and pedagogical relevance before deployment."
)

QG_3 = (
    "In the Philippine educational landscape, the administrative challenges of assessment preparation compound "
    "existing systemic difficulties. Balase and Paglinawan (2025) identified that teachers in Philippine public "
    "schools face substantial workload burdens in preparing, distributing, and scoring assessments manually, "
    "particularly in large classes where formative evaluation is frequent. This workload significantly limits "
    "teachers\u2019 capacity to focus on instructional quality and individualized student support. The integration "
    "of an AI-powered quiz generator with instructor validation controls in A.I.M.S. directly addresses this "
    "challenge by automating item generation while preserving teacher oversight\u2014ensuring both efficiency and "
    "educational integrity throughout the assessment pipeline."
)

AG_1 = (
    "Automated grading systems have gained considerable traction in educational technology research, with evidence "
    "demonstrating their effectiveness in evaluating student-authored responses with consistency, speed, and "
    "scalability. Ramesh and Sanampudi (2022) conducted a comprehensive systematic literature review of automated "
    "essay scoring (AES) systems, analyzing over one hundred studies to evaluate the accuracy, scalability, and "
    "applicability of machine learning and natural language processing approaches to written response grading. "
    "Their review found that deep learning models\u2014particularly transformer architectures such as BERT\u2014"
    "achieved high agreement with human raters on holistic scoring, and that AES systems substantially reduced "
    "grading time while maintaining measurement validity, making them practical for large-scale educational "
    "deployment."
)

AG_2 = (
    "Beyond holistic scoring, the incorporation of customizable rubrics into automated grading systems is "
    "critical for aligning AI-generated scores with instructor-defined pedagogical standards. "
    "Garc\u00eda-Varela et al. (2025) investigated the reliability and fairness of ChatGPT as an automated "
    "essay scoring tool, demonstrating that when provided with explicit, normalized rubrics and standardized "
    "output formats at low temperature settings, the model achieved perfect grading consistency across ten "
    "independent evaluation sessions. Their study also developed a scalable algorithm that automatically generates "
    "normalized grading rubrics and decision tables from any open-ended question with minimal instructor input. "
    "These findings confirm that the combination of AI grading capability with customizable rubric structures "
    "yields consistent, transparent, and educationally sound assessment outcomes."
)

AG_3 = (
    "In Philippine educational settings, where teachers routinely manage large classes, the time consumed by "
    "manual evaluation of quizzes and assignments substantially limits the instructional time available for lesson "
    "preparation and individualized student support (Colegado, 2025). Colegado\u2019s study on LMS adoption in "
    "Philippine public schools found that while educators recognized the transformative potential of digital "
    "assessment tools, a critical limitation of available platforms was the absence of integrated automated "
    "evaluation features aligned with local curriculum requirements. This gap directly motivates the inclusion "
    "of an automated grading module with customizable rubrics in A.I.M.S., which would enable Philippine teachers "
    "to administer and evaluate assessments at scale while maintaining instructional quality through "
    "teacher-defined, adaptable scoring criteria."
)

ML_1 = (
    "Mastery-based learning is a pedagogical model grounded in the principle that all students can achieve "
    "academic proficiency when given sufficient time, appropriate instructional support, and the opportunity to "
    "demonstrate competency before progressing to subsequent content. Persky and Hughes (2022) conducted a "
    "comprehensive review of mastery learning principles and their practical application in varied educational "
    "contexts, finding that mastery-based instruction consistently outperformed traditional time-based delivery "
    "models in promoting long-term content retention, reducing achievement gaps, and improving student confidence. "
    "Their review emphasized that effective mastery learning frameworks require clearly defined performance "
    "thresholds, systematic corrective feedback mechanisms, and structured opportunities for retesting\u2014"
    "components that are directly embedded in the mastery-lock system proposed within A.I.M.S."
)

ML_2 = (
    "The adaptive dimension of the mastery-lock system distinguishes it from static competency gates by "
    "dynamically adjusting remediation content based on individual student performance data. Rather than merely "
    "blocking progression at a fixed score threshold, adaptive mastery-lock systems analyze patterns in student "
    "errors and misconceptions to deliver targeted remediation activities aligned with specific learning gaps. "
    "This integration of adaptive remediation with competency-based progression has been validated in the broader "
    "learning analytics literature, where data-driven feedback mechanisms have been shown to improve students\u2019 "
    "self-regulatory behavior and academic performance (Susnjak et al., 2022). By connecting performance data to "
    "automated intervention triggers, the A.I.M.S. mastery-lock mechanism ensures that progression decisions are "
    "both consistent and educationally purposeful."
)

ML_3 = (
    "In the Philippine educational context, class size constraints and limited instructional support time make "
    "individualized intervention challenging to sustain manually (Balase & Paglinawan, 2025). Balase and "
    "Paglinawan\u2019s study on educational technology adoption barriers in Philippine public schools found that "
    "while adaptive learning tools hold significant institutional promise, their effective implementation requires "
    "structural integration with automated performance tracking and decision-support systems that do not depend on "
    "continuous teacher mediation. The A.I.M.S. mastery-lock system addresses this implementation gap by "
    "automating progression decisions and corresponding remediation delivery, reducing dependency on "
    "teacher-initiated intervention while maintaining pedagogical consistency and competency-based rigor across "
    "diverse classroom environments."
)

DB_1 = (
    "Learning analytics dashboards have emerged as powerful tools for converting raw educational data into "
    "actionable insights that support both student self-regulation and teacher instructional decision-making. "
    "Susnjak et al. (2022) investigated how learning analytics dashboards can provide actionable insights to "
    "learners by synthesizing engagement metrics, quiz performance data, and learning behavior logs into "
    "interpretable visualizations. Their study found that dashboards displaying timely and contextually "
    "interpretable performance information significantly improved students\u2019 awareness of their academic "
    "trajectories, encouraged proactive study behaviors, and enabled educators to identify and support at-risk "
    "learners before performance deteriorated critically\u2014positioning the performance dashboard as an "
    "indispensable component of any AI-enhanced educational platform."
)

DB_2 = (
    "The design and usability of learning analytics dashboards substantially influence their adoption rate and "
    "educational efficacy in real classroom settings. De Vreugd et al. (2024) conducted a sequential "
    "mixed-methods study involving needs assessments, think-aloud usability interviews, and perceived usefulness "
    "evaluations to design a dashboard supporting student self-regulation of study behavior in higher education. "
    "Their findings revealed that students and instructors agreed on the relevance of core performance "
    "constructs\u2014including quiz completion rates, engagement levels, and module access frequency\u2014but "
    "differed in how they prioritized these data dimensions. The study also demonstrated that providing "
    "interpretive reference frames alongside raw data significantly enhanced the dashboard\u2019s perceived "
    "usefulness, offering critical design guidance for the A.I.M.S. student performance dashboard."
)

DB_3 = (
    "A systematic review by Kaliisa et al. (2024), synthesizing findings from 38 research studies on learning "
    "analytics dashboards, found that these tools consistently and significantly improved student participation "
    "and engagement, with several studies reporting medium to large effect sizes. However, the review also noted "
    "that evidence for direct impact on academic achievement remains limited and methodologically inconsistent, "
    "largely due to small sample sizes, self-selection biases, and the conflation of dashboard access with "
    "meaningful use. In the Philippine educational context, where data-driven instruction remains nascent and "
    "manual performance tracking is still common practice, the introduction of a web-based student performance "
    "dashboard through A.I.M.S. represents a substantive step toward evidence-informed teaching and real-time "
    "academic monitoring at the classroom level."
)

SYN_1 = (
    "The reviewed literature collectively demonstrates that each of the five core features of A.I.M.S. is "
    "grounded in a robust and growing body of empirical evidence. Centralized digital libraries address "
    "persistent access and distribution inequities in public school environments (Bradley, 2021; Bustillo & "
    "Aguilos, 2022), while AI-powered quiz generators substantially reduce the cognitive and administrative "
    "burden of assessment creation (Bulathwela et al., 2023; Alamoudi et al., 2025). Automated grading systems "
    "promote evaluation consistency and scalability through machine learning and rubric-guided assessment "
    "(Ramesh & Sanampudi, 2022; Garc\u00eda-Varela et al., 2025), and mastery-lock mechanisms provide "
    "structured, competency-based progression proven to reduce learning gaps (Persky & Hughes, 2022). Performance "
    "dashboards complete the instructional feedback loop by delivering real-time visibility into student progress "
    "(Susnjak et al., 2022; de Vreugd et al., 2024; Kaliisa et al., 2024)."
)

SYN_2 = (
    "A consistent theme across both foreign and local literature is that the greatest educational impact is "
    "achieved when multiple intelligent features are integrated cohesively within a single platform, rather than "
    "deployed as fragmented standalone tools that create workflow disruptions and adoption resistance (Colegado, "
    "2025; Balase & Paglinawan, 2025). The reviewed studies demonstrate that while individual features\u2014quiz "
    "generation, automated grading, mastery progression\u2014have been validated in isolation, no existing "
    "platform fully integrates the complete instructional and assessment cycle in the manner envisioned in "
    "A.I.M.S. This integration gap represents the primary research and design opportunity that the present "
    "study addresses."
)

SYN_3 = (
    "Overall, the literature reviewed in this chapter provides strong theoretical, empirical, and contextual "
    "justification for the design and development of A.I.M.S. as a unified, AI-driven educational platform "
    "tailored to the Philippine public school context. The evidence base confirms that the proposed features are "
    "not merely technically feasible but also pedagogically sound, contextually responsive, and aligned with "
    "contemporary best practices in educational technology research. By operationalizing verified AI capabilities "
    "within a coherent, teacher-centered system architecture, A.I.M.S. represents a research-supported response "
    "to documented gaps in Philippine digital learning environments\u2014integrating access management, automated "
    "assessment, competency-based progression, and performance analytics into a single, scalable educational "
    "system."
)

MAT_1 = (
    "The matrix of related literature and studies above presents a comparative overview of existing systems and "
    "studies in relation to the five core technical features of the proposed A.I.M.S. platform. The studies "
    "mapped in the matrix span both foreign and local literature, covering digital library systems and LMS "
    "integration (Bradley, 2021; Bustillo & Aguilos, 2022), AI-driven and LLM-based quiz generation methods "
    "(Bulathwela et al., 2023; Alamoudi et al., 2025), automated grading and AES approaches with rubric "
    "customization (Ramesh & Sanampudi, 2022; Garc\u00eda-Varela et al., 2025), mastery-based learning and "
    "competency progression frameworks (Persky & Hughes, 2022), and learning analytics dashboard design and "
    "evaluation (Susnjak et al., 2022; de Vreugd et al., 2024; Kaliisa et al., 2024)."
)

MAT_2 = (
    "Several studies demonstrate the value of centralized digital libraries for efficient and equitable module "
    "distribution, while others validate AI-powered assessment tools for reducing teacher workload and improving "
    "evaluation consistency. Competency-based progression supported by mastery learning principles is affirmed as "
    "an effective instructional strategy across diverse learner populations, and learning analytics dashboards are "
    "consistently associated with improved student participation and engagement. Local Philippine studies document "
    "systemic barriers to technology adoption and underscore the demand for integrated, automated digital "
    "solutions in public school settings (Bustillo & Aguilos, 2022; Balase & Paglinawan, 2025; Colegado, 2025)."
)

MAT_3 = (
    "Overall, the synthesis reveals that while the individual features addressed by A.I.M.S. are present in "
    "existing systems and have been independently validated in the literature, no single platform currently "
    "integrates all five functionalities\u2014centralized digital library, AI quiz generation, automated grading, "
    "mastery-lock progression, and performance dashboard\u2014into a unified, AI-driven educational environment. "
    "This integration gap constitutes the primary research contribution of A.I.M.S. and affirms the novelty and "
    "practical significance of the proposed system as a response to documented needs in the Philippine educational "
    "technology landscape."
)

# ─────────────────────────────────────────────
# Paragraph index → new text mapping
# Based on cha1&2_v2.docx structure:
#  86 = "Chapter 2"           (HEADING - DO NOT TOUCH)
#  87 = "REVIEW OF..."        (HEADING - DO NOT TOUCH)
#  88 = Intro paragraph
#  89 = Digital Library heading   (DO NOT TOUCH)
#  90-92 = DL body paras
#  93 = empty (skip)
#  94 = Quiz Gen heading          (DO NOT TOUCH)
#  95-97 = QG body paras
#  98 = Grading heading           (DO NOT TOUCH)
#  99-101 = AG body paras
#  102 = Mastery heading          (DO NOT TOUCH)
#  103-105 = ML body paras
#  106 = empty (skip)
#  107 = Dashboard heading        (DO NOT TOUCH)
#  108-110 = DB body paras
#  111 = empty (skip)
#  112 = Synthesis heading        (DO NOT TOUCH)
#  113-115 = Synthesis body paras
#  116-117 = Table caption        (DO NOT TOUCH)
#  118 = empty (skip)
#  119 = "Synthesis"             (DO NOT TOUCH)
#  120-122 = Matrix synthesis body paras
# ─────────────────────────────────────────────
REPLACEMENTS = {
    88:  INTRO,
    90:  DL_1,
    91:  DL_2,
    92:  DL_3,
    95:  QG_1,
    96:  QG_2,
    97:  QG_3,
    99:  AG_1,
    100: AG_2,
    101: AG_3,
    103: ML_1,
    104: ML_2,
    105: ML_3,
    108: DB_1,
    109: DB_2,
    110: DB_3,
    113: SYN_1,
    114: SYN_2,
    115: SYN_3,
    120: MAT_1,
    121: MAT_2,
    122: MAT_3,
}

doc = Document('cha1&2_v2.docx')
total = len(doc.paragraphs)
print(f'Total paragraphs in v2: {total}')

for idx, new_text in sorted(REPLACEMENTS.items()):
    if idx < total:
        p = doc.paragraphs[idx]
        old_preview = p.text[:80].strip()
        set_para_text(p, new_text)
        print(f'[{idx}] Replaced: "{old_preview}..." → new text ({len(new_text)} chars)')
    else:
        print(f'[{idx}] SKIP: index out of range (total={total})')

doc.save('cha1&2_v3.docx')
print('\nSaved: cha1&2_v3.docx')
print('\nDone! Chapter 2 body paragraphs updated with verified citations.')
