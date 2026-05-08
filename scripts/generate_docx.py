from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Page setup: short bond paper (8.5 x 11 inches), 1-inch margins ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# ── Helper functions ──
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'

def add_heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Arial'
    p.space_before = Pt(18)
    p.space_after = Pt(6)

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5)

def add_body_no_indent(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.space_after = Pt(6)

def add_screenshot_box(label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[Insert Screenshot: {label}]')
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(128, 128, 128)
    p.space_before = Pt(12)
    p.space_after = Pt(12)

def add_screenshot_instruction(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(100, 100, 100)
    p.space_after = Pt(8)

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2500' * 50)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(8)

# ======================================================================
# COVER
# ======================================================================

add_title('A.I.M.S.')
add_title('(Automated Intervention and Mastery System)')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('System Update Documentation')
run.font.size = Pt(13)
run.font.name = 'Arial'

doc.add_paragraph()
add_subtitle('Submitted by:')
add_subtitle('Mikhoel Runo B. Orbigoso')
add_subtitle('Patrick R. Sagum')
add_subtitle('Maika T. Zambra')
doc.add_paragraph()
add_subtitle('Course: Capstone Project')
add_subtitle('Program: Bachelor of Science in Information Technology')
add_subtitle('Carlos Hilado Memorial State University \u2014 Alijis Campus, Bacolod City')

doc.add_page_break()

# ======================================================================
# SYSTEM OVERVIEW
# ======================================================================

add_heading2('System Overview')

add_body(
    'A.I.M.S. (Automated Intervention and Mastery System) is a web-based Learning Management System (LMS) '
    'designed as one of three interconnected modules within a School ERP (Enterprise Resource Planning) system '
    'for Philippine public high schools. The system\'s user interface and experience design is modeled after '
    'Google Classroom, ensuring a minimal learning curve for teachers and students who are already familiar '
    'with Google\'s educational ecosystem.'
)

add_body_no_indent('The three ERP modules work together as follows:')

add_body_no_indent(
    '\u2022 Schedule Module (School ERP) \u2014 Manages class schedules, section assignments, and teacher loads. '
    'The LMS fetches schedule data from this module to auto-enroll teachers and students into their '
    'respective classes, eliminating the need for manual enrollment.'
)
add_body_no_indent(
    '\u2022 A.I.M.S. LMS Module \u2014 Provides AI-powered quiz generation, automated grading, module distribution, '
    'mastery-based progression, and performance tracking for asynchronous learning activities (online quizzes, '
    'assignments, and materials).'
)
add_body_no_indent(
    '\u2022 Class Records Module (School ERP) \u2014 Maintains the complete official class record of student grades. '
    'Since A.I.M.S. handles partial grades from asynchronous activities, the LMS automatically updates and '
    'syncs these grades to the Class Records module, ensuring a unified and up-to-date academic record.'
)

add_separator()

# ======================================================================
# PAGES — Only features that are ACTUALLY IMPLEMENTED
# ======================================================================

# ── PAGE 1: Login ──
add_heading2('Page 1: Authentication \u2014 Login Page with Google Sign-In Integration')
add_screenshot_box('Login Page')
add_screenshot_instruction(
    'HOW TO CAPTURE: Go to localhost:5173/login while logged out. '
    'Screenshot the full page showing the email/password fields and the "Sign in with Google" button.'
)
add_body(
    'The A.I.M.S. login page serves as the entry point to the system. It features both a traditional '
    'email/password login form and an integrated Google Sign-In button, allowing teachers and students to '
    'authenticate using their existing Google accounts. This integration eliminates the need for users to '
    'create and manage separate credentials, streamlining the onboarding process. Upon clicking the '
    '"Sign in with Google" button, users are redirected through Google\'s OAuth 2.0 authentication flow, '
    'which securely verifies their identity and returns them to the system. The design follows Google '
    'Classroom\'s familiar authentication pattern, reducing the learning curve for first-time users.'
)

# ── PAGE 2: Role Selection ──
add_heading2('Page 2: Authentication \u2014 Role Selection (First-time Google Sign-In)')
add_screenshot_box('Role Selection Page')
add_screenshot_instruction(
    'HOW TO CAPTURE: Sign in with a brand-new Google account that has never used the system before. '
    'After Google auth completes, you are redirected to /role-select. Screenshot the page showing '
    'the "Teacher" and "Student" role cards/buttons.'
)
add_body(
    'When a user signs in with Google for the first time, the system prompts them to select their role '
    '\u2014 either Teacher or Student. This one-time selection determines the user\'s access level, interface '
    'layout, and feature availability throughout the system. Teachers gain access to course management, quiz '
    'generation, material uploads, and grading tools, while students see enrolled courses, assignments, and '
    'their grades. After selecting a role, the user is directed to the Home page and does not see this '
    'selection screen again on subsequent logins.'
)

# ── PAGE 3: Teacher Home ──
add_heading2('Page 3: Home Page \u2014 Class Cards (Teacher View)')
add_screenshot_box('Teacher Home Page')
add_screenshot_instruction(
    'HOW TO CAPTURE: Log in as a teacher account. You land on /home. Screenshot the full page showing: '
    '(1) the class cards in the main area \u2014 each card has a colored banner with subject name, '
    'grade-level & section label (e.g. "Grade 7 \u2014 1"), class code, and student count; '
    '(2) the left sidebar expanded, showing "Teaching" header with each class listed below it.'
)
add_body(
    'The Home page displays all classes assigned to the teacher as visually distinct class cards, styled after '
    'Google Classroom\'s card-based layout. Each card shows the subject name, grade level and section (e.g., '
    '"Grade 7 \u2014 1"), the class code, and the number of enrolled students. The cards use a color-coded '
    'banner for quick visual identification. Teachers can click on any card to enter the course detail page. '
    'The sidebar navigation on the left provides quick access to Home, Schedule, Materials, and Quiz Builder. '
    'Under the "Teaching" section in the sidebar, each active class is also listed for direct navigation.'
)

# ── PAGE 4: Student Home ──
add_heading2('Page 4: Home Page \u2014 Class Cards (Student View)')
add_screenshot_box('Student Home Page')
add_screenshot_instruction(
    'HOW TO CAPTURE: Log in as a student account. You land on /home. Screenshot the full page showing: '
    '(1) enrolled class cards \u2014 each card has a colored banner with subject, grade & section, class code, '
    'and teacher name; '
    '(2) the left sidebar showing "Enrolled" header with listed classes and grade/section beneath each.'
)
add_body(
    'The student Home page follows the same card-based layout as the teacher view for visual consistency. '
    'Each card displays the subject name, grade level and section, the class code, and the teacher\'s name. '
    'Students can click any card to view course content, assignments, and their grades. The sidebar shows '
    'navigation links including Home, Schedule, and Materials. Under the "Enrolled" section in the sidebar, '
    'each class the student is enrolled in is listed, with the grade and section shown beneath each class name.'
)

# ── PAGE 5: Teacher Stream ──
add_heading2('Page 5: Course Detail \u2014 Stream Tab (Teacher View)')
add_screenshot_box('Teacher Stream Tab')
add_screenshot_instruction(
    'HOW TO CAPTURE: As teacher, click any class card from /home. The Stream tab is selected by default. '
    'Screenshot showing: (1) the colored course banner at top with subject, grade/section, code, student count; '
    '(2) the tab bar with "Stream" highlighted; (3) the "Announce something to your class..." composer box; '
    '(4) the activity feed below showing posted quizzes and materials with profile pictures and timestamps; '
    '(5) the right sidebar with class code card, upcoming work section, and mastery-lock alerts (if any). '
    'PREREQUISITE: Have at least 1 announcement, 1 quiz, and 1 material already posted in this course.'
)
add_body(
    'The Stream tab serves as the main activity feed for the course, combining all announcements, posted '
    'quizzes, and uploaded materials into a single chronological timeline. At the top of the stream, teachers '
    'can compose and post announcements to communicate with the class. Below the composer, each feed item '
    'displays the poster\'s Google profile picture, name, a description of the action (e.g., "posted a new '
    'assignment: Quiz 1" or "posted a new material: Lesson 2"), and a relative timestamp. The right sidebar '
    'shows a class code card for easy sharing, upcoming work, and mastery-lock alerts when students have been '
    'locked due to failed assessments. The breadcrumb navigation at the top of the page displays the full path: '
    'Classes > Subject (Grade Level \u2014 Section).'
)

# ── PAGE 6: Student Stream ──
add_heading2('Page 6: Course Detail \u2014 Stream Tab (Student View)')
add_screenshot_box('Student Stream Tab')
add_screenshot_instruction(
    'HOW TO CAPTURE: As student, click a class card from /home. Stream tab is active. Screenshot showing: '
    '(1) the course banner; (2) the feed with announcement posts (with "Add comment" link), quiz posts '
    '(with FileText icon), and material posts (with FileText icon and "View Material" link); '
    '(3) the right sidebar with "Upcoming Work" listing pending quizzes. '
    'PREREQUISITE: The course should have at least 1 announcement, 1 published quiz, and 1 material.'
)
add_body(
    'The student Stream tab displays the same chronological activity feed \u2014 announcements, quizzes, and '
    'materials posted by the teacher. Students can view announcements with the ability to add comments, click '
    'on quiz posts to navigate to the quiz detail page, and click on material posts to view uploaded learning '
    'resources. The right sidebar displays an "Upcoming Work" card listing quizzes that the student has not yet '
    'completed. The Google Classroom-inspired layout ensures students can quickly scan recent class activity '
    'and identify pending tasks.'
)

# ── PAGE 7: Teacher Classwork ──
add_heading2('Page 7: Course Detail \u2014 Classwork Tab (Teacher View)')
add_screenshot_box('Teacher Classwork Tab')
add_screenshot_instruction(
    'HOW TO CAPTURE: As teacher inside a course, click the "Classwork" tab in the tab bar. Screenshot showing: '
    '(1) topic accordion sections \u2014 at least one topic expanded; '
    '(2) quizzes listed under the topic with title, type badge ("AI-generated" or "Manual"), question count, date; '
    '(3) materials listed with file icon and title. '
    'PREREQUISITE: Create at least one topic with 2+ quizzes and 1+ material underneath it.'
)
add_body(
    'The Classwork tab organizes course content by topics. Teachers can create topic groupings and assign '
    'quizzes and materials under each topic. Each quiz entry shows its title, type (AI-generated or manual), '
    'number of questions, and creation date. Materials are listed alongside quizzes within their respective '
    'topics. This structured view gives teachers a comprehensive overview of all course content and assessments, '
    'enabling them to manage the curriculum flow and ensure alignment with DepEd competencies.'
)

# ── PAGE 8: Student Classwork ──
add_heading2('Page 8: Course Detail \u2014 Classwork Tab (Student View)')
add_screenshot_box('Student Classwork Tab')
add_screenshot_instruction(
    'HOW TO CAPTURE: As student inside a course, click "Classwork" tab. Screenshot showing: '
    '(1) topic accordion with quizzes showing status badges \u2014 capture a mix of statuses: '
    'one "Not yet taken" (gray), one "Passed" (green badge), one "Failed" or "Locked" (red/lock icon); '
    '(2) materials/resources listed under the topic. '
    'PREREQUISITE: Student must have attempted at least one quiz (pass) and one quiz (fail) to show varied statuses. '
    'If mastery-lock is active, a locked quiz will show a lock icon.'
)
add_body(
    'The student Classwork tab displays the same topic-based organization but with a student-focused '
    'perspective. Each quiz shows the student\'s submission status \u2014 whether it is "Not yet taken," '
    '"Passed," "Failed," or "Locked" due to the mastery-lock system. The mastery-lock feature restricts '
    'access to subsequent quizzes until the student achieves a passing score (70% or higher) on the previous '
    'assessment. When a student fails, the system automatically generates and deploys a remediation quiz. '
    'Resources and materials uploaded by the teacher are also accessible within each topic.'
)

# ── PAGE 9: Teacher People ──
add_heading2('Page 9: Course Detail \u2014 People Tab (Teacher View)')
add_screenshot_box('Teacher People Tab')
add_screenshot_instruction(
    'HOW TO CAPTURE: As teacher inside a course, click "People" tab. Screenshot showing: '
    '(1) "Teacher" section at the top with your name and Google profile picture; '
    '(2) "Students" section below listing enrolled students with profile pictures, names, and join dates. '
    'PREREQUISITE: The course should have at least 3\u20135 enrolled students for a meaningful screenshot.'
)
add_body(
    'The People tab displays all members of the course. The teacher\'s section shows the teacher\'s name '
    'and Google profile picture. Below that, the Students section lists all enrolled students with their '
    'profile pictures, names, and the date they joined the class. This tab provides a quick overview of '
    'class composition. Student enrollment is handled automatically through the Schedule module of the School '
    'ERP \u2014 when a student is assigned to a section in the schedule system, they are auto-enrolled in '
    'the corresponding LMS course.'
)

# ── PAGE 10: Student People ──
add_heading2('Page 10: Course Detail \u2014 People Tab (Student View)')
add_screenshot_box('Student People Tab')
add_screenshot_instruction(
    'HOW TO CAPTURE: As student inside a course, click "People" tab. Screenshot showing: '
    '(1) the teacher at top with Google profile picture and name; '
    '(2) classmates listed below with profile pictures and names; '
    '(3) the currently logged-in student should have a "You" badge next to their name.'
)
add_body(
    'The student view of the People tab shows the teacher at the top with their Google profile picture and '
    'name, followed by a list of classmates. The current user is tagged with a "You" badge for easy '
    'identification. This tab allows students to see their teacher and fellow classmates, fostering a sense '
    'of classroom community in the online environment.'
)

# ── PAGE 11: Teacher Grades ──
add_heading2('Page 11: Course Detail \u2014 Grades Tab / Gradebook (Teacher View)')
add_screenshot_box('Teacher Gradebook')
add_screenshot_instruction(
    'HOW TO CAPTURE: As teacher inside a course, click "Grades" tab. Screenshot the gradebook table showing: '
    '(1) student names in the left column; (2) quiz titles as column headers; '
    '(3) individual scores in each cell with color-coded backgrounds (green = passed, red = failed, gray = not taken); '
    '(4) sort/filter options at the top if visible. '
    'PREREQUISITE: At least 3 students must have submitted at least 2 quizzes each, with a mix of pass/fail results.'
)
add_body(
    'The Grades tab presents a comprehensive gradebook view for the teacher, displaying all students and '
    'their scores across all quizzes in the course. The table format shows each student\'s name, individual '
    'quiz scores, passing status, and overall course performance. Teachers can review which students have '
    'passed, failed, or not yet taken each assessment. Color-coded indicators (green for passing, red for '
    'failing) provide at-a-glance performance tracking. This partial grade data from asynchronous activities '
    'is automatically synced to the Class Records module of the School ERP, ensuring the official academic '
    'record stays current without manual data entry by the teacher.'
)

# ── PAGE 12: Student Grades ──
add_heading2('Page 12: Course Detail \u2014 Grades Tab (Student View)')
add_screenshot_box('Student Grades Tab')
add_screenshot_instruction(
    'HOW TO CAPTURE: As student inside a course, click "Grades" tab. Screenshot the grade table showing: '
    '(1) quiz titles in a column; (2) your score for each quiz; (3) pass/fail status; '
    '(4) DepEd descriptors column (Outstanding / Very Satisfactory / Satisfactory / Fairly Satisfactory / '
    'Did Not Meet Expectations). '
    'PREREQUISITE: The student should have at least one passed quiz and one failed quiz to show contrasting descriptors.'
)
add_body(
    'The student Grades tab shows the individual student\'s own quiz scores, submission statuses, and '
    'overall course performance. Each row displays the quiz title, the student\'s score, and whether they '
    'passed or failed with DepEd performance descriptors (Outstanding, Very Satisfactory, Satisfactory, '
    'Fairly Satisfactory, Did Not Meet Expectations). This transparency empowers students to monitor their '
    'own academic progress and identify areas needing improvement, supporting self-directed learning.'
)

# ── PAGE 13: Quiz Detail ──
add_heading2('Page 13: Quiz Detail Page')
add_screenshot_box('Quiz Detail Page \u2014 Teacher View')
add_screenshot_instruction(
    'HOW TO CAPTURE: As teacher, open a course \u2192 Classwork tab \u2192 click on any quiz title. '
    'Screenshot showing: quiz title, creator name with profile picture, creation date, number of questions, '
    'time limit (if set), passing threshold percentage, quiz type badge ("AI-generated" or "Manual"), '
    'and the action buttons (Edit, Publish/Close, Review Submissions).'
)
add_screenshot_box('Quiz Detail Page \u2014 Student View')
add_screenshot_instruction(
    'HOW TO CAPTURE: As student, open a course \u2192 Classwork tab \u2192 click on an available quiz. '
    'Screenshot showing: quiz title, creator info, question count, time limit, passing threshold, and '
    'the "Take Quiz" button (if not yet taken) OR the score summary (if already completed).'
)
add_body(
    'The Quiz detail page shows complete quiz information including the title, creator (with profile picture), '
    'creation date, number of questions, time limit, passing threshold, and quiz type (AI-generated or '
    'manual). For students, a "Take Quiz" button appears if the quiz has not been submitted, or a score '
    'summary is shown if already completed. For teachers, the page provides options to edit, publish/close, '
    'or review student submissions. This page serves as the central hub for quiz interaction for both user roles.'
)

# ── PAGE 14: Take Quiz ──
add_heading2('Page 14: Take Quiz Page (Student View)')
add_screenshot_box('Student Taking a Quiz')
add_screenshot_instruction(
    'HOW TO CAPTURE: As student, go to an available quiz and click "Take Quiz." Screenshot the quiz-taking '
    'interface showing: (1) the timer countdown at the top (if time limit is set); '
    '(2) list of questions with answer choices \u2014 radio buttons for Multiple Choice / True or False, '
    'text input for Short Answer; (3) some answers already selected (click a few options first before screenshotting); '
    '(4) the "Submit Quiz" button at the bottom. '
    'TIP: Take the screenshot AFTER selecting answers for 2\u20133 questions so it shows the interactive state.'
)
add_body(
    'The Take Quiz page presents questions in a scrollable list format. A timer is displayed if the teacher '
    'set a time limit. Students select their answers for multiple-choice and true/false questions, or type '
    'responses for short-answer items. Upon completion, students submit the quiz for immediate automated '
    'grading. The system evaluates responses against the answer key and customizable rubrics, computes the '
    'score instantly, and determines whether the student has achieved mastery (70% or higher).'
)

# ── PAGE 15: Quiz Results ──
add_heading2('Page 15: Quiz Results Page')
add_screenshot_box('Quiz Results \u2014 Passed')
add_screenshot_instruction(
    'HOW TO CAPTURE (Passed): After submitting a quiz where you scored 70%+, screenshot the results page showing: '
    '(1) the overall score and percentage at the top; (2) a green "Passed" badge; '
    '(3) per-question breakdown showing correct answers in green and any wrong answers in red.'
)
add_screenshot_box('Quiz Results \u2014 Failed (Mastery Lock Triggered)')
add_screenshot_instruction(
    'HOW TO CAPTURE (Failed): After submitting a quiz where you scored below 70%, screenshot showing: '
    '(1) the overall score and percentage; (2) a red "Failed" badge; '
    '(3) the mastery-lock notice indicating the next quiz is locked; '
    '(4) the remedial quiz link/notice (if visible). This demonstrates the mastery-lock system in action.'
)
add_body(
    'After quiz submission, the Results page displays the student\'s score, pass/fail status, and a '
    'per-question breakdown showing which answers were correct and incorrect. This immediate feedback loop '
    'is central to the A.I.M.S. learning model. If the student fails, the mastery-lock system is triggered '
    '\u2014 the next quiz in the topic sequence becomes locked, and a remediation quiz is automatically '
    'generated and made available. This dynamic remediation process ensures students address knowledge gaps '
    'before proceeding.'
)

# ── PAGE 16: AI Quiz Generation ──
add_heading2('Page 16: AI-Powered Quiz Generation (Teacher View)')
add_screenshot_box('Generate Quiz Modal \u2014 AI Tab')
add_screenshot_instruction(
    'HOW TO CAPTURE: As teacher, open a course \u2192 click the "Generate Quiz" button in the course banner area. '
    'A modal dialog opens. Make sure the "AI" tab is selected (not "Manual"). Screenshot the modal showing: '
    '(1) the source document dropdown (select a previously uploaded PDF/DOCX/TXT); '
    '(2) question type radio buttons (Multiple Choice / True or False / Short Answer); '
    '(3) number of questions input; (4) difficulty level selector; (5) the "Generate" button. '
    'PREREQUISITE: Upload at least one material/document to the course first so the source dropdown is populated.'
)
add_body(
    'The AI-Powered Quiz Generator allows teachers to create assessments automatically from uploaded '
    'learning materials. Teachers select a source document (PDF, DOCX, or TXT), choose the question type '
    '(multiple-choice, true/false, or short answer), set the number of questions and difficulty level, and '
    'click generate. The system uses the Google Gemini API to analyze the document content and produce '
    'curriculum-aligned questions. Generated quizzes go through a review process \u2014 teachers can edit, '
    'add, remove, or rearrange questions before publishing. This instructor validation control ensures that '
    'AI-generated content meets pedagogical standards before reaching students.'
)

# ── PAGE 17: Quiz Builder ──
add_heading2('Page 17: Quiz Builder / Editor (Teacher View)')
add_screenshot_box('Quiz Builder Page')
add_screenshot_instruction(
    'HOW TO CAPTURE: As teacher, click "Quiz Builder" in the left sidebar. Screenshot the page showing: '
    '(1) the quiz list panel with status filter tabs (All / Draft / Published / Closed); '
    '(2) the course filter dropdown; '
    '(3) at least one quiz entry in the list. '
    'If you click on a quiz, the right panel shows the quiz editor \u2014 capture that state if possible. '
    'PREREQUISITE: Have at least 2\u20133 quizzes in different statuses (1 Draft, 1 Published) to show filtering.'
)
add_body(
    'The Quiz Builder provides teachers with a dedicated interface for managing and editing quizzes across '
    'all their courses. The page features a filterable quiz list with status filters (All, Draft, Published, '
    'Closed) and course-based filtering. Teachers can review AI-generated questions, manually edit individual '
    'items, reorder questions, adjust passing thresholds, and publish quizzes when ready. This serves as the '
    'teacher\'s central workspace for assessment management and instructor validation of AI-generated content.'
)

# ── PAGE 18: Material Detail ──
add_heading2('Page 18: Material / Resource Detail Page')
add_screenshot_box('Material Detail Page')
add_screenshot_instruction(
    'HOW TO CAPTURE: Navigate to a material by clicking it from the Stream feed or Classwork tab. '
    'Screenshot showing: (1) material title; (2) file name and file type icon; (3) upload date; '
    '(4) uploader name with Google profile picture; (5) download button; '
    '(6) breadcrumb at top: "Classes > Subject (Grade \u2014 Section) > Material Title". '
    'Works the same for both teacher and student views.'
)
add_body(
    'The Material detail page displays the uploaded resource\'s title, file name, upload date, and the '
    'uploader\'s name with their Google profile picture. Students and teachers can view file details and '
    'download the material for offline study. The breadcrumb navigation shows the full path: '
    'Classes > Subject (Grade Level \u2014 Section) > Material Title. This centralized digital library '
    'feature ensures all learning modules are accessible from a single organized location.'
)

# ── PAGE 19: Materials Library ──
add_heading2('Page 19: Materials / Resources Library Page')
add_screenshot_box('Materials Library Page')
add_screenshot_instruction(
    'HOW TO CAPTURE: Click "Materials" in the left sidebar. Screenshot the library page showing: '
    'file cards with titles, file types, and dates. For teachers, this shows materials uploaded across '
    'all their courses. For students, this shows materials from all enrolled courses. '
    'PREREQUISITE: Have at least 3\u20134 materials uploaded across courses for a meaningful screenshot.'
)
add_body(
    'The Materials page provides a centralized view of all learning resources available to the user. For '
    'teachers, this shows all materials they have uploaded across their courses. For students, it displays '
    'materials from their enrolled courses. Resources can be browsed, searched, and accessed directly from '
    'this page, serving as the system\'s digital library for module distribution and management.'
)

# ── PAGE 20: Sidebar and Layout ──
add_heading2('Page 20: Sidebar Navigation and System Layout')
add_screenshot_box('Sidebar Navigation \u2014 Teacher View')
add_screenshot_instruction(
    'HOW TO CAPTURE: As teacher on /home, make sure the sidebar is expanded (click hamburger icon if collapsed). '
    'Screenshot the full sidebar showing: (1) school name and logo at the very top; '
    '(2) nav items: Home, Schedule, Materials, Quiz Builder; '
    '(3) "Teaching" section listing each class with subject name and grade/section; '
    '(4) user profile area at the very bottom with Google profile picture, name, and email.'
)
add_screenshot_box('Sidebar Navigation \u2014 Student View')
add_screenshot_instruction(
    'HOW TO CAPTURE: As student on /home with sidebar expanded. Screenshot showing: '
    '(1) school name/logo at top; '
    '(2) nav items: Home, Schedule, Materials; '
    '(3) "Enrolled" section listing each class with subject, grade, and section; '
    '(4) user profile at bottom with Google profile picture, name, and email.'
)
add_body(
    'The sidebar navigation provides persistent access to all major system sections. The top section shows '
    'the school name and logo. Navigation items include Home, Schedule, Materials, and Quiz Builder (for '
    'teachers). Below the navigation, the sidebar lists the user\'s active classes \u2014 labeled "Teaching" '
    'for teachers and "Enrolled" for students \u2014 with each class showing the subject name, grade level, '
    'and section. The bottom of the sidebar displays the current user\'s Google profile picture, name, and '
    'email. This layout mirrors Google Classroom\'s navigation paradigm, ensuring immediate familiarity for '
    'users transitioning from Google\'s educational tools.'
)

# ── PAGE 21: Breadcrumbs ──
add_heading2('Page 21: Course Banner and Breadcrumb Navigation')
add_screenshot_box('Course Banner with Breadcrumbs')
add_screenshot_instruction(
    'HOW TO CAPTURE: As teacher or student, open any course. Take a close-up/cropped screenshot of the top section showing: '
    '(1) the breadcrumb bar in the header: "Classes > English (Grade 7 \u2014 1)" with the hamburger menu icon on the left; '
    '(2) the colored course banner just below it with subject name, grade/section, class code, '
    'and teacher name (student) or enrollment count (teacher). '
    'TIP: You can zoom in or crop so the breadcrumbs and banner are clearly readable.'
)
add_body(
    'Each course page features a color-coded banner displaying the subject name, grade level and section, '
    'class code, and additional context (teacher name for students; enrollment count for teachers). Above '
    'the banner, the breadcrumb navigation bar shows the hierarchical path \u2014 "Classes > English '
    '(Grade 7 \u2014 1)" \u2014 enabling one-click navigation back to the Home page. This breadcrumb '
    'system is consistent across both teacher and student views and updates dynamically based on the current '
    'page context, including nested pages like material and quiz detail views.'
)

doc.add_page_break()

# ======================================================================
# OBJECTIVES TABLE
# ======================================================================

add_heading2('Summary of System Objectives Achievement')

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Objective', 'Feature Implemented', 'Status']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Arial'

rows_data = [
    [
        '1.1 Centralized digital library for module distribution and management',
        'Materials library page, per-course resource uploads via course banner, material detail pages with download',
        'Implemented',
    ],
    [
        '1.2 AI-powered quiz generator with instructor validation controls',
        'Google Gemini API integration, quiz generation from uploaded documents (PDF/DOCX/TXT), teacher review/edit in Quiz Builder before publish',
        'Implemented',
    ],
    [
        '1.3 Automated grading module with customizable rubrics',
        'Instant scoring on quiz submission, configurable passing thresholds per quiz, per-question score breakdown on results page',
        'Implemented',
    ],
    [
        '1.4 Mastery-lock system for automated student remediation',
        'Sequential quiz locking on failure, auto-generated remediation quizzes, 70% mastery threshold, lock indicators in Classwork tab',
        'Implemented',
    ],
    [
        '1.5 Student performance dashboard for real-time academic tracking',
        'Per-student grades tab with DepEd descriptors, teacher gradebook with color-coded scores, submission status tracking',
        'Implemented',
    ],
]

for r_idx, row_data in enumerate(rows_data):
    for c_idx, text in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Arial'

for row in table.rows:
    row.cells[0].width = Inches(2.2)
    row.cells[1].width = Inches(3.0)
    row.cells[2].width = Inches(1.3)

doc.add_paragraph()

# ======================================================================
# ERP INTEGRATION
# ======================================================================

add_heading2('ERP Module Integration Summary')

add_body(
    'The A.I.M.S. LMS does not operate in isolation \u2014 it is part of a three-module School ERP ecosystem:'
)

add_body_no_indent(
    '1. Schedule Module \u2192 LMS: Class schedules, teacher assignments, and section rosters are fetched from '
    'the ERP Schedule module. The LMS uses this data to automatically create courses and enroll teachers and '
    'students, ensuring accurate and up-to-date class composition without manual intervention.'
)

add_body_no_indent(
    '2. LMS \u2192 Class Records Module: Partial grades generated from asynchronous activities (online quizzes '
    'and assignments) within the LMS are automatically pushed to the ERP Class Records module. This ensures '
    'that the teacher\'s official class record reflects all student performance data \u2014 both from '
    'face-to-face and online activities \u2014 in a single consolidated record.'
)

add_body_no_indent(
    '3. Bidirectional Sync: The three modules maintain data consistency through automatic synchronization, '
    'reducing administrative burden and ensuring accuracy across the school\'s academic management system.'
)

# ── Save ──
output_path = r'C:\Users\parus\Desktop\AIMS DOCU\System-Update-Documentation.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
print('Done! 21 pages, 24 screenshot placeholders with capture instructions.')
