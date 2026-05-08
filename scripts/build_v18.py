#!/usr/bin/env python3
"""build_v18.py

Builds papers/AIMS/revisions/cha1&2_v18.docx from v17.

Scope (minimal, literature-alignment cleanup):
- Soften Objective 1.4 wording so the literature supports targeted remediation
  principles rather than directly claiming proof of AI-generated remedial quiz
  mechanics.
- Soften Objective 1.5 wording so mastery locks are presented as A.I.M.S.'
  implementation of mastery learning principles, not a directly evidenced
  mechanism from the cited studies.
- Align the synthesis paragraph with those safer claims.

Constraints:
- Preserve formatting through run-preserving replacements.
- Do not introduce the word "dashboard".
- Never overwrite existing revision files.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v17.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v18.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    m = _leading_ws_re.match(text)
    return m.group(0) if m else ""


def replace_paragraph_text_preserve_runs(paragraph, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    prefix = leading_ws(paragraph.text)
    final_text = prefix + new_text

    if not paragraph.runs:
        paragraph.add_run(final_text)
        return

    paragraph.runs[0].text = final_text
    for run in paragraph.runs[1:]:
        run.text = ""


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(
    doc: Document,
    needle: str,
    start: int = 0,
    end: int | None = None,
) -> int:
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    ch2_start = find_paragraph_index(doc, "Chapter 2")
    ref_start = find_paragraph_index(doc, "REFERENCES")

    p116_i = find_paragraph_index_contains(
        doc,
        "Dynamic remediation is a core element of mastery-oriented instruction",
        start=ch2_start,
        end=ref_start,
    )
    p116_new = (
        "Dynamic remediation is a core element of mastery-oriented instruction, which requires systematic corrective feedback mechanisms and structured opportunities for retesting when learners do not meet performance thresholds (Persky & Hughes, 2022). "
        "In the context of A.I.M.S., the proposed dynamic remedial quiz generator applies this principle by using identified learning gaps to support the drafting of targeted remedial assessment items for teacher-reviewed intervention and re-assessment. "
        "Local evidence further affirms that structured remediation interventions can yield measurable learning gains in Philippine public school contexts (Jumao-as et al., 2025)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p116_i], p116_new)

    p119_i = find_paragraph_index_contains(
        doc,
        "Mastery-based learning is a pedagogical model grounded in the principle",
        start=ch2_start,
        end=ref_start,
    )
    p119_new = (
        "Mastery-based learning is a pedagogical model grounded in the principle that all students can achieve academic proficiency when given sufficient time, appropriate instructional support, and the opportunity to demonstrate competency before progressing to subsequent content. "
        "Persky and Hughes (2022) conducted a comprehensive review of mastery learning principles and their practical application in varied educational contexts, finding that mastery-based instruction consistently outperformed traditional time-based delivery models in promoting long-term content retention, reducing achievement gaps, and improving student confidence. "
        "Their review emphasized that effective mastery learning frameworks require clearly defined performance thresholds, systematic corrective feedback mechanisms, and structured opportunities for retesting. These principles provide the pedagogical basis for the proposed mastery-based learning module in A.I.M.S., where progression controls are used as a system design mechanism to enforce competency-based sequencing."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p119_i], p119_new)

    p121_i = find_paragraph_index_contains(
        doc,
        "Empirical evidence from the Philippine basic education context further affirms the value of structured remediation",
        start=ch2_start,
        end=ref_start,
    )
    p121_new = (
        "Empirical evidence from the Philippine basic education context further affirms the value of structured remediation in supporting learner progression. "
        "Jumao-as et al. (2025) investigated the effectiveness of a reading remediation program implemented in a DepEd elementary school in Ozamiz City, Misamis Occidental, finding that Grade 4 pupils identified as having frustrated reading levels through the Philippine Informal Reading Inventory (PHIL-IRI) advanced to the instructional reading level following the systematic remediation program. "
        "The study demonstrated a statistically significant improvement in reading comprehension, validating that targeted, competency-specific intervention embedded within a structured instructional cycle produces measurable learning gains. These findings support the rationale for linking learner progression in A.I.M.S. to demonstrated mastery and structured remediation, even though the platform's specific lock and override controls remain part of the system's implementation design."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p121_i], p121_new)

    p124_i = find_paragraph_index_contains(
        doc,
        "The reviewed literature collectively demonstrates that each of the five core features of A.I.M.S.",
        start=ch2_start,
        end=ref_start,
    )
    p124_new = (
        "The reviewed literature collectively demonstrates that each of the five core features of A.I.M.S. is grounded in a robust and growing body of empirical evidence. "
        "Centralized learning portal with material distribution and task submission supports structured access to learning materials and course activities in public school environments (Dahal & Manandhar, 2024; Bustillo & Aguilos, 2022), while AI-Powered Quiz Generator with instructor validation and approval controls reduces the cognitive and administrative burden of assessment creation through validated neural and LLM-based generation methods (Bulathwela et al., 2023; Alamoudi et al., 2025). "
        "Automated assessment and grading engine supporting teacher validation promotes evaluation consistency and scalability through machine learning and rubric-guided assessment (Ramesh & Sanampudi, 2022; García-Varela et al., 2025). "
        "Dynamic remedial quiz generator for targeted student intervention is supported by literature on corrective feedback, targeted remediation, and structured intervention cycles aligned with identified learning gaps (Persky & Hughes, 2022; Jumao-as et al., 2025). "
        "Finally, Mastery-based learning module with configurable progression locks is pedagogically grounded in mastery learning literature that emphasizes criterion-referenced progression, corrective feedback, and remediation before advancement (Persky & Hughes, 2022; Jumao-as et al., 2025)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p124_i], p124_new)

    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "dashboard" in full_text.lower():
        raise RuntimeError("Forbidden word 'dashboard' found after updates")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()