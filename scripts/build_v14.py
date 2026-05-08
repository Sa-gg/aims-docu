#!/usr/bin/env python3
"""build_v14.py

Builds papers/AIMS/revisions/cha1&2_v14.docx from v13.

Requested edits (Chapter 2 only):
1) Chapter 2 feature headings: convert to Title Case.
2) Table 1 (Matrix) feature column headers: shorten to core nouns.

Additional correctness edits (Chapter 2 only, minimal):
- Soften/align claims for Balase & Paglinawan (2025), Colegado (2025), and Ramesh & Sanampudi (2022)
  so they do not over-claim beyond abstract-level support.

Constraints:
- Preserve formatting; do not introduce the word “dashboard”.
- Do not overwrite existing revision files.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v13.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v14.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    m = _leading_ws_re.match(text)
    return m.group(0) if m else ""


def replace_paragraph_text_preserve_runs(p, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    prefix = leading_ws(p.text)
    final_text = prefix + new_text

    if not p.runs:
        p.add_run(final_text)
        return

    p.runs[0].text = final_text
    for r in p.runs[1:]:
        r.text = ""


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(doc: Document, needle: str, start: int = 0, end: int | None = None) -> int:
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def ensure_all_runs_bold(p) -> None:
    for r in p.runs:
        if r.text.strip():
            r.font.bold = True


def set_cell_text_preserve_runs(cell, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    if not cell.paragraphs:
        cell.text = new_text
        return
    p = cell.paragraphs[0]
    replace_paragraph_text_preserve_runs(p, new_text)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    # Boundaries
    ch2_start = find_paragraph_index(doc, "Chapter 2")
    ref_start = find_paragraph_index(doc, "REFERENCES")

    # ── 1) Chapter 2 headings -> Title Case (explicit mapping to avoid mistakes) ──
    headings_map = {
        "Centralized learning portal with material distribution and task submission":
            "Centralized Learning Portal with Material Distribution and Task Submission",
        "AI-Powered Quiz Generator with instructor validation and approval controls":
            "AI-Powered Quiz Generator with Instructor Validation and Approval Controls",
        "Automated assessment and grading engine supporting teacher validation":
            "Automated Assessment and Grading Engine Supporting Teacher Validation",
        "Dynamic remedial quiz generator for targeted student intervention":
            "Dynamic Remedial Quiz Generator for Targeted Student Intervention",
        "Mastery-based learning module with configurable progression locks":
            "Mastery-Based Learning Module with Configurable Progression Locks",
    }

    for i in range(ch2_start, ref_start):
        t = doc.paragraphs[i].text.strip()
        if t in headings_map:
            replace_paragraph_text_preserve_runs(doc.paragraphs[i], headings_map[t])
            ensure_all_runs_bold(doc.paragraphs[i])

    # ── 2) Table 1 matrix column headers -> shortened nouns ──────────────────
    if not doc.tables:
        raise RuntimeError("No tables found; expected Matrix table")

    tbl = doc.tables[0]
    header_row = tbl.rows[1]  # row 2 in Word

    short_headers = [
        "Centralized Learning Portal",
        "AI-Powered Quiz Generator",
        "Automated Grading Engine",
        "Dynamic Remedial Quizzes",
        "Mastery-Based Progression",
    ]

    for col_i, text in enumerate(short_headers, start=1):
        set_cell_text_preserve_runs(header_row.cells[col_i], text)

    # ── 3) Minimal claim-alignment edits (paragraph text only) ───────────────
    # Balase & Paglinawan (2025) paragraph about assessment scoring/class size.
    p108_i = find_paragraph_index_contains(
        doc,
        "In the Philippine educational landscape, the administrative challenges of assessment preparation",
        start=ch2_start,
        end=ref_start,
    )
    p108_new = (
        "In the Philippine educational landscape, the integration of learning management systems in public schools is often constrained by contextual and operational barriers. "
        "Balase and Paglinawan (2025) reported that teachers recognized the DepEd Learning Management System as beneficial for digital instruction and engagement, but also emphasized persistent barriers such as unstable connectivity, limited technical support, digital literacy gaps, and workload management concerns. "
        "These constraints underscore the importance of designing A.I.M.S. features to streamline teacher workflows and reduce additional administrative overhead, while maintaining instructor validation controls to preserve pedagogical integrity."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p108_i], p108_new)

    # Ramesh & Sanampudi (2022) paragraph with over-specific AES details.
    p111_i = find_paragraph_index_contains(
        doc,
        "Automated grading systems have gained considerable traction",
        start=ch2_start,
        end=ref_start,
    )
    p111_new = (
        "Automated grading systems have gained considerable traction in educational technology research as a response to the time demands and consistency limitations of manual evaluation. "
        "Ramesh and Sanampudi (2022) presented a systematic literature review of automated essay scoring (AES) systems, synthesizing research on machine learning and natural language processing approaches for grading written responses and outlining key challenges for reliable deployment in educational contexts. "
        "This body of work supports the inclusion of an automated assessment and grading engine in A.I.M.S., particularly when paired with teacher validation mechanisms and rubric-guided scoring to align automated outputs with instructor expectations."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p111_i], p111_new)

    # Colegado (2025) paragraph currently over-claims manual evaluation + urban confinement.
    p113_i = find_paragraph_index_contains(
        doc,
        "In Philippine educational settings, where teachers routinely manage large classes",
        start=ch2_start,
        end=ref_start,
    )
    p113_new = (
        "In Philippine educational settings, persistent infrastructural and capacity constraints continue to shape how digital tools are adopted and sustained. "
        "Colegado (2025), in a scoping review of digital innovations across Philippine K�12 science education, reported that teachers relied heavily on accessible platforms and that learning management systems such as Google Classroom and Quipper were used to provide structure for content delivery and assessment. "
        "However, the review emphasized ongoing challenges including poor connectivity, unequal device access, and limited teacher preparation, particularly in rural and underserved contexts. "
        "These findings highlight the need for A.I.M.S. to support structured assessment workflows and teacher-defined scoring criteria within realistic implementation constraints in Philippine public schools."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p113_i], p113_new)

    # Balase & Paglinawan (2025) paragraph currently mentions class size/adaptive tools/decision-support.
    p120_i = find_paragraph_index_contains(
        doc,
        "In the Philippine educational context, class size constraints",
        start=ch2_start,
        end=ref_start,
    )
    p120_new = (
        "In the Philippine educational context, the sustained implementation of digital learning platforms is influenced by practical constraints that limit teachers� capacity to manage additional system processes. "
        "Balase and Paglinawan (2025) highlighted barriers to effective LMS integration in public schools, including unstable connectivity, limited technical support, digital literacy gaps, and workload management concerns. "
        "Accordingly, the mastery-based learning module in A.I.M.S. should be designed to support competency-based progression in a way that minimizes added teacher workload while maintaining clear performance thresholds and instructor oversight."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p120_i], p120_new)

    # Synthesis paragraph: remove over-claim about fragmented standalone tools from Balase/Colegado.
    p125_i = find_paragraph_index_contains(
        doc,
        "A consistent theme across both foreign and local literature",
        start=ch2_start,
        end=ref_start,
    )
    p125_new = (
        "A consistent theme across both foreign and local literature is that the effectiveness of educational technology solutions depends not only on technical capability but also on contextual fit with teacher workflows and implementation constraints. "
        "Local evidence highlights persistent barriers such as unstable connectivity, limited technical support, digital literacy gaps, and workload management concerns that shape platform adoption in Philippine public schools (Balase & Paglinawan, 2025; Colegado, 2025). "
        "Across the reviewed studies, core functionalities such as centralized content distribution, AI-assisted assessment creation, automated grading, targeted remediation, and mastery-based progression are commonly investigated as distinct solutions. "
        "This indicates a continuing research and design opportunity to integrate the complete instructional and assessment cycle within a single, teacher-centered platform, as envisioned in A.I.M.S."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p125_i], p125_new)

    # Safety check: forbidden term must not appear
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "dashboard" in full_text.lower():
        raise RuntimeError("Forbidden word 'dashboard' found after updates")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()
