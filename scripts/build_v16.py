#!/usr/bin/env python3
"""build_v16.py

Builds papers/AIMS/revisions/cha1&2_v16.docx from v15.

Scope (minimal, content-correctness + compliance):
- Replace Bradley (2021) with a 2021–2026 source whose DOI metadata is compliant.
  - Bradley DOI 10.46328/ijte.36 is dated 2020-12-20 in Crossref.
  - Replacement: Dahal & Manandhar (2024) DOI 10.25082/amler.2024.01.001.
- Update the Objective 1.1 narrative to match the replacement abstract (avoid over-claims).
- Add a local Philippine citation directly inside Objective 1.4 subsection
  (Jumao-as et al., 2025) to satisfy local/foreign balance per subsection.
- Update Table 1 row label and REFERENCES entry accordingly.

Constraints:
- Preserve formatting (run-preserving replacements).
- Do not introduce the word “dashboard”.
- Never overwrite existing revision files.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v15.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v16.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    m = _leading_ws_re.match(text)
    return m.group(0) if m else ""


def replace_paragraph_text_preserve_runs(p, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    prefix = leading_ws(p.text)
    final_text = prefix + new_text

    if not p.runs:
        p.add_run(final_text)
        return

    p.runs[0].text = final_text
    for r in p.runs[1:]:
        r.text = ""


def set_cell_text_preserve_runs(cell, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    if not cell.paragraphs:
        cell.text = new_text
        return
    replace_paragraph_text_preserve_runs(cell.paragraphs[0], new_text)


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(
    doc: Document,
    needle: str,
    start: int = 0,
    end: int | None = None,
) -> int:
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    # Boundaries (used to restrict some edits to Chapter 2)
    ch2_start = find_paragraph_index(doc, "Chapter 2")
    ref_start = find_paragraph_index(doc, "REFERENCES")

    # Replacement source
    new_paren = "(Dahal & Manandhar, 2024)"
    new_narr = "Dahal and Manandhar (2024)"

    # ── Replace the earlier short definition sentence (occurs before Ch2) ───
    p_def_i = find_paragraph_index_contains(
        doc,
        "Centralized Learning Portal. Contextually, this is an integrated online platform",
        start=0,
        end=ch2_start,
    )
    p_def_new = (
        "Centralized Learning Portal. Contextually, this is an integrated online platform within a learning management system that "
        "supports structured content delivery and access to learning resources, as well as course activities and assessments "
        f"{new_paren}."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p_def_i], p_def_new)

    # ── Objective 1.1 subsection (Chapter 2) ────────────────────────────────
    p101_i = find_paragraph_index_contains(
        doc,
        "The centralized learning portal has emerged as a foundational component",
        start=ch2_start,
        end=ref_start,
    )
    p101_new = (
        "The centralized learning portal is a foundational component of modern learning management systems (LMS), providing a structured environment for organizing and delivering course resources. "
        "Dahal and Manandhar (2024) described an LMS as a digital platform for developing, delivering, and managing courses, learning resources, learning activities, and assessments, and noted that LMS adoption can improve content deliverability, accessibility, and retrievability. "
        "They also emphasized that treating an LMS solely as a storage system can limit its capacity to support meaningful learning, which highlights the importance of designing portals that support both resource management and active learning workflows." 
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p101_i], p101_new)

    p103_i = find_paragraph_index_contains(
        doc,
        "Web-based module management systems address these challenges",
        start=ch2_start,
        end=ref_start,
    )
    p103_new = (
        "Web-based module management systems address access and distribution constraints by enabling instructors to upload, categorize, and update learning materials electronically, while students retrieve resources through a standardized platform. "
        "Dahal and Manandhar (2024) highlighted that resources and learning management and content management are key success factors for LMS implementation, while also identifying assessment and learning engagement as areas that can fail when platforms are used only as passive storage. "
        "Accordingly, the centralized learning portal in A.I.M.S. should be designed not only for material distribution but also to support task submission and teacher-managed assessment workflows in a manner that remains practical for real-world classroom implementation." 
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p103_i], p103_new)

    # ── Objective 1.4 subsection: add local PH citation inside the subsection ─
    p116_i = find_paragraph_index_contains(
        doc,
        "Dynamic remediation is a core element of mastery-oriented instruction",
        start=ch2_start,
        end=ref_start,
    )
    p116_new = (
        "Dynamic remediation is a core element of mastery-oriented instruction, which requires systematic corrective feedback mechanisms and structured opportunities for retesting when learners do not meet performance thresholds (Persky & Hughes, 2022). "
        "In the context of A.I.M.S., a dynamic remedial quiz generator operationalizes this principle by producing targeted remedial assessment items based on identified learning gaps to support timely intervention and mastery attainment. "
        "Local evidence further affirms that structured remediation interventions can yield measurable learning gains in Philippine public school contexts (Jumao-as et al., 2025)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p116_i], p116_new)

    # ── Synthesis paragraphs that cite Bradley ──────────────────────────────
    p124_i = find_paragraph_index_contains(
        doc,
        "The reviewed literature collectively demonstrates that each of the five core features of A.I.M.S.",
        start=ch2_start,
        end=ref_start,
    )
    p124_old = doc.paragraphs[p124_i].text
    p124_new = p124_old.replace("(Bradley, 2021;", f"({new_paren[1:-1]};")
    # Also fix any standalone Bradley cites in this paragraph
    p124_new = p124_new.replace("Bradley, 2021", "Dahal & Manandhar, 2024")
    # Keep claim strength aligned: remove explicit 'inequities' phrasing tied to the foreign LMS source
    p124_new = p124_new.replace(
        "addresses persistent access and distribution inequities in public school environments",
        "supports structured access to learning materials and course activities in public school environments",
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p124_i], p124_new)

    p131_i = find_paragraph_index_contains(
        doc,
        "The matrix of related literature and studies above presents a comparative overview",
        start=ch2_start,
        end=ref_start,
    )
    p131_old = doc.paragraphs[p131_i].text
    p131_new = p131_old.replace("(Bradley, 2021;", f"({new_paren[1:-1]};")
    p131_new = p131_new.replace("Bradley, 2021", "Dahal & Manandhar, 2024")
    replace_paragraph_text_preserve_runs(doc.paragraphs[p131_i], p131_new)

    # ── Update any remaining (Bradley, 2021) occurrences in Chapter 2 ───────
    for i in range(ch2_start, ref_start):
        t = doc.paragraphs[i].text
        if "Bradley" not in t:
            continue
        t = t.replace("(Bradley, 2021)", new_paren)
        t = t.replace("Bradley (2021)", new_narr)
        t = t.replace("Bradley, 2021", "Dahal & Manandhar, 2024")
        if t != doc.paragraphs[i].text:
            replace_paragraph_text_preserve_runs(doc.paragraphs[i], t)

    # ── Update Table 1 cell that mentions Bradley ──────────────────────────
    if not doc.tables:
        raise RuntimeError("No tables found; expected Matrix table")

    replaced_cells = 0
    for row in doc.tables[0].rows:
        for cell in row.cells:
            cell_text = " ".join(p.text for p in cell.paragraphs)
            if "Bradley" not in cell_text:
                continue
            # Replace the entire cell content to avoid partial title mismatch
            set_cell_text_preserve_runs(
                cell,
                "The reality of e-Learning: Success and failure of learning management system – Dahal & Manandhar (2024)",
            )
            replaced_cells += 1

    # ── Replace REFERENCES entry ───────────────────────────────────────────
    ref_i = find_paragraph_index_contains(doc, "Bradley, V. M.", start=ref_start, end=len(doc.paragraphs))
    ref_new = (
        "Dahal, N., & Manandhar, N. (2024). The reality of e-learning: Success and failure of learning management system. "
        "Advances in Mobile Learning Educational Research, 4(1), 903–910. https://doi.org/10.25082/amler.2024.01.001"
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[ref_i], ref_new)

    # ── Safety checks ──────────────────────────────────────────────────────
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "dashboard" in full_text.lower():
        raise RuntimeError("Forbidden word 'dashboard' found after updates")
    if "Bradley" in full_text:
        raise RuntimeError("'Bradley' still present after replacement")

    doc.save(DST)
    print(f"Saved: {DST} (table cells updated: {replaced_cells})")


if __name__ == "__main__":
    main()
