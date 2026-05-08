#!/usr/bin/env python3
"""build_v11.py

Builds papers/AIMS/revisions/cha1&2_v11.docx from v10.
Edits are intentionally narrow:
  1) Chapter 1 -> Background of the Study: rewrite for clarity + add explicit research gap.
  2) Chapter 1 -> Objective of the Study: rewrite into proper "To ..." format and align ISO/IEC 25010 phrasing.

This script does not change formatting rules; it only updates paragraph text.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v10.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v11.docx")


def sanitize_xml_text(text: str) -> str:
    # Remove control characters not allowed in XML 1.0
    # Keep: tab (0x09), LF (0x0A), CR (0x0D)
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    m = _leading_ws_re.match(text)
    return m.group(0) if m else ""


def set_paragraph_text_preserve_leading_ws(p, new_text: str) -> None:
    prefix = leading_ws(p.text)
    p.text = sanitize_xml_text(prefix + new_text.strip())


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    # ── Background of the Study ─────────────────────────────────────────────
    bg_heading_i = find_paragraph_index(doc, "Background of the Study")
    obj_heading_i = find_paragraph_index(doc, "Objective of the Study")

    # Expect 4 body paragraphs under Background before the blank line.
    bg_targets = [bg_heading_i + 1, bg_heading_i + 2, bg_heading_i + 3, bg_heading_i + 4]

    bg_new = [
        (
            "Technology has become increasingly integrated in education systems to support instruction, assessment, "
            "and learner support. In particular, educational technologies powered by artificial intelligence (AI) "
            "have demonstrated potential to enable personalized learning, adaptive feedback, and more efficient "
            "assessment workflows that can help educators respond to diverse learner needs and improve academic "
            "performance (Merino-Campos, 2025; Marcos, 2026)."
        ),
        (
            "In the Philippine basic education context, the adoption of flexible learning modalities increased the "
            "use of technology in teaching and learning. However, persistent challenges such as the digital divide, "
            "uneven access to learning resources, inadequate infrastructure, and time-intensive manual processes for "
            "module distribution and assessment continue to affect many public schools (Bustillo & Aguilos, 2022; "
            "Balase & Paglinawan, 2025). Local evidence further indicates that limited connectivity and varying levels "
            "of digital competencies can prevent teachers and learners from maximizing the functionality of learning "
            "management systems and other digital tools (Colegado, 2025)."
        ),
        (
            "These conditions often result in fragmented workflows where learning materials, quiz creation, scoring, "
            "and follow-up interventions are handled through separate tools or manual procedures. Consequently, "
            "feedback and remediation may be delayed, and teachers may have limited capacity to provide targeted "
            "support for learners who have not yet mastered required competencies. A key research and implementation "
            "gap is the limited availability of a locally contextualized, integrated web-based platform for public "
            "secondary schools that unifies centralized module distribution, instructor-validated AI-assisted quiz "
            "generation, automated assessment with teacher validation, and dynamic remediation linked to mastery-based "
            "progression controls within a single system."
        ),
        (
            "In response to this gap, the study presents A.I.M.S. (Automated Intervention and Mastery System), an "
            "AI-enabled web-based educational platform designed to support public school teachers by centralizing "
            "module distribution, enabling instructor-validated AI-powered quiz generation, automating assessment and "
            "grading with teacher validation, and generating dynamic remedial quizzes linked to mastery-based "
            "progression locks. Through these integrated features, A.I.M.S. aims to reduce administrative workload, "
            "improve the timeliness of feedback and interventions, and strengthen instructional decision-making in "
            "public secondary school settings."
        ),
    ]

    if not (bg_targets[-1] < obj_heading_i):
        raise RuntimeError("Unexpected paragraph layout between Background and Objective sections")

    for target_i, new_text in zip(bg_targets, bg_new, strict=True):
        set_paragraph_text_preserve_leading_ws(doc.paragraphs[target_i], new_text)

    # ── Objectives of the Study ─────────────────────────────────────────────
    sig_heading_i = find_paragraph_index(doc, "Significance of the Study")
    obj_start = obj_heading_i + 1
    obj_end = sig_heading_i - 1

    # Rewrite only if we have the expected paragraph block size (v10 layout).
    # v10 objective block is paragraphs 30..47 (1-based), which is 18 paragraphs.
    expected_count = 18
    actual_count = obj_end - obj_start + 1
    if actual_count < expected_count:
        raise RuntimeError(
            f"Objective block shorter than expected: {actual_count} < {expected_count}"
        )

    obj_lines = [
        (
            "The general objective of the study is to design and develop a web-based application titled "
            "\"A.I.M.S. (Automated Intervention and Mastery System): A Web-Based Educational Platform with AI-Driven "
            "Assessment and Dynamic Remediation.\""
        ),
        "Specifically, the study aims to:",
        "1. To design and develop a web-based application with the following core technical features:",
        "1.1. Centralized learning portal with material distribution and task submission;",
        "1.2. AI-Powered Quiz Generator with instructor validation and approval controls;",
        "1.3. Automated assessment and grading engine supporting teacher validation;",
        "1.4. Dynamic remedial quiz generator for targeted student intervention; and",
        "1.5. Mastery-based learning module with configurable progression locks.",
        "2. To implement the developed system and conduct functional testing to verify the correct operation of the features and workflows.",
        "3. To evaluate the system based on ISO/IEC 25010 standards in terms of:",
        "3.1. Functional suitability",
        "3.2. Performance efficiency",
        "3.3. Usability",
        "3.4. Compatibility",
        "3.5. Reliability",
        "3.6. Security",
        "3.7. Maintainability",
        "3.8. Portability",
    ]

    for offset, line in enumerate(obj_lines):
        set_paragraph_text_preserve_leading_ws(doc.paragraphs[obj_start + offset], line)

    # If there are extra paragraphs in the objective block (should not happen in v10), blank them.
    for i in range(obj_start + len(obj_lines), obj_end + 1):
        doc.paragraphs[i].text = ""

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()
