#!/usr/bin/env python3
"""build_v12.py

Builds papers/AIMS/revisions/cha1&2_v12.docx from v11.

Goal: Align Chapter 1 module descriptions with the intended A.I.M.S. lifecycle
(teacher validation/approval flows, task submission logging, configurable mastery lock
thresholds + override), without inserting the lifecycle steps into the paper.

Edits are intentionally narrow and text-only:
  - Chapter 1 -> Significance of the Study (Teachers/Students sentences)
  - Chapter 1 -> Scope and Limitation of the Study (one paragraph)
  - Chapter 1 -> Definition of Terms (operational definitions for the 5 core modules)
  - Terminology consistency: ISO 25010 -> ISO/IEC 25010 (term label)

This script never overwrites existing revisions.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v11.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v12.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    m = _leading_ws_re.match(text)
    return m.group(0) if m else ""


def set_paragraph_text_preserve_leading_ws(p, new_text: str) -> None:
    prefix = leading_ws(p.text)
    p.text = sanitize_xml_text(prefix + new_text.strip())


def find_paragraph_exact(doc: Document, needle: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_contains(doc: Document, needle: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if needle in p.text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def update_labeled_paragraph(doc: Document, heading: str, label: str, new_body: str) -> None:
    """Update a paragraph that starts with a bold label run like 'Teachers.'"""
    start = find_paragraph_exact(doc, heading)
    end = len(doc.paragraphs)
    for i in range(start + 1, end):
        t = doc.paragraphs[i].text.strip()
        if t == "":
            continue
        # Stop once we reach the next major heading
        if t in {"Scope and Limitation of the Study", "Definition of Terms", "Chapter 2"}:
            break
        if t.startswith(label):
            p = doc.paragraphs[i]
            if not p.runs:
                set_paragraph_text_preserve_leading_ws(p, label + " " + new_body)
                return
            # Preserve label run formatting (typically bold)
            p.runs[0].text = label
            if len(p.runs) == 1:
                p.add_run(" " + new_body)
            else:
                p.runs[1].text = " " + new_body
                for r in p.runs[2:]:
                    r.text = ""
            return
    raise ValueError(f"Could not locate labeled paragraph {label!r} under {heading!r}")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    # ── Significance (alignment: submissions, validation, remediation, mastery locks) ──
    update_labeled_paragraph(
        doc,
        heading="Significance of the Study",
        label="Teachers.",
        new_body=(
            "The system helps teachers centralize learning modules, collect task submissions with recorded timestamps, "
            "and streamline assessment through instructor-validated AI-assisted quiz drafting and automated scoring. "
            "It also supports timely intervention by surfacing mastery status and pending remedial activities for teacher review and approval, "
            "thereby reducing administrative workload and allowing teachers to focus more on instruction delivery and learner support."
        ),
    )

    update_labeled_paragraph(
        doc,
        heading="Significance of the Study",
        label="Students.",
        new_body=(
            "A.I.M.S. enables students to access learning modules, submit tasks, and take assessments through a unified online portal. "
            "It provides immediate feedback for objective items, and when mastery is not achieved, it supports guided remediation through teacher-approved targeted remedial quizzes. "
            "This encourages self-directed learning and helps learners identify areas for improvement and mastery of content."
        ),
    )

    # ── Scope (add explicit approval/threshold/override alignment without adding lifecycle steps) ──
    scope_i = find_paragraph_exact(doc, "Scope and Limitation of the Study")
    scope_body_i = scope_i + 2  # v11: heading, then setting paragraph, then scope paragraph
    scope_text = doc.paragraphs[scope_body_i].text
    if "core technical scope" not in scope_text:
        # fallback: find by phrase
        scope_body_i = find_paragraph_contains(doc, "Its core technical scope")

    set_paragraph_text_preserve_leading_ws(
        doc.paragraphs[scope_body_i],
        (
            "Its core technical scope includes a centralized learning portal for material distribution and task submission, "
            "an AI-powered quiz generator with strict instructor validation controls, and an automated assessment and grading engine supporting teacher validation. "
            "To address learning gaps, it features a dynamic remedial quiz generator for targeted student intervention, queued for instructor approval prior to deployment, "
            "and driven by a mastery-based learning module with teacher-configurable progression locks, passing thresholds, and instructor override controls. "
            "The system is specifically designed to serve public secondary schools and may require structural modifications for tertiary or private institutions. "
            "Furthermore, while the quiz generation utilizes generative Large Language Models (LLMs), instructor validation remains a mandatory system requirement to ensure all materials meet contextual DepEd curriculum standards."
        ),
    )

    # ── Definitions (Operational) ──
    # Centralized Learning Portal (operational)
    p = doc.paragraphs[66]  # 1-based 67
    if not p.text.startswith("Operationally, it is the repository feature"):
        # robust locate
        idx = find_paragraph_contains(doc, "Operationally, it is the repository feature in A.I.M.S.")
        p = doc.paragraphs[idx]
    set_paragraph_text_preserve_leading_ws(
        p,
        (
            "Operationally, it is the centralized learning portal feature in A.I.M.S. that allows teachers to upload and organize learning materials and tasks by subject and grade level, "
            "enables students to securely access and download modules, and supports online task submission with automatic logging of submission timestamps for teacher review."
        ),
    )

    # AI-Powered Quiz Generator (operational)
    idx = find_paragraph_contains(doc, "uses Google Gemini API")
    set_paragraph_text_preserve_leading_ws(
        doc.paragraphs[idx],
        (
            "Operationally, in this study, it refers to the module in A.I.M.S. that uses the Google Gemini API to draft objective quiz items (e.g., multiple-choice) "
            "from teacher-selected learning materials or prompts, including distractors and answer keys, and requires teacher review, editing, and approval before a quiz is deployed to students."
        ),
    )

    # Automated Assessment and Grading Engine (operational)
    idx = find_paragraph_contains(doc, "evaluates quiz responses according")
    set_paragraph_text_preserve_leading_ws(
        doc.paragraphs[idx],
        (
            "Operationally, it is the system component in A.I.M.S. that administers quizzes, records responses securely, and automatically computes preliminary scores for objective items. "
            "For responses requiring human judgment, it supports teacher validation using customizable rubrics before final scores are logged and made visible to students as feedback."
        ),
    )

    # Mastery-Based Learning Module (operational)
    idx = find_paragraph_contains(doc, "conditional progression feature")
    set_paragraph_text_preserve_leading_ws(
        doc.paragraphs[idx],
        (
            "Operationally, it is the conditional progression feature in A.I.M.S. that restricts access to subsequent lessons until students meet a teacher-defined passing threshold "
            "on prerequisite assessments. When mastery is not achieved, progression remains locked until the learner completes teacher-approved remedial activities, unless the teacher manually overrides the lock."
        ),
    )

    # Dynamic Remedial Quiz Generator (operational)
    idx = find_paragraph_contains(doc, "generates remedial quizzes")
    set_paragraph_text_preserve_leading_ws(
        doc.paragraphs[idx],
        (
            "Operationally, it is the A.I.M.S. component that generates targeted remedial quizzes based on a learner's assessment results and DepEd learning competencies using generative Large Language Models (LLMs). "
            "Remedial quizzes are saved as pending interventions and require instructor validation and approval before they are deployed to the student for re-assessment."
        ),
    )

    # ISO term label consistency (keep bold label run)
    idx = find_paragraph_exact(doc, "Definition of Terms")
    iso_i = find_paragraph_contains(doc, "ISO 25010.")
    iso_p = doc.paragraphs[iso_i]
    if iso_p.runs:
        iso_p.runs[0].text = "ISO/IEC 25010."
        if len(iso_p.runs) > 1:
            iso_p.runs[1].text = iso_p.runs[1].text.replace("ISO 25010", "ISO/IEC 25010")
    else:
        set_paragraph_text_preserve_leading_ws(
            iso_p,
            (
                "ISO/IEC 25010. Contextually, it is the international software quality standard that defines eight characteristics for evaluating system and software products: "
                "functional suitability, performance efficiency, compatibility, usability, reliability, security, maintainability, and portability (ISO/IEC, 2011)."
            ),
        )

    # Also update the operational ISO paragraph if present
    try:
        iso_op_i = find_paragraph_contains(doc, "Operationally, it is the standardized evaluation framework")
        # keep as-is but fix A.I.M.S.' system -> A.I.M.S. system
        t = doc.paragraphs[iso_op_i].text
        if "A.I.M.S.'" in t:
            doc.paragraphs[iso_op_i].text = t.replace("A.I.M.S.' system", "A.I.M.S. system")
    except ValueError:
        pass

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()
