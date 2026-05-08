#!/usr/bin/env python3
"""build_v19.py

Builds papers/AIMS/revisions/cha1&2_v19.docx from v18.

Scope (minimal, plan execution):
- Add Simon et al. (2025) to strengthen Objective 1.1 portal / submission support.
- Add Gorgun & Bulut (2024) to strengthen Objective 1.2 teacher-validation
  and question-quality framing.
- Add Hegde et al. (2024) to strengthen Objective 1.4 targeted remediation.
- Add Toti et al. (2023) to strengthen Objective 1.5 mastery progression.
- Update the Chapter 2 synthesis paragraph accordingly.
- Add matching REFERENCES entries.

Explicit exclusion:
- Do NOT add Quah et al. (2024) because A.I.M.S. auto-scores objective items,
  while essay / subjective tasks require teacher manual grading.

Constraints:
- Preserve formatting through run-preserving replacements.
- Do not introduce the word "dashboard".
- Never overwrite existing revision files.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v18.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v19.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    match = _leading_ws_re.match(text)
    return match.group(0) if match else ""


def replace_paragraph_text_preserve_runs(paragraph, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    prefix = leading_ws(paragraph.text)
    final_text = prefix + new_text

    if not paragraph.runs:
        paragraph.add_run(final_text)
        return

    paragraph.runs[0].text = final_text
    for run in paragraph.runs[1:]:
        run.text = ""


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(
    doc: Document,
    needle: str,
    start: int = 0,
    end: int | None = None,
) -> int:
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    ch2_start = find_paragraph_index(doc, "Chapter 2")
    ref_start = find_paragraph_index(doc, "REFERENCES")

    p103_i = find_paragraph_index_contains(
        doc,
        "Web-based module management systems address access and distribution constraints",
        start=ch2_start,
        end=ref_start,
    )
    p103_new = (
        "Web-based module management systems address access and distribution constraints by enabling instructors to upload, categorize, and update learning materials electronically, while students retrieve resources through a standardized platform. "
        "Simon et al. (2025) further observed that learning management systems function as active environments for information dissemination, assignments, feedback, and classroom management, which supports the role of A.I.M.S. as both a content access and submission platform. "
        "Dahal and Manandhar (2024) highlighted that resources and learning management and content management are key success factors for LMS implementation, while also identifying assessment and learning engagement as areas that can fail when platforms are used only as passive storage. "
        "Accordingly, the centralized learning portal in A.I.M.S. should be designed not only for material distribution but also to support task submission and teacher-managed assessment workflows in a manner that remains practical for real-world classroom implementation."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p103_i], p103_new)

    p107_i = find_paragraph_index_contains(
        doc,
        "Recent research has further validated the use of large language models",
        start=ch2_start,
        end=ref_start,
    )
    p107_new = (
        "Recent research has further validated the use of large language models (LLMs) in automated question generation for educational domains. "
        "Gorgun and Bulut (2024) noted that although automated question generation can efficiently produce large volumes of assessment items, gaps in standardized quality evaluation and usability review still hinder seamless educational deployment. "
        "Alamoudi et al. (2025) compared three methodological approaches-template-based structured ontology generation, LLM-based structured ontology generation, and LLM-based flat concept list generation-using BERT Precision, Recall, F1-score, and semantic similarity as performance metrics. "
        "Their findings indicate that the LLM-based flat concept list approach achieved the highest semantic similarity (0.567) and the most balanced precision-recall performance. Significantly, the authors noted that unstructured LLM generation occasionally produced hallucinated or topically unrelated questions, reinforcing the necessity of instructor validation controls in any AI-powered quiz generation system to ensure accuracy, curriculum alignment, and pedagogical relevance before deployment."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p107_i], p107_new)

    p116_i = find_paragraph_index_contains(
        doc,
        "Dynamic remediation is a core element of mastery-oriented instruction",
        start=ch2_start,
        end=ref_start,
    )
    p116_new = (
        "Dynamic remediation is a core element of mastery-oriented instruction, which requires systematic corrective feedback mechanisms and structured opportunities for retesting when learners do not meet performance thresholds (Persky & Hughes, 2022). "
        "Hegde et al. (2024) provided additional empirical support for this approach by showing that underperforming students improved after tailored remediation that incorporated quizzes, assignments, and focused review strategies following the identification of learning deficiencies. "
        "In the context of A.I.M.S., the proposed dynamic remedial quiz generator applies this principle by using identified learning gaps to support the drafting of targeted remedial assessment items for teacher-reviewed intervention and re-assessment. "
        "Local evidence further affirms that structured remediation interventions can yield measurable learning gains in Philippine public school contexts (Jumao-as et al., 2025)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p116_i], p116_new)

    p119_i = find_paragraph_index_contains(
        doc,
        "Mastery-based learning is a pedagogical model grounded in the principle",
        start=ch2_start,
        end=ref_start,
    )
    p119_new = (
        "Mastery-based learning is a pedagogical model grounded in the principle that all students can achieve academic proficiency when given sufficient time, appropriate instructional support, and the opportunity to demonstrate competency before progressing to subsequent content. "
        "Persky and Hughes (2022) conducted a comprehensive review of mastery learning principles and their practical application in varied educational contexts, finding that mastery-based instruction consistently outperformed traditional time-based delivery models in promoting long-term content retention, reducing achievement gaps, and improving student confidence. "
        "Their review emphasized that effective mastery learning frameworks require clearly defined performance thresholds, systematic corrective feedback mechanisms, and structured opportunities for retesting. "
        "Toti et al. (2023) further demonstrated that students in a mastery learning framework can prove mastery of individual units over multiple attempts, showing how progression can be tied to demonstrated competence rather than one-time completion. "
        "These principles provide the pedagogical basis for the proposed mastery-based learning module in A.I.M.S., where progression controls are used as a system design mechanism to enforce competency-based sequencing."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p119_i], p119_new)

    p124_i = find_paragraph_index_contains(
        doc,
        "The reviewed literature collectively demonstrates that each of the five core features of A.I.M.S.",
        start=ch2_start,
        end=ref_start,
    )
    p124_new = (
        "The reviewed literature collectively demonstrates that each of the five core features of A.I.M.S. is grounded in a robust and growing body of empirical evidence. "
        "Centralized learning portal with material distribution and task submission supports structured access to learning materials and course activities in public school environments (Dahal & Manandhar, 2024; Bustillo & Aguilos, 2022; Simon et al., 2025), while AI-Powered Quiz Generator with instructor validation and approval controls reduces the cognitive and administrative burden of assessment creation through validated neural and LLM-based generation methods together with quality review requirements for educational deployment (Bulathwela et al., 2023; Alamoudi et al., 2025; Gorgun & Bulut, 2024). "
        "Automated assessment and grading engine supporting teacher validation promotes evaluation consistency and scalability through machine learning and rubric-guided assessment for objective scoring workflows while retaining teacher review for subjective responses (Ramesh & Sanampudi, 2022; García-Varela et al., 2025). "
        "Dynamic remedial quiz generator for targeted student intervention is supported by literature on corrective feedback, targeted remediation, and structured intervention cycles aligned with identified learning gaps (Persky & Hughes, 2022; Hegde et al., 2024; Jumao-as et al., 2025). "
        "Finally, Mastery-based learning module with configurable progression locks is pedagogically grounded in mastery learning literature that emphasizes criterion-referenced progression, repeated attempts, corrective feedback, and remediation before advancement (Persky & Hughes, 2022; Toti et al., 2023; Jumao-as et al., 2025)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p124_i], p124_new)

    ref_style = doc.paragraphs[ref_start + 1].style if ref_start + 1 < len(doc.paragraphs) else None
    new_refs = [
        "Gorgun, G., & Bulut, O. (2024). Exploring quality criteria and evaluation methods in automated question generation: A comprehensive survey. Education and Information Technologies, 29(18), 24111-24142. https://doi.org/10.1007/s10639-024-12771-3",
        "Hegde, S. V., Shetty, P. P., Senthilkumar, M., Kandimalla, R., Bernhardt, G. V., Pinto, J. R. T., Mahadevan, R., Kotian, S. M., & Rashmi, K. S. (2024). Bridging knowledge gaps: Impact of remedial classes on first-year medical students in biochemistry - a cross-sectional study. BMC Medical Education, 24(1), Article 1375. https://doi.org/10.1186/s12909-024-06243-y",
        "Simon, P. D., Jiang, J., Fryer, L. K., King, R. B., & Frondozo, C. E. (2025). An assessment of learning management system use in higher education: Perspectives from a comprehensive sample of teachers and students. Technology, Knowledge and Learning, 30(2), 741-767. https://doi.org/10.1007/s10758-024-09734-5",
        "Toti, G., Chen, G., & Gonzalez, S. (2023). Teaching CS1 with a mastery learning framework: Impact on students' learning and engagement. In Proceedings of the 2023 Conference on Innovation and Technology in Computer Science Education V. 1 (pp. 540-546). https://doi.org/10.1145/3587102.3588844",
    ]
    full_text_before_refs = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    for ref in new_refs:
        if ref.split(". ", 1)[0] in full_text_before_refs:
            raise RuntimeError(f"Reference appears to already exist: {ref}")
        doc.add_paragraph(sanitize_xml_text(ref), style=ref_style)

    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "dashboard" in full_text.lower():
        raise RuntimeError("Forbidden word 'dashboard' found after updates")
    if "Quah" in full_text:
        raise RuntimeError("Excluded Quah source was introduced unexpectedly")

    doc.save(DST)
    print(f"Saved: {DST} (added references: {len(new_refs)})")


if __name__ == "__main__":
    main()