#!/usr/bin/env python3
"""build_v20.py

Builds papers/AIMS/revisions/cha1&2_v20.docx from v19.

Scope:
- Realign Objective 1.3 with actual A.I.M.S. behavior.
- Remove essay-scoring literature that implies automated essay grading.
- Replace it with sources that better support automated assessment for
  structured tasks and teacher oversight.
- Update the Chapter 2 synthesis and matrix entries for Objective 1.3.
- Replace the matching REFERENCES entries and append one new reference.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v19.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v20.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    match = _leading_ws_re.match(text)
    return match.group(0) if match else ""


def replace_paragraph_text_preserve_runs(paragraph, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    prefix = leading_ws(paragraph.text)
    final_text = prefix + new_text

    if not paragraph.runs:
        paragraph.add_run(final_text)
        return

    paragraph.runs[0].text = final_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_cell_text_preserve_runs(cell, new_text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    replace_paragraph_text_preserve_runs(cell.paragraphs[0], new_text)
    for paragraph in cell.paragraphs[1:]:
        for run in paragraph.runs:
            run.text = ""


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(
    doc: Document,
    needle: str,
    start: int = 0,
    end: int | None = None,
) -> int:
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    ch2_start = find_paragraph_index(doc, "Chapter 2")
    ref_start = find_paragraph_index(doc, "REFERENCES")

    p111_i = find_paragraph_index_contains(
        doc,
        "Automated grading systems have gained considerable traction",
        start=ch2_start,
        end=ref_start,
    )
    p111_new = (
        "Automated assessment systems are increasingly used to reduce manual checking time and to provide immediate feedback in structured learning tasks. "
        "Zampirolli et al. (2021) reported that automated assessment in a large-scale introductory programming course supported continuous assessment through immediate and automatic feedback, and the classes that adopted the approach achieved higher pass rates than those that did not. "
        "Although their context focused on programming education, the study demonstrates how automated assessment can efficiently evaluate structured student outputs at scale while still supporting instructional decision-making. "
        "This supports the inclusion of an automated assessment and grading engine in A.I.M.S. for objective items generated through the platform's quiz workflow, where multiple-choice responses can be scored automatically and presented to teachers for review."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p111_i], p111_new)

    p112_i = find_paragraph_index_contains(
        doc,
        "Beyond holistic scoring, the incorporation of customizable rubrics",
        start=ch2_start,
        end=ref_start,
    )
    p112_new = (
        "At the same time, the use of AI in educational assessment must be framed carefully in terms of validity, scope, and oversight. "
        "Gardner et al. (2021) argued that AI-enabled educational assessment should be examined critically in relation to efficacy and validity rather than treated as an unconditional replacement for human judgment. "
        "Likewise, Ouyang et al. (2023), in a systematic review of AI-driven educational assessment in STEM education, showed that automated assessment has become a sustained area of research across formative and summative contexts. "
        "Taken together, these studies support a teacher-validation model for A.I.M.S., in which automation is applied to objective-response workflows while teachers retain responsibility for reviewing results and manually evaluating subjective or essay-based responses."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p112_i], p112_new)

    p113_i = find_paragraph_index_contains(
        doc,
        "In Philippine educational settings, persistent infrastructural and capacity constraints",
        start=ch2_start,
        end=ref_start,
    )
    p113_new = (
        "In Philippine educational settings, persistent infrastructural and capacity constraints continue to shape how digital tools are adopted and sustained. "
        "Colegado (2025), in a scoping review of digital innovations across Philippine K-12 science education, reported that teachers relied heavily on accessible platforms and that learning management systems such as Google Classroom and Quipper were used to provide structure for content delivery and assessment. "
        "However, the review emphasized ongoing challenges including poor connectivity, unequal device access, and limited teacher preparation, particularly in rural and underserved contexts. "
        "These findings highlight the need for A.I.M.S. to support structured assessment workflows in which objective quiz items can be scored automatically while teachers retain control over scoring criteria and manual evaluation for subjective responses within realistic implementation constraints in Philippine public schools."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p113_i], p113_new)

    p124_i = find_paragraph_index_contains(
        doc,
        "The reviewed literature collectively demonstrates that each of the five core features of A.I.M.S.",
        start=ch2_start,
        end=ref_start,
    )
    p124_new = (
        "The reviewed literature collectively demonstrates that each of the five core features of A.I.M.S. is grounded in a robust and growing body of empirical evidence. "
        "Centralized learning portal with material distribution and task submission supports structured access to learning materials and course activities in public school environments (Dahal & Manandhar, 2024; Bustillo & Aguilos, 2022; Simon et al., 2025), while AI-Powered Quiz Generator with instructor validation and approval controls reduces the cognitive and administrative burden of assessment creation through validated neural and LLM-based generation methods together with quality review requirements for educational deployment (Bulathwela et al., 2023; Alamoudi et al., 2025; Gorgun & Bulut, 2024). "
        "Automated assessment and grading engine supporting teacher validation is supported by literature showing that automated assessment can provide immediate feedback and scalable evaluation in structured tasks, while broader reviews of AI-enabled assessment emphasize validity and human oversight in educational use (Zampirolli et al., 2021; Gardner et al., 2021; Ouyang et al., 2023). "
        "Dynamic remedial quiz generator for targeted student intervention is supported by literature on corrective feedback, targeted remediation, and structured intervention cycles aligned with identified learning gaps (Persky & Hughes, 2022; Hegde et al., 2024; Jumao-as et al., 2025). "
        "Finally, Mastery-based learning module with configurable progression locks is pedagogically grounded in mastery learning literature that emphasizes criterion-referenced progression, repeated attempts, corrective feedback, and remediation before advancement (Persky & Hughes, 2022; Toti et al., 2023; Jumao-as et al., 2025)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p124_i], p124_new)

    p131_i = find_paragraph_index_contains(
        doc,
        "The matrix of related literature and studies above presents a comparative overview",
        start=ch2_start,
        end=ref_start,
    )
    p131_new = (
        "The matrix of related literature and studies above presents a comparative overview of existing systems and studies in relation to the five core technical features of the proposed A.I.M.S. platform. "
        "The studies mapped in the matrix span both foreign and local literature, covering Centralized learning portal with material distribution and task submission (Dahal & Manandhar, 2024; Bustillo & Aguilos, 2022), AI-Powered Quiz Generator with instructor validation and approval controls (Bulathwela et al., 2023; Alamoudi et al., 2025), Automated assessment and grading engine supporting teacher validation (Zampirolli et al., 2021; Ouyang et al., 2023; Colegado, 2025), Dynamic remedial quiz generator for targeted student intervention (Persky & Hughes, 2022; Jumao-as et al., 2025), and Mastery-based learning module with configurable progression locks (Persky & Hughes, 2022; Jumao-as et al., 2025)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p131_i], p131_new)

    table = doc.tables[0]
    replace_cell_text_preserve_runs(
        table.rows[5].cells[0],
        "A Systematic Review of AI-Driven Educational Assessment in STEM Education - Ouyang et al. (2023)",
    )
    replace_cell_text_preserve_runs(
        table.rows[7].cells[0],
        "An Experience of Automated Assessment in a Large-Scale Introduction Programming Course - Zampirolli et al. (2021)",
    )
    replace_cell_text_preserve_runs(table.rows[2].cells[3], "✗")
    replace_cell_text_preserve_runs(table.rows[5].cells[3], "✓")
    replace_cell_text_preserve_runs(table.rows[6].cells[3], "✗")
    replace_cell_text_preserve_runs(table.rows[7].cells[3], "✓")
    replace_cell_text_preserve_runs(table.rows[10].cells[1], "✗")
    replace_cell_text_preserve_runs(table.rows[10].cells[3], "✓")

    ramesh_i = find_paragraph_index_contains(doc, "Ramesh, D., & Sanampudi, S. K.", start=ref_start)
    ouyang_ref = (
        "Ouyang, F., Dinh, T. A., & Xu, W. (2023). A systematic review of AI-driven educational assessment in STEM education. "
        "Journal for STEM Education Research, 6(3), 383-426. https://doi.org/10.1007/s41979-023-00112-x"
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[ramesh_i], ouyang_ref)

    garcia_i = find_paragraph_index_contains(doc, "Garc", start=ref_start)
    zampirolli_ref = (
        "Zampirolli, F. A., Borovina Josko, J. M., Venero, M. L. F., Kobayashi, G., Fraga, F. J., Goya, D., & Savegnago, H. R. (2021). "
        "An experience of automated assessment in a large-scale introduction programming course. Computer Applications in Engineering Education, 29(5), 1284-1299. https://doi.org/10.1002/cae.22385"
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[garcia_i], zampirolli_ref)

    ref_style = doc.paragraphs[ref_start + 1].style if ref_start + 1 < len(doc.paragraphs) else None
    gardner_ref = (
        "Gardner, J., O'Leary, M., & Yuan, L. (2021). Artificial intelligence in educational assessment: 'Breakthrough? Or buncombe and ballyhoo?' "
        "Journal of Computer Assisted Learning, 37(5), 1207-1216. https://doi.org/10.1111/jcal.12577"
    )
    full_text_before_append = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "Gardner, J., O'Leary, M., & Yuan, L. (2021)." not in full_text_before_append:
        doc.add_paragraph(sanitize_xml_text(gardner_ref), style=ref_style)

    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "automated essay scoring" in full_text.lower():
        raise RuntimeError("Essay-scoring wording remains after Objective 1.3 realignment")
    if "García-Varela" in full_text or "Ramesh, D., & Sanampudi" in full_text:
        raise RuntimeError("Old Objective 1.3 essay-scoring references remain")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()