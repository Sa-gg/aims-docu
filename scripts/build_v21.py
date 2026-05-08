#!/usr/bin/env python3
"""build_v21.py

Builds papers/AIMS/revisions/cha1&2_v21.docx from v20.

Scope:
- Fix the stray old DOI artifact left in the Ouyang reference paragraph.
- Preserve all Objective 1.3 content changes introduced in v20.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v20.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v21.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(doc: Document, needle: str, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)
    ref_start = find_paragraph_index(doc, "REFERENCES")
    ouyang_i = find_paragraph_index_contains(doc, "Ouyang, F., Dinh, T. A., & Xu, W. (2023).", start=ref_start)

    clean_ref = (
        "Ouyang, F., Dinh, T. A., & Xu, W. (2023). A systematic review of AI-driven educational assessment in STEM education. "
        "Journal for STEM Education Research, 6(3), 383-426. https://doi.org/10.1007/s41979-023-00112-x"
    )

    # Force-reset the full paragraph text to remove leftover hyperlink-field text.
    doc.paragraphs[ouyang_i].text = sanitize_xml_text(clean_ref)

    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "10.1007/s10462-021-10068-2" in full_text:
        raise RuntimeError("Old Ramesh DOI still present after reference cleanup")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()