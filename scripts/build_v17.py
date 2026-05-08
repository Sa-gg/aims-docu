#!/usr/bin/env python3
"""build_v17.py

Builds papers/AIMS/revisions/cha1&2_v17.docx from v16.

Scope (corrective, minimal):
- Fix REFERENCES entry for Dahal & Manandhar (2024) to remove the stray appended DOI
  "https://doi.org/10.46328/ijte.36" that remained due to hyperlink field remnants.

Constraints:
- Preserve existing content; do not introduce the word “dashboard”.
- Never overwrite existing revision files.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v16.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v17.docx")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    # Locate REFERENCES heading
    ref_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "REFERENCES")

    target_old = "https://doi.org/10.46328/ijte.36"
    target_key = "Dahal, N., & Manandhar, N. (2024)."

    fixed = 0
    for i in range(ref_start + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if target_key not in p.text:
            continue

        clean_ref = (
            "Dahal, N., & Manandhar, N. (2024). The reality of e-learning: Success and failure of learning management system. "
            "Advances in Mobile Learning Educational Research, 4(1), 903–910. https://doi.org/10.25082/amler.2024.01.001"
        )

        # Force reset paragraph content to avoid leftover hyperlink field runs.
        # This keeps the paragraph style but replaces its run structure.
        p.text = clean_ref

        # If the old DOI still appears after reset, something is very wrong.
        if target_old in p.text:
            raise RuntimeError("Old Bradley DOI still present after forced reference reset")

        fixed += 1
        break

    if fixed != 1:
        raise RuntimeError("Expected to fix exactly 1 Dahal & Manandhar reference paragraph")

    # Safety checks
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "dashboard" in full_text.lower():
        raise RuntimeError("Forbidden word 'dashboard' found")
    if target_old in full_text:
        raise RuntimeError("Old Bradley DOI still present in document")

    doc.save(DST)
    print(f"Saved: {DST} (fixed references: {fixed})")


if __name__ == "__main__":
    main()
