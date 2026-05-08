#!/usr/bin/env python3
"""build_v23.py

Builds papers/AIMS/revisions/cha1&2_v23.docx from v22.

Scope:
- Remove hyperlink-field artifacts left in reordered reference paragraphs.
- Preserve the corrected matrix table and matrix synthesis from v22.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v22.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v23.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


CLEAN_REFERENCES = [
    "Alamoudi, S., Al Khuzayem, L. A., & Jamal, A. (2025). Optimizing automated question generation for educational assessments: A semantic analysis of LLMs with structured and unstructured ontologies. Engineering, Technology & Applied Science Research, 15(3), 23664–23671. https://doi.org/10.48084/etasr.10662",
    "Balase, M. M. H., & Paglinawan, J. L. (2025). Voices from the field: Teachers' narratives and reflections on integrating the DepEd Learning Management System (LMS) in teaching and professional development. International Journal of Scientific and Management Research, 8(10), 147–159. https://doi.org/10.37502/ijsmr.2025.81011",
    "Bulathwela, S., Pérez-Ortiz, M., Holloway, C., Cukurova, M., & Shawe-Taylor, J. (2023). Review on neural question generation for education purposes. International Journal of Artificial Intelligence in Education, 33(4), 691–717. https://doi.org/10.1007/s40593-023-00374-x",
    "Bustillo, E., & Aguilos, M. (2022). The challenges of modular learning in the wake of COVID-19: A digital divide in the Philippine countryside revealed. Education Sciences, 12(7), 449. https://doi.org/10.3390/educsci12070449",
    "Colegado, J. C. (2025). Digital innovations in science education in the Philippines: A scoping review of teaching practices and tools. International Journal of Research and Innovation in Social Science, 9(3S), 6549–6556. https://doi.org/10.47772/IJRISS.2025.903SEDU0479",
    "Dahal, N., & Manandhar, N. (2024). The reality of e-learning: Success and failure of learning management system. Advances in Mobile Learning Educational Research, 4(1), 903–910. https://doi.org/10.25082/amler.2024.01.001",
    "Gardner, J., O'Leary, M., & Yuan, L. (2021). Artificial intelligence in educational assessment: 'Breakthrough? Or buncombe and ballyhoo?' Journal of Computer Assisted Learning, 37(5), 1207–1216. https://doi.org/10.1111/jcal.12577",
    "Gorgun, G., & Bulut, O. (2024). Exploring quality criteria and evaluation methods in automated question generation: A comprehensive survey. Education and Information Technologies, 29(18), 24111–24142. https://doi.org/10.1007/s10639-024-12771-3",
    "Hegde, S. V., Shetty, P. P., Senthilkumar, M., Kandimalla, R., Bernhardt, G. V., Pinto, J. R. T., Mahadevan, R., Kotian, S. M., & Rashmi, K. S. (2024). Bridging knowledge gaps: Impact of remedial classes on first-year medical students in biochemistry - a cross-sectional study. BMC Medical Education, 24(1), Article 1375. https://doi.org/10.1186/s12909-024-06243-y",
    "International Organization for Standardization. (2011). Systems and software engineering—Systems and software Quality Requirements and Evaluation (SQuaRE)—System and software quality models (ISO/IEC Standard No. 25010:2011). https://www.iso.org/standard/35733.html",
    "Jumao-as, J., Fuentes, J. L., Lacar, L., Fuertes, G. J., Erag, C., & Tagaylo, C. (2025). Reading remediation program and Grade 4 pupils' reading comprehension. Asian Journal of Education and Social Studies, 51(6), 365–377. https://doi.org/10.9734/ajess/2025/v51i62000",
    "Marcos, L. T. (2026). A systematic review on artificial intelligence in education: Opportunities, challenges, and ethical implications. Preprints. https://doi.org/10.20944/preprints202601.0448.v1",
    "Merino-Campos, C. (2025). The impact of artificial intelligence on personalized learning in higher education: A systematic review. Trends in Higher Education, 4(2), Article 17. https://doi.org/10.3390/higheredu4020017",
    "Ouyang, F., Dinh, T. A., & Xu, W. (2023). A systematic review of AI-driven educational assessment in STEM education. Journal for STEM Education Research, 6(3), 383–426. https://doi.org/10.1007/s41979-023-00112-x",
    "Persky, A. M., & Hughes, M. L. (2022). A practical review of mastery learning. American Journal of Pharmaceutical Education, 86(10), Article 8906. https://doi.org/10.5688/ajpe8906",
    "Simon, P. D., Jiang, J., Fryer, L. K., King, R. B., & Frondozo, C. E. (2025). An assessment of learning management system use in higher education: Perspectives from a comprehensive sample of teachers and students. Technology, Knowledge and Learning, 30(2), 741–767. https://doi.org/10.1007/s10758-024-09734-5",
    "Toti, G., Chen, G., & Gonzalez, S. (2023). Teaching CS1 with a mastery learning framework: Impact on students' learning and engagement. In Proceedings of the 2023 Conference on Innovation and Technology in Computer Science Education V. 1 (pp. 540–546). https://doi.org/10.1145/3587102.3588844",
    "Zampirolli, F. A., Borovina Josko, J. M., Venero, M. L. F., Kobayashi, G., Fraga, F. J., Goya, D., & Savegnago, H. R. (2021). An experience of automated assessment in a large-scale introduction programming course. Computer Applications in Engineering Education, 29(5), 1284–1299. https://doi.org/10.1002/cae.22385",
]


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)
    ref_start = find_paragraph_index(doc, "REFERENCES")

    ref_paragraphs = [p for p in doc.paragraphs[ref_start + 1:] if p.text.strip()]
    if len(ref_paragraphs) != len(CLEAN_REFERENCES):
        raise RuntimeError("Unexpected reference paragraph count in source document")

    for paragraph, text in zip(ref_paragraphs, CLEAN_REFERENCES):
        paragraph.text = sanitize_xml_text(text)

    refs_blob = "\n".join(p.text for p in ref_paragraphs)
    doi_count = refs_blob.count("https://doi.org/")
    if doi_count != len(ref_paragraphs) - 1:
        # One reference is an ISO standard URL rather than a DOI URL.
        raise RuntimeError("Unexpected DOI URL count after reference cleanup")
    if "https://doi.org/10.48084/etasr.10662https://doi.org/10.48084/etasr.10662" in refs_blob:
        raise RuntimeError("Duplicate DOI artifact remains in references")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()