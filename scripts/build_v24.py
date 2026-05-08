#!/usr/bin/env python3
"""build_v24.py

Builds papers/AIMS/revisions/cha1&2_v24.docx from v23.

Scope:
- Realign Chapter 1 wording with the updated Objective 1.3 behavior.
- Remove the stale Ramesh & Sanampudi (2022) contextual citation from Chapter 1.
- Clarify that automated scoring is limited to objective quiz items, while
  subjective responses remain teacher-reviewed.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v23.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v24.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    match = _leading_ws_re.match(text)
    return match.group(0) if match else ""


def replace_paragraph_text_preserve_runs(paragraph, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    prefix = leading_ws(paragraph.text)
    final_text = prefix + new_text

    if not paragraph.runs:
        paragraph.add_run(final_text)
        return

    paragraph.runs[0].text = final_text
    for run in paragraph.runs[1:]:
        run.text = ""


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(
    doc: Document,
    needle: str,
    start: int = 0,
    end: int | None = None,
) -> int:
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    chapter1 = find_paragraph_index(doc, "Chapter 1")
    chapter2 = find_paragraph_index(doc, "Chapter 2")

    p27_i = find_paragraph_index_contains(
        doc,
        "In response to this gap, the study presents A.I.M.S.",
        start=chapter1,
        end=chapter2,
    )
    p27_new = (
        "In response to this gap, the study presents A.I.M.S. (Automated Intervention and Mastery System), an AI-enabled web-based educational platform designed to support public school teachers by centralizing module distribution, enabling instructor-validated AI-powered quiz generation, automating the scoring of objective quiz items with teacher validation, and generating dynamic remedial quizzes linked to mastery-based progression locks. "
        "Through these integrated features, A.I.M.S. aims to reduce administrative workload, improve the timeliness of feedback and interventions, and strengthen instructional decision-making in public secondary school settings."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p27_i], p27_new)

    p51_i = find_paragraph_index_contains(
        doc,
        "Teachers. The system helps teachers centralize learning modules",
        start=chapter1,
        end=chapter2,
    )
    p51_new = (
        "Teachers. The system helps teachers centralize learning modules, collect task submissions with recorded timestamps, and streamline assessment through instructor-validated AI-assisted quiz drafting and automated scoring of objective quiz items. "
        "It also supports timely intervention by surfacing mastery status and pending remedial activities for teacher review and approval, thereby reducing administrative workload and allowing teachers to focus more on instruction delivery and learner support."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p51_i], p51_new)

    p58_i = find_paragraph_index_contains(
        doc,
        "Its core technical scope includes a centralized learning portal",
        start=chapter1,
        end=chapter2,
    )
    p58_new = (
        "Its core technical scope includes a centralized learning portal for material distribution and task submission, an AI-powered quiz generator with strict instructor validation controls, and an automated assessment and grading engine supporting teacher validation. "
        "Within this engine, automatic scoring is limited to objective quiz items, while responses requiring human judgment remain teacher-reviewed. "
        "To address learning gaps, it features a dynamic remedial quiz generator for targeted student intervention, queued for instructor approval prior to deployment, and driven by a mastery-based learning module with teacher-configurable progression locks, passing thresholds, and instructor override controls. "
        "The system is specifically designed to serve public secondary schools and may require structural modifications for tertiary or private institutions. Furthermore, while the quiz generation utilizes generative Large Language Models (LLMs), instructor validation remains a mandatory system requirement to ensure all materials meet contextual DepEd curriculum standards."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p58_i], p58_new)

    p64_i = find_paragraph_index_contains(
        doc,
        "Automated Assessment and Grading Engine. Contextually",
        start=chapter1,
        end=chapter2,
    )
    p64_new = (
        "Automated Assessment and Grading Engine. Contextually, it refers to AI-assisted and automated assessment approaches that support the evaluation of structured student outputs and improve the efficiency of educational assessment workflows, while emphasizing validity, scope, and human oversight in actual instructional use (Ouyang et al., 2023; Gardner et al., 2021)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p64_i], p64_new)

    full_ch1 = "\n".join(p.text for p in doc.paragraphs[chapter1:chapter2])
    if "Ramesh & Sanampudi" in full_ch1 or "Ramesh and Sanampudi" in full_ch1:
        raise RuntimeError("Stale Ramesh citation remains in Chapter 1")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()