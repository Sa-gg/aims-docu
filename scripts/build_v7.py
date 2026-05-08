#!/usr/bin/env python3
"""
build_v7.py
Builds cha1&2_v7.docx from v6 with ATLAS-matching Chapter 2 spacing:
  - Section headings: LEFT-aligned (ATLAS Heading 1 style behaviour)
  - Proper empty separator paragraphs between every RRL section
  - Clean empty paragraphs (no stray fi/sa values)
Output: papers/AIMS/revisions/cha1&2_v7.docx
"""
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SRC = "papers/AIMS/revisions/cha1&2_v6.docx"
DST = "papers/AIMS/revisions/cha1&2_v7.docx"

doc = Document(SRC)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def _is_section_heading(p):
    """True if paragraph is a Chapter-2 RRL section heading
    (bold run, no first_line_indent, non-empty, not a chapter-level title,
    not a table label, not the synthesis subheading after the table)."""
    txt = p.text.strip()
    if not txt:
        return False
    pf = p.paragraph_format
    runs_bold = any(r.font.bold for r in p.runs if r.text.strip())
    if not runs_bold:
        return False
    if pf.first_line_indent and pf.first_line_indent > 0:
        return False  # body paragraph
    # Exclude chapter-level headings and table labels
    excluded = {
        'Chapter 2',
        'REVIEW OF RELATED LITERATURE AND STUDIES',
        'Table 1',
        'Matrix of Related Literature and Studies',
    }
    if txt in excluded:
        return False
    return True


def _clean_pPr(p):
    """Remove all w:pPr children except w:pStyle, then reapply desired props."""
    pPr = p._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._p.insert(0, pPr)
    # Remove spacing, ind, jc, textAlignment that we want to control
    for tag in ('w:spacing', 'w:ind', 'w:jc'):
        el = pPr.find(qn(tag))
        if el is not None:
            pPr.remove(el)
    return pPr


def set_heading_format(p):
    """Apply ATLAS-matching section heading format: left-aligned, bold TNR 12pt, double, sa=0."""
    pf = p.paragraph_format
    pf.alignment        = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after       = Pt(0)
    pf.space_before      = Pt(0)


def set_empty_para_format(p):
    """Clean up stray formatting on empty separator paragraphs."""
    pf = p.paragraph_format
    pf.alignment         = None
    pf.first_line_indent = None
    pf.space_after       = Pt(0)
    pf.space_before      = None
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.line_spacing      = 2.0


def make_empty_para(doc_obj):
    """Create a clean empty paragraph with double spacing, sa=0."""
    p_new = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    # Double line spacing
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'),     '480')
    sp.set(qn('w:lineRule'), 'auto')
    sp.set(qn('w:after'),    '0')
    pPr.append(sp)
    p_new.append(pPr)
    return p_new


# ─── Find Chapter 2 range ────────────────────────────────────────────────────

ch2_para_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Chapter 2':
        ch2_para_idx = i
        break

print(f"Chapter 2 starts at paragraph index {ch2_para_idx}")

# ─── Fix existing paragraph formatting ───────────────────────────────────────

section_heading_indices = []

for i in range(ch2_para_idx, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    txt = p.text.strip()

    if _is_section_heading(p):
        set_heading_format(p)
        section_heading_indices.append(i)
        print(f"  Fixed heading [{i}]: {txt[:60]!r}")
    elif not txt:
        set_empty_para_format(p)

print(f"\nSection headings found: {section_heading_indices}")

# ─── Insert missing empty separator paragraphs ───────────────────────────────
# ATLAS pattern: every RRL section ends with an empty paragraph before the next heading
# We need an empty para BEFORE each section heading (except the very first one)
# Check by looking at the body children order

body = doc.element.body

def get_preceding_body_child(p_element):
    """Return the body child immediately before this paragraph element."""
    children = list(body)
    idx = children.index(p_element)
    if idx == 0:
        return None
    return children[idx - 1]

# Work backwards through section headings so indices stay valid
for para_i in reversed(section_heading_indices[1:]):  # skip the FIRST section heading
    p = doc.paragraphs[para_i]
    prev_child = get_preceding_body_child(p._p)
    if prev_child is None:
        continue
    # Check if preceding body child is already an empty paragraph
    tag = prev_child.tag.split('}')[-1]
    if tag == 'p':
        prev_text = ''.join(
            t.text or '' for t in prev_child.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
            )
        ).strip()
        if prev_text == '':
            print(f"  Empty para already exists before heading [{para_i}] — skipping insert")
            continue

    # Insert empty paragraph before this heading
    new_empty = make_empty_para(doc)
    p._p.addprevious(new_empty)
    print(f"  Inserted empty para before heading [{para_i}]: {p.text.strip()[:50]!r}")

# ─── Save ────────────────────────────────────────────────────────────────────

doc.save(DST)
print(f"\nSaved → {DST}")

# ─── Verify ──────────────────────────────────────────────────────────────────

doc2 = Document(DST)
total_h = 0
print("\n=== v7 Chapter 2 structure ===")

in_ch2 = False
for i, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if txt == 'Chapter 2':
        in_ch2 = True
    if not in_ch2:
        continue

    pf = p.paragraph_format
    align = pf.alignment
    fi = pf.first_line_indent
    sa = pf.space_after
    run_bold = any(r.font.bold for r in p.runs if r.text.strip())
    tag = '📌 HEADING' if _is_section_heading(p) else ('⬜ empty' if not txt else '  body')
    if txt == 'Chapter 2' or txt == 'REVIEW OF RELATED LITERATURE AND STUDIES':
        tag = '🔷 CH TITLE'
    print(f"  [{i:3d}] {tag:12s} align={str(align):<25s} fi={str(fi):<8s} sa={str(sa):<8s} bold={run_bold!s:<5} | {txt[:55]!r}")

print(f"\nTotal tables: {len(doc2.tables)}")
print(f"Total paragraphs: {len(doc2.paragraphs)}")
print("Done.")
