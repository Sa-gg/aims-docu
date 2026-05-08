#!/usr/bin/env python3
"""build_v22.py

Builds papers/AIMS/revisions/cha1&2_v22.docx from v21.

Scope:
- Clean APA formatting details in recently added references.
- Alphabetize the REFERENCES section.
- Correct stale matrix feature mappings.
- Add newer key sources to the matrix table.
- Update the matrix synthesis paragraph to match the table.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v21.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v22.docx")


def sanitize_xml_text(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if (ord(ch) >= 0x20) or (ch in ("\t", "\n", "\r"))
    )


_leading_ws_re = re.compile(r"^\s+")


def leading_ws(text: str) -> str:
    match = _leading_ws_re.match(text)
    return match.group(0) if match else ""


def replace_paragraph_text_preserve_runs(paragraph, new_text: str) -> None:
    new_text = sanitize_xml_text(new_text.strip())
    prefix = leading_ws(paragraph.text)
    final_text = prefix + new_text

    if not paragraph.runs:
        paragraph.add_run(final_text)
        return

    paragraph.runs[0].text = final_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_cell_text_preserve_runs(cell, new_text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    replace_paragraph_text_preserve_runs(cell.paragraphs[0], new_text)
    for paragraph in cell.paragraphs[1:]:
        for run in paragraph.runs:
            run.text = ""


def find_paragraph_index(doc: Document, needle: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == needle:
            return i
    raise ValueError(f"Paragraph not found: {needle!r}")


def find_paragraph_index_contains(
    doc: Document,
    needle: str,
    start: int = 0,
    end: int | None = None,
) -> int:
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if needle in doc.paragraphs[i].text:
            return i
    raise ValueError(f"Paragraph containing {needle!r} not found")


def normalize_sort_key(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", " ", text).strip().lower()


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    ch2_start = find_paragraph_index(doc, "Chapter 2")
    ref_start = find_paragraph_index(doc, "REFERENCES")

    p131_i = find_paragraph_index_contains(
        doc,
        "The matrix of related literature and studies above presents a comparative overview",
        start=ch2_start,
        end=ref_start,
    )
    p131_new = (
        "The matrix of related literature and studies above presents a comparative overview of existing systems and studies in relation to the five core technical features of the proposed A.I.M.S. platform. "
        "The studies mapped in the matrix span both foreign and local literature, covering Centralized learning portal with material distribution and task submission (Dahal & Manandhar, 2024; Bustillo & Aguilos, 2022; Simon et al., 2025), AI-Powered Quiz Generator with instructor validation and approval controls (Alamoudi et al., 2025; Bulathwela et al., 2023; Gorgun & Bulut, 2024; Balase & Paglinawan, 2025), Automated assessment and grading engine supporting teacher validation (Zampirolli et al., 2021; Ouyang et al., 2023; Gardner et al., 2021; Colegado, 2025), Dynamic remedial quiz generator for targeted student intervention (Persky & Hughes, 2022; Hegde et al., 2024; Jumao-as et al., 2025), and Mastery-based learning module with configurable progression locks (Persky & Hughes, 2022; Balase & Paglinawan, 2025; Jumao-as et al., 2025; Toti et al., 2023)."
    )
    replace_paragraph_text_preserve_runs(doc.paragraphs[p131_i], p131_new)

    table = doc.tables[0]

    # Fix stale mappings left from earlier revisions.
    replace_cell_text_preserve_runs(table.rows[4].cells[1], "✗")
    replace_cell_text_preserve_runs(table.rows[4].cells[2], "✓")
    replace_cell_text_preserve_runs(table.rows[6].cells[2], "✗")
    replace_cell_text_preserve_runs(table.rows[6].cells[4], "✓")
    replace_cell_text_preserve_runs(table.rows[6].cells[5], "✓")
    replace_cell_text_preserve_runs(table.rows[9].cells[1], "✗")
    replace_cell_text_preserve_runs(table.rows[9].cells[2], "✓")

    first_column_texts = {
        " | ".join(par.text.strip() for par in row.cells[0].paragraphs if par.text.strip()): row
        for row in table.rows
    }
    new_matrix_rows = [
        ("An assessment of learning management system use in higher education – Simon et al. (2025)", ["✓", "✗", "✗", "✗", "✗"]),
        ("Exploring quality criteria and evaluation methods in automated question generation – Gorgun & Bulut (2024)", ["✗", "✓", "✗", "✗", "✗"]),
        ("Artificial intelligence in educational assessment: 'Breakthrough? Or buncombe and ballyhoo?' – Gardner et al. (2021)", ["✗", "✗", "✓", "✗", "✗"]),
        ("Bridging knowledge gaps: Impact of remedial classes on first-year medical students in biochemistry – Hegde et al. (2024)", ["✗", "✗", "✗", "✓", "✗"]),
        ("Teaching CS1 with a mastery learning framework – Toti et al. (2023)", ["✗", "✗", "✗", "✗", "✓"]),
    ]
    for label, features in new_matrix_rows:
        if label in first_column_texts:
            continue
        row = table.add_row()
        replace_cell_text_preserve_runs(row.cells[0], label)
        for cell_index, value in enumerate(features, start=1):
            replace_cell_text_preserve_runs(row.cells[cell_index], value)

    clean_refs = {
        "Zampirolli, F. A.,": (
            "Zampirolli, F. A., Borovina Josko, J. M., Venero, M. L. F., Kobayashi, G., Fraga, F. J., Goya, D., & Savegnago, H. R. (2021). "
            "An experience of automated assessment in a large-scale introduction programming course. Computer Applications in Engineering Education, 29(5), 1284–1299. https://doi.org/10.1002/cae.22385"
        ),
        "Ouyang, F.,": (
            "Ouyang, F., Dinh, T. A., & Xu, W. (2023). A systematic review of AI-driven educational assessment in STEM education. "
            "Journal for STEM Education Research, 6(3), 383–426. https://doi.org/10.1007/s41979-023-00112-x"
        ),
        "Gorgun, G.,": (
            "Gorgun, G., & Bulut, O. (2024). Exploring quality criteria and evaluation methods in automated question generation: A comprehensive survey. "
            "Education and Information Technologies, 29(18), 24111–24142. https://doi.org/10.1007/s10639-024-12771-3"
        ),
        "Simon, P. D.,": (
            "Simon, P. D., Jiang, J., Fryer, L. K., King, R. B., & Frondozo, C. E. (2025). An assessment of learning management system use in higher education: Perspectives from a comprehensive sample of teachers and students. "
            "Technology, Knowledge and Learning, 30(2), 741–767. https://doi.org/10.1007/s10758-024-09734-5"
        ),
        "Toti, G.,": (
            "Toti, G., Chen, G., & Gonzalez, S. (2023). Teaching CS1 with a mastery learning framework: Impact on students' learning and engagement. "
            "In Proceedings of the 2023 Conference on Innovation and Technology in Computer Science Education V. 1 (pp. 540–546). https://doi.org/10.1145/3587102.3588844"
        ),
        "Gardner, J.,": (
            "Gardner, J., O'Leary, M., & Yuan, L. (2021). Artificial intelligence in educational assessment: 'Breakthrough? Or buncombe and ballyhoo?' "
            "Journal of Computer Assisted Learning, 37(5), 1207–1216. https://doi.org/10.1111/jcal.12577"
        ),
    }

    ref_paragraphs = [p for p in doc.paragraphs[ref_start + 1:] if p.text.strip()]
    ref_texts = []
    for paragraph in ref_paragraphs:
        text = paragraph.text.strip()
        for prefix, clean_text in clean_refs.items():
            if text.startswith(prefix):
                text = clean_text
                break
        ref_texts.append(text)

    ref_texts = sorted(ref_texts, key=normalize_sort_key)
    if len(ref_texts) != len(ref_paragraphs):
        raise RuntimeError("Reference paragraph count changed unexpectedly")

    for paragraph, text in zip(ref_paragraphs, ref_texts):
        replace_paragraph_text_preserve_runs(paragraph, text)

    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    table_text = "\n".join(
        " | ".join(par.text.strip() for par in cell.paragraphs if par.text.strip())
        for row in table.rows
        for cell in row.cells
    )
    required_matrix_labels = [
        "Simon et al. (2025)",
        "Gorgun & Bulut (2024)",
        "Gardner et al. (2021)",
        "Hegde et al. (2024)",
        "Toti et al. (2023)",
    ]
    for label in required_matrix_labels:
        if label not in table_text:
            raise RuntimeError(f"Missing matrix label after update: {label}")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()