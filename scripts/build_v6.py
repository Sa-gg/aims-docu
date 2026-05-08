#!/usr/bin/env python3
"""
build_v6.py
Replaces the placeholder matrix table in cha1&2_v5.docx with an
ATLAS-format TableGrid: 2-row rotated header + 14 data rows.
Output: papers/AIMS/revisions/cha1&2_v6.docx
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "papers/AIMS/revisions/cha1&2_v5.docx"
DST = "papers/AIMS/revisions/cha1&2_v6.docx"

# Column widths (dxa) — total must equal 8637 (same as ATLAS)
# Col 0 = Studies/System column, Cols 1-5 = 5 feature columns
COL_W = [5048, 720, 717, 718, 718, 716]
assert sum(COL_W) == 8637, f"Column widths sum to {sum(COL_W)}, expected 8637"

# AIMS module names for the 5 rotated column headers
MODULES = [
    "Centralized Digital Library (1.1)",
    "AI-Powered Quiz Generator (1.2)",
    "Automated Grading Module (1.3)",
    "Mastery-Lock System (1.4)",
    "Student Performance Dashboard (1.5)",
]

# 14 study rows: (full title – Author(s) (Year), [1.1, 1.2, 1.3, 1.4, 1.5])
STUDIES = [
    # ---- Foreign Studies ----
    ("Optimizing Automated Question Generation for Educational Assessments – Alamoudi et al. (2025)",
     [0, 1, 1, 0, 0]),
    ("Learning Management System (LMS) Use with Online Instruction – Bradley (2021)",
     [1, 0, 0, 0, 0]),
    ("Review on Neural Question Generation for Education Purposes – Bulathwela et al. (2023)",
     [1, 1, 0, 0, 0]),
    ("Learning Analytics Dashboard Design and Evaluation to Support Student Self-Regulation – de Vreugd et al. (2024)",
     [0, 0, 0, 0, 1]),
    ("ChatGPT as a Stable and Fair Tool for Automated Essay Scoring – García-Varela et al. (2025)",
     [0, 0, 1, 0, 0]),
    ("Have Learning Analytics Dashboards Lived Up to the Hype? A Systematic Review – Kaliisa et al. (2024)",
     [0, 0, 0, 0, 1]),
    ("A Practical Review of Mastery Learning – Persky & Hughes (2022)",
     [0, 1, 1, 0, 0]),
    ("An Automated Essay Scoring Systems: A Systematic Literature Review – Ramesh & Sanampudi (2022)",
     [0, 0, 1, 0, 0]),
    ("Learning Analytics Dashboard: A Tool for Providing Actionable Insights to Learners – Susnjak et al. (2022)",
     [0, 0, 0, 1, 1]),
    # ---- Local (Philippine) Studies ----
    ("The Challenges of Modular Learning in the Wake of COVID-19: A Digital Divide in the Philippine Countryside Revealed – Bustillo & Aguilos (2022)",
     [1, 0, 0, 0, 0]),
    ("Voices from the Field: Teachers' Narratives and Reflections on Integrating the DepEd LMS – Balase & Paglinawan (2025)",
     [1, 0, 0, 1, 0]),
    ("Digital Innovations in Science Education in the Philippines: A Scoping Review – Colegado (2025)",
     [1, 0, 0, 0, 0]),
    ("Smart eLearning: A Framework Development of a Web Portal for Data-Driven Assessment – Gajardo & Balahadia (2025)",
     [0, 0, 0, 0, 1]),
    ("Reading Remediation Program and Grade 4 Pupils' Reading Comprehension – Jumao-as et al. (2025)",
     [0, 0, 0, 1, 0]),
]

# Height for header rows (dxa) — matching ATLAS
ROW0_H = 1500   # "RRL | Features" row — taller to show vertical text
ROW1_H = 2890   # Rotated-header row (individual module names)

# ─── XML helpers ────────────────────────────────────────────────────────────

def _borders():
    """Standard black single-line borders on all four sides (ATLAS style)."""
    b = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        s = OxmlElement(f'w:{side}')
        s.set(qn('w:val'),   'single')
        s.set(qn('w:sz'),    '4')
        s.set(qn('w:space'), '0')
        s.set(qn('w:color'), '000000')
        b.append(s)
    return b


def _tcW(w):
    e = OxmlElement('w:tcW')
    e.set(qn('w:w'),    str(w))
    e.set(qn('w:type'), 'dxa')
    return e


def _vAlign(val='center'):
    e = OxmlElement('w:vAlign')
    e.set(qn('w:val'), val)
    return e


def _spacing_zero():
    e = OxmlElement('w:spacing')
    e.set(qn('w:after'), '0')
    return e


def _run(text, font='Times New Roman', sz=24, bold=False, sym=False):
    """Build a w:r element."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rF = OxmlElement('w:rFonts')
    if sym:
        # Segoe UI Symbol for ✓/✗
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rF.set(qn(attr), 'Segoe UI Symbol')
    else:
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rF.set(qn(attr), font)
    rPr.append(rF)

    if bold:
        rPr.append(OxmlElement('w:b'))

    for tag in ('w:sz', 'w:szCs'):
        e = OxmlElement(tag)
        e.set(qn('w:val'), str(sz))
        rPr.append(e)

    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    if text != text.strip():
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r


def _para(runs, center=False):
    """Build a w:p with given run elements."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pPr.append(_spacing_zero())
    if center:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
    p.append(pPr)
    for r in runs:
        p.append(r)
    return p


def _trPr_height(h, rule='atLeast'):
    trPr = OxmlElement('w:trPr')
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),   str(h))
    trH.set(qn('w:hRule'), rule)
    trPr.append(trH)
    return trPr


def _tblPr():
    """Table-level properties — mirrors ATLAS TableGrid format."""
    tblPr = OxmlElement('w:tblPr')

    ts = OxmlElement('w:tblStyle')
    ts.set(qn('w:val'), 'TableGrid')
    tblPr.append(ts)

    tW = OxmlElement('w:tblW')
    tW.set(qn('w:w'),    '8637')
    tW.set(qn('w:type'), 'dxa')
    tblPr.append(tW)

    # Explicit table borders
    tb = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        s = OxmlElement(f'w:{side}')
        s.set(qn('w:val'),   'single')
        s.set(qn('w:sz'),    '4')
        s.set(qn('w:space'), '0')
        s.set(qn('w:color'), '000000')
        tb.append(s)
    tblPr.append(tb)

    # Cell margins matching ATLAS (top=62, left=108, bottom=0, right=0)
    cm = OxmlElement('w:tblCellMar')
    margins = [('top','62'), ('left','108'), ('bottom','0'), ('right','0')]
    for side, val in margins:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    val)
        m.set(qn('w:type'), 'dxa')
        cm.append(m)
    tblPr.append(cm)

    return tblPr


def _tblGrid():
    tg = OxmlElement('w:tblGrid')
    for w in COL_W:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tg.append(gc)
    return tg


# ─── Cell builders ──────────────────────────────────────────────────────────

def normal_cell(width, text, bold=False, center=False, font='Times New Roman', sz=24):
    """Plain text cell."""
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    tcPr.append(_tcW(width))
    tcPr.append(_borders())
    tcPr.append(_vAlign())
    tc.append(tcPr)
    tc.append(_para([_run(text, font=font, sz=sz, bold=bold)], center=center))
    return tc


def gridspan_cell(width, span, text, bold=False, center=True):
    """Cell spanning multiple columns (for 'Features' header)."""
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    tcPr.append(_tcW(width))
    gs = OxmlElement('w:gridSpan')
    gs.set(qn('w:val'), str(span))
    tcPr.append(gs)
    tcPr.append(_borders())
    tcPr.append(_vAlign())
    tc.append(tcPr)
    tc.append(_para([_run(text, bold=bold)], center=center))
    return tc


def rotated_cell(width, text):
    """Cell with bottom-to-top vertical text direction (ATLAS rotated header style)."""
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    tcPr.append(_tcW(width))
    tcPr.append(_borders())
    # textDirection btLr = bottom-left to top-right = 90° counterclockwise
    td = OxmlElement('w:textDirection')
    td.set(qn('w:val'), 'btLr')
    tcPr.append(td)
    tcPr.append(_vAlign())
    tc.append(tcPr)
    # Use sz=15 (7.5pt) matching ATLAS rotated header font size
    tc.append(_para([_run(text, sz=15, bold=True)]))
    return tc


def check_cell(width, mark):
    """Cell with ✓ or ✗ in Segoe UI Symbol font, centered."""
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    tcPr.append(_tcW(width))
    tcPr.append(_borders())
    tcPr.append(_vAlign())
    tc.append(tcPr)
    tc.append(_para([_run(mark, sz=24, sym=True)], center=True))
    return tc


# ─── Table builders ─────────────────────────────────────────────────────────

def rotated_span_cell(width, span, text):
    """Multi-column spanning cell with vertical (btLr) text — for 'Features' header."""
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    tcPr.append(_tcW(width))
    gs = OxmlElement('w:gridSpan')
    gs.set(qn('w:val'), str(span))
    tcPr.append(gs)
    tcPr.append(_borders())
    td = OxmlElement('w:textDirection')
    td.set(qn('w:val'), 'btLr')
    tcPr.append(td)
    tcPr.append(_vAlign())
    tc.append(tcPr)
    tc.append(_para([_run(text, sz=15, bold=True)]))
    return tc


def build_header_table():
    """2-row header: Row0 = RRL(vertical)|Features(vertical, span5);  Row1 = Studies/System | 5 rotated cols."""
    tbl = OxmlElement('w:tbl')
    tbl.append(_tblPr())
    tbl.append(_tblGrid())

    # Row 0 ─ "RRL" (vertical) | "Features" (vertical, spans cols 1-5)
    tr0 = OxmlElement('w:tr')
    tr0.append(_trPr_height(ROW0_H))
    tr0.append(rotated_cell(COL_W[0], 'RRL'))
    feature_width = sum(COL_W[1:])
    tr0.append(rotated_span_cell(feature_width, 5, 'Features'))
    tbl.append(tr0)

    # Row 1 ─ "Studies/System" | 5 rotated module names
    tr1 = OxmlElement('w:tr')
    tr1.append(_trPr_height(ROW1_H))
    tr1.append(normal_cell(COL_W[0], 'Studies/System', bold=True))
    for i, mod in enumerate(MODULES):
        tr1.append(rotated_cell(COL_W[i + 1], mod))
    tbl.append(tr1)

    return tbl


def build_data_table():
    """14 data rows — one per study."""
    tbl = OxmlElement('w:tbl')
    tbl.append(_tblPr())
    tbl.append(_tblGrid())

    for study_name, marks in STUDIES:
        tr = OxmlElement('w:tr')
        tr.append(normal_cell(COL_W[0], study_name))
        for i, m in enumerate(marks):
            symbol = '✓' if m else '✗'
            tr.append(check_cell(COL_W[i + 1], symbol))
        tbl.append(tr)

    return tbl


# ─── Main ────────────────────────────────────────────────────────────────────

doc = Document(SRC)
body = doc.element.body
children = list(body)

old_tbl = doc.tables[0]._tbl
old_idx = children.index(old_tbl)
print(f"Old placeholder table at body[{old_idx}] — removing it.")
body.remove(old_tbl)

header_tbl = build_header_table()
data_tbl   = build_data_table()

# Insert header first, then data table immediately after
body.insert(old_idx,     header_tbl)
body.insert(old_idx + 1, data_tbl)

doc.save(DST)
print(f"Saved → {DST}")

# ── Verification ──
doc2 = Document(DST)
print(f"\nTables in v6: {len(doc2.tables)}")
for ti, t in enumerate(doc2.tables):
    print(f"\n  Table {ti + 1}: {len(t.rows)} rows × {len(t.columns)} cols")
    for ri in range(len(t.rows)):
        cells = [c.text.strip() for c in t.rows[ri].cells]
        line  = " | ".join(f"{c[:30]!r}" for c in cells)
        print(f"    Row[{ri:02d}]: {line}")

print("\nDone.")
