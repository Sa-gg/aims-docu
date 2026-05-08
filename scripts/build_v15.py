#!/usr/bin/env python3
"""build_v15.py

Builds papers/AIMS/revisions/cha1&2_v15.docx from v14.

Scope: Chapter 2 cleanup only.
- Replaces U+FFFD (�) characters introduced in v14 edits with safe ASCII equivalents.

This is a corrective follow-up to keep the Word output clean.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

SRC = Path("papers/AIMS/revisions/cha1&2_v14.docx")
DST = Path("papers/AIMS/revisions/cha1&2_v15.docx")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if DST.exists():
        raise FileExistsError(DST)

    doc = Document(SRC)

    ch2_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Chapter 2")
    ref_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "REFERENCES")

    replacements = {
        "K\ufffd12": "K-12",
        "teachers\ufffd capacity": "teachers' capacity",
    }

    fixed = 0
    for i in range(ch2_start, ref_start):
        t = doc.paragraphs[i].text
        if "\ufffd" not in t:
            continue
        new_t = t
        for old, new in replacements.items():
            new_t = new_t.replace(old, new)
        if new_t != t:
            # Replace paragraph text via runs to preserve formatting
            p = doc.paragraphs[i]
            if not p.runs:
                p.add_run(new_t)
            else:
                p.runs[0].text = new_t
                for r in p.runs[1:]:
                    r.text = ""
            fixed += 1

    # Safety: ensure no U+FFFD remains in Chapter 2
    for i in range(ch2_start, ref_start):
        if "\ufffd" in doc.paragraphs[i].text:
            raise RuntimeError("U+FFFD still present in Chapter 2 after attempted fixes")

    doc.save(DST)
    print(f"Saved: {DST} (fixed paragraphs: {fixed})")


if __name__ == "__main__":
    main()
